#!/usr/bin/env python3
"""GUI-facing control plane for the portable digital-office agent system."""

from __future__ import annotations

import argparse
import base64
import binascii
import copy
import http.server
import hashlib
import hmac
import ipaddress
import math
import datetime as dt
import importlib.util
import json
import mimetypes
import os
import re
import shutil
import subprocess
import sys
import urllib.error
import urllib.parse
import urllib.request
import uuid
import tarfile
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

try:
    import fcntl
    _HAS_FCNTL = True
except ImportError:
    _HAS_FCNTL = False


class JsonFileLock:
    """Cross-process advisory lock for JSON file writes."""

    def __init__(self, lock_path: Path):
        self.lock_path = lock_path
        self._fd: Any = None

    def acquire(self) -> None:
        self.lock_path.parent.mkdir(parents=True, exist_ok=True)
        self._fd = os.open(str(self.lock_path), os.O_RDWR | os.O_CREAT, 0o644)
        if _HAS_FCNTL:
            fcntl.flock(self._fd, fcntl.LOCK_EX)
        else:
            # Windows fallback: not ideal, but prevents same-process races
            import msvcrt
            msvcrt.locking(self._fd, msvcrt.LK_LOCK, 1)

    def release(self) -> None:
        if self._fd is not None:
            if _HAS_FCNTL:
                fcntl.flock(self._fd, fcntl.LOCK_UN)
            else:
                import msvcrt
                msvcrt.locking(self._fd, msvcrt.LK_UNLCK, 1)
            os.close(self._fd)
            self._fd = None

    def __enter__(self) -> "JsonFileLock":
        self.acquire()
        return self

    def __exit__(self, *args: Any) -> None:
        self.release()


def write_json_locked(path: Path, data: dict[str, Any]) -> None:
    lock_file = path.with_name(f".{path.name}.lock")
    with JsonFileLock(lock_file):
        path.parent.mkdir(parents=True, exist_ok=True)
        tmp = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
        try:
            with tmp.open("w", encoding="utf-8") as handle:
                json.dump(data, handle, ensure_ascii=False, indent=2, sort_keys=True)
                handle.write("\n")
            os.replace(tmp, path)
        finally:
            if tmp.exists():
                tmp.unlink()


def write_json_unlocked(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    try:
        with tmp.open("w", encoding="utf-8") as handle:
            json.dump(data, handle, ensure_ascii=False, indent=2, sort_keys=True)
            handle.write("\n")
        os.replace(tmp, path)
    finally:
        if tmp.exists():
            tmp.unlink()


TEXT_EXTS = {".md", ".txt", ".json", ".csv"}
DOCX_EXTS = {".docx"}
PDF_EXTS = {".pdf"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}
DEFAULT_RESTORE_MAX_MEMBERS = 100_000
DEFAULT_RESTORE_MAX_UNCOMPRESSED_BYTES = 20 * 1024 * 1024 * 1024
CUSTOM_AGENT_STATUSES = {"active", "inactive", "archived"}
AGENT_ADMIN_ROLES = {"owner", "admin", "enterprise_admin", "system_admin"}
LOOP_STAGES = ["context", "decide", "act", "evaluate"]
LOOP_STAGE_ALIASES = {
    "perceive": "context",
    "reason": "decide",
    "plan": "decide",
    "execute": "act",
    "reflect": "evaluate",
    "iterate": "evaluate",
}
LOOP_STAGE_CHOICES = sorted(set(LOOP_STAGES) | set(LOOP_STAGE_ALIASES))
LOOP_STATUS_BY_STAGE = {
    "context": "context_loading",
    "decide": "deciding",
    "act": "acting",
    "evaluate": "evaluating",
}
LOOP_CONTROL_DECISIONS = {"continue", "replan", "retry", "wait_human", "complete", "fail", "cancel", "budget_exhausted"}
ITERATION_DECISION_TO_STATUS = {
    "confirm": "confirmed_for_application",
    "tune": "needs_tuning",
    "pause": "paused_by_user",
    "reject": "rejected",
}
STATUS_LABELS = {
    "received": "接收需求",
    "in_progress": "正在推动需求",
    "completed": "已完成需求",
    "downloaded_deployed": "已下载部署",
    "pending_user_confirmation": "等待用户确认",
    "confirmed_for_activation": "已确认，等待注册部署",
    "needs_tuning": "通过对话微调需求",
    "paused_by_user": "暂不处理",
    "needs_more_info": "需要补充信息",
    "rejected": "暂不受理",
    "send_failed": "发送失败",
}
TASK_STATUSES = {
    "queued",
    "running",
    "waiting_approval",
    "waiting_human_judgment",
    "blocked",
    "completed",
    "failed",
    "cancelled",
}
WORKFLOW_CONTROL_ACTIONS = {"run", "pause", "stop", "resume"}
CANVAS_NODE_TYPES = {
    "start",
    "agent_task",
    "text_instruction",
    "file_ref",
    "folder_ref",
    "knowledge_scope",
    "approval_gate",
    "human_input",
    "output_artifact",
    "merge_summary",
    "condition",
    "parallel_group",
}
SIMPLE_CANVAS_NODE_TYPES = {"agent_task", "text_instruction", "file_ref", "folder_ref", "knowledge_scope", "approval_gate", "output_artifact"}
NODE_IO_TYPES = {
    "start": {"outputs": {"control"}},
    "text_instruction": {"outputs": {"instruction"}},
    "file_ref": {"outputs": {"file_reference"}},
    "folder_ref": {"outputs": {"folder_reference"}},
    "knowledge_scope": {"outputs": {"knowledge_scope"}},
    "human_input": {"outputs": {"human_input", "instruction", "file_reference"}},
    "agent_task": {"inputs": {"control", "instruction", "file_reference", "folder_reference", "knowledge_scope", "human_input", "artifact"}, "outputs": {"artifact", "report", "decision"}},
    "approval_gate": {"inputs": {"artifact", "proposal", "decision"}, "outputs": {"approved", "rejected"}},
    "output_artifact": {"inputs": {"artifact", "report", "approved", "human_input"}},
    "merge_summary": {"inputs": {"artifact", "report"}, "outputs": {"report", "artifact"}},
    "condition": {"inputs": {"decision", "artifact", "approved"}, "outputs": {"decision"}},
    "parallel_group": {"inputs": {"control", "instruction"}, "outputs": {"artifact", "report"}},
}
KNOWLEDGE_SPACE_TYPES = {"personal", "project", "team", "company", "workflow_artifacts", "shared_with_me"}
KNOWLEDGE_RESOURCE_TYPES = {"folder", "item"}
KNOWLEDGE_PERMISSIONS = {"read", "write", "manage"}
KNOWLEDGE_TARGET_TYPES = {"user", "role", "agent", "project", "workflow"}
APPROVAL_STATUSES = {"pending", "approved", "rejected", "cancelled", "expired"}
APPROVAL_DECISION_TO_STATUS = {
    "approve": "approved",
    "reject": "rejected",
    "cancel": "cancelled",
}
JUDGMENT_STATUSES = {"pending", "approved", "rejected", "needs_evidence", "cancelled", "expired"}
JUDGMENT_BLOCKING_STATUSES = {"pending", "needs_evidence"}
JUDGMENT_DECISION_TO_STATUS = {
    "approve": "approved",
    "reject": "rejected",
    "request_evidence": "needs_evidence",
    "revise_scope": "pending",
    "cancel": "cancelled",
}
PASSING_GATE_STATUSES = {"passed", "success", "completed", "approved"}
RULE_PROPOSAL_STATUSES = {"pending_user_confirmation", "approved", "needs_tuning", "rejected", "applied"}
RULE_PROPOSAL_DECISION_TO_STATUS = {
    "approve": "approved",
    "tune": "needs_tuning",
    "reject": "rejected",
}
COORDINATION_MODES = {
    "single_agent",
    "secretary_centralized",
    "sequential_specialist_chain",
    "parallel_expert_dag",
    "human_gated",
}
TENANT_ADMIN_ROLES = {"owner", "enterprise_admin"}
ROLE_ACTIONS = {
    "owner": {"*"},
    "enterprise_admin": {
        "approval.create",
        "approval.decide",
        "audit.read",
        "notification.read",
        "project.manage",
        "rule.manage",
        "knowledge.manage",
        "knowledge.read",
        "release.approve",
        "support.grant",
        "task.manage",
        "task.read",
        "workflow.cancel",
        "workflow.control",
        "workflow.edit",
        "workflow.read",
        "workflow.resume",
        "workflow.retry",
        "workflow.start",
        "workbench.read",
    },
    "project_manager": {
        "agent.delegate",
        "approval.create",
        "approval.decide",
        "knowledge.manage",
        "knowledge.read",
        "notification.read",
        "rule.manage",
        "task.manage",
        "task.read",
        "workflow.cancel",
        "workflow.control",
        "workflow.edit",
        "workflow.read",
        "workflow.resume",
        "workflow.retry",
        "workflow.start",
        "workbench.read",
    },
    "professional_reviewer": {
        "approval.decide",
        "knowledge.read",
        "notification.read",
        "regulated_output.approve",
        "task.read",
        "workflow.read",
        "workbench.read",
    },
    "member": {
        "agent.delegate",
        "approval.create",
        "knowledge.read",
        "notification.read",
        "task.read",
        "workflow.read",
        "workflow.start",
        "workbench.read",
    },
    "viewer": {"knowledge.read", "notification.read", "task.read", "workflow.read", "workbench.read"},
}
ONBOARDING_FIELDS = [
    "assistant_style",
    "address_style",
    "language",
    "initiative_level",
    "pushback_style",
    "approval_strictness",
    "memory_mode",
    "work_mode",
]
ONBOARDING_OUTPUTS = {
    "preferences_json": "settings/user-preferences.json",
    "preferences_markdown": "settings/user-preferences.md",
}


def maybe_reexec_venv() -> None:
    if os.environ.get("OFFICE_SYSTEM_NO_VENV"):
        return
    script_root = Path(__file__).resolve().parent.parent
    venv_python = script_root / ".venv" / "bin" / "python"
    if not venv_python.exists():
        return
    if Path(sys.executable).resolve() == venv_python.resolve():
        return
    os.execv(str(venv_python), [str(venv_python), __file__, *sys.argv[1:]])


maybe_reexec_venv()


def system_root() -> Path:
    configured = os.environ.get("DIGITAL_OFFICE_SYSTEM_HOME")
    if configured:
        return Path(configured).expanduser()
    return Path(__file__).resolve().parent.parent


def now_iso() -> str:
    return dt.datetime.now(dt.timezone.utc).isoformat()


def canonical_hash(value: Any) -> str:
    if isinstance(value, bytes):
        payload = value
    elif isinstance(value, str):
        payload = value.encode("utf-8")
    else:
        payload = json.dumps(value, ensure_ascii=False, sort_keys=True).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def slugify(value: str) -> str:
    value = value.strip().lower()
    value = re.sub(r"[^a-z0-9\u4e00-\u9fff._-]+", "-", value)
    value = value.strip("-._")
    return value or uuid.uuid4().hex[:8]


SAFE_COMPONENT_RE = re.compile(r"^[A-Za-z0-9_\u4e00-\u9fff][A-Za-z0-9._\-\u4e00-\u9fff]{0,127}$")
SAFE_CLAIM_RE = re.compile(r"^[^\x00-\x1f\x7f]{1,256}$")


def safe_component(value: str, label: str) -> str:
    value = str(value or "").strip()
    if not SAFE_COMPONENT_RE.fullmatch(value) or "/" in value or "\\" in value or value in {".", ".."}:
        print(f"office-system: invalid {label}: {value!r}", file=sys.stderr)
        raise SystemExit(2)
    return value


def safe_claim(value: str | None, label: str, required: bool = True) -> str:
    value = str(value or "").strip()
    if not value and not required:
        return ""
    if not SAFE_CLAIM_RE.fullmatch(value):
        print(f"office-system: invalid {label}: {value!r}", file=sys.stderr)
        raise SystemExit(2)
    return value


def registered_agent(root: Path, value: str) -> str:
    agent = safe_component(value, "agent id")
    registry = effective_agents_registry(root)
    if agent not in registry.get("agents", {}):
        print(f"office-system: unknown agent id: {agent}", file=sys.stderr)
        raise SystemExit(2)
    return agent


def read_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def custom_agents_path(root: Path) -> Path:
    return root / "settings" / "custom-agents.registry.json"


def empty_custom_agents_registry() -> dict[str, Any]:
    return {
        "version": "1.0.0",
        "kind": "digital-office-custom-agents",
        "updated_at": "",
        "agents": {},
        "tombstones": [],
    }


def load_custom_agents_registry(root: Path) -> dict[str, Any]:
    path = custom_agents_path(root)
    if not path.exists():
        return empty_custom_agents_registry()
    registry = read_json(path)
    if registry.get("kind") != "digital-office-custom-agents" or not isinstance(registry.get("agents"), dict):
        print("office-system: invalid custom Agent registry", file=sys.stderr)
        raise SystemExit(2)
    registry.setdefault("tombstones", [])
    return registry


def effective_agents_registry(root: Path, *, include_inactive: bool = False) -> dict[str, Any]:
    registry = copy.deepcopy(read_json(root / "agents.registry.json"))
    custom = load_custom_agents_registry(root)
    for agent_id, record in custom.get("agents", {}).items():
        status = str(record.get("status", "inactive"))
        if status != "active" and not include_inactive:
            continue
        config = record.get("config", {})
        if isinstance(config, dict):
            registry.setdefault("agents", {})[agent_id] = copy.deepcopy(config)
    return registry


def require_agent_admin(role: str) -> str:
    role = safe_component(role, "role")
    if role not in AGENT_ADMIN_ROLES:
        print("office-system: Agent lifecycle changes require an administrator role", file=sys.stderr)
        raise SystemExit(3)
    return role


def write_json(path: Path, data: dict[str, Any]) -> None:
    write_json_locked(path, data)


def append_jsonl(path: Path, event: dict[str, Any]) -> None:
    lock_file = path.with_name(f".{path.name}.lock")
    with JsonFileLock(lock_file):
        append_jsonl_unlocked(path, event)


def append_jsonl_unlocked(path: Path, event: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(event, ensure_ascii=False, sort_keys=True) + "\n")
        handle.flush()
        os.fsync(handle.fileno())


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def append_log(root: Path, event: dict[str, Any]) -> None:
    event = {"time": now_iso(), **event}
    log = root / "logs" / "office-system.jsonl"
    append_jsonl(log, event)


def onboarding_presets(root: Path) -> dict[str, Any]:
    return read_json(root / "onboarding.presets.json")


def onboarding_preferences_path(root: Path) -> Path:
    return root / ONBOARDING_OUTPUTS["preferences_json"]


def onboarding_preferences_markdown_path(root: Path) -> Path:
    return root / ONBOARDING_OUTPUTS["preferences_markdown"]


def current_onboarding_preferences(root: Path) -> dict[str, Any] | None:
    path = onboarding_preferences_path(root)
    if not path.exists():
        return None
    return read_json(path)


def validate_onboarding_choices(presets: dict[str, Any], choices: dict[str, str]) -> None:
    fields = presets.get("fields", {})
    for field in ONBOARDING_FIELDS:
        if field not in fields:
            print(f"office-system: onboarding presets missing field: {field}", file=sys.stderr)
            raise SystemExit(2)
        value = choices.get(field)
        if value not in (fields.get(field, {}).get("choices") or {}):
            print(f"office-system: invalid onboarding choice for {field}: {value!r}", file=sys.stderr)
            raise SystemExit(2)


def choice_summary(presets: dict[str, Any], field: str, value: str) -> dict[str, str]:
    option = presets["fields"][field]["choices"][value]
    return {
        "field": field,
        "value": value,
        "label": option.get("label", value),
        "description": option.get("description", ""),
    }


def render_onboarding_markdown(preferences: dict[str, Any]) -> str:
    choices = preferences.get("choices", {})
    summaries = preferences.get("choice_summaries", [])
    lines = [
        "# Digital Office User Preferences",
        "",
        "This file is generated by the GUI onboarding or settings flow. It is local runtime state and must not be committed to the public repository.",
        "",
        "## Identity",
        "",
        f"- Tenant: {preferences.get('tenant_id', '') or 'default'}",
        f"- User: {preferences.get('user_id', '') or 'default'}",
        f"- Company: {preferences.get('company_name', '') or 'not set'}",
        f"- Secretary display name: {preferences.get('secretary_name', '') or 'Digital Office Secretary'}",
        "",
        "## Selected Behavior",
        "",
    ]
    for item in summaries:
        lines.append(f"- {item['field']}: {item['label']} ({item['value']})")
        if item.get("description"):
            lines.append(f"  - {item['description']}")
    lines.extend(
        [
            "",
            "## Runtime Instruction",
            "",
            "- Treat these choices as user preferences, not immutable system law.",
            "- Safety, authorization, approval, knowledge-source priority, and production harness gates override persona preferences.",
            "- If the user changes preferences in the GUI, reload this file before continuing the conversation.",
            "- Do not infer personal preferences that were not selected or explicitly confirmed by the user.",
            "",
            "## Compact Summary",
            "",
            f"- assistant_style: {choices.get('assistant_style')}",
            f"- address_style: {choices.get('address_style')}",
            f"- language: {choices.get('language')}",
            f"- initiative_level: {choices.get('initiative_level')}",
            f"- pushback_style: {choices.get('pushback_style')}",
            f"- approval_strictness: {choices.get('approval_strictness')}",
            f"- memory_mode: {choices.get('memory_mode')}",
            f"- work_mode: {choices.get('work_mode')}",
        ]
    )
    tone_note = preferences.get("tone_note", "")
    if tone_note:
        lines.extend(["", "## User Tone Note", "", tone_note])
    return "\n".join(lines) + "\n"


def onboarding_options(args: argparse.Namespace) -> int:
    root = system_root()
    presets = onboarding_presets(root)
    current = current_onboarding_preferences(root)
    agents = agent_summaries(root)
    projects = project_summaries(root, 1000)
    payload = {
        "presets": presets,
        "current": current,
        "configured": current is not None,
        "outputs": ONBOARDING_OUTPUTS,
    }
    print(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def onboarding_status(args: argparse.Namespace) -> int:
    root = system_root()
    current = current_onboarding_preferences(root)
    payload = {
        "configured": current is not None,
        "preferences_json": str(onboarding_preferences_path(root)),
        "preferences_markdown": str(onboarding_preferences_markdown_path(root)),
        "preferences": current or {},
    }
    print(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True))
    return 0 if current is not None else 1


def onboarding_apply(args: argparse.Namespace) -> int:
    if not args.confirmed:
        print("office-system: preference changes require --confirmed after GUI user review", file=sys.stderr)
        return 2
    root = system_root()
    presets = onboarding_presets(root)
    defaults = presets.get("default_choices", {})
    current = current_onboarding_preferences(root) or {}
    current_choices = current.get("choices", {})
    choices = {
        field: getattr(args, field.replace("-", "_"), None) or current_choices.get(field) or defaults.get(field)
        for field in ONBOARDING_FIELDS
    }
    validate_onboarding_choices(presets, choices)
    preferences = {
        "version": "1.0.0",
        "kind": "digital-office-user-preferences",
        "configured": True,
        "configured_at": current.get("configured_at") or now_iso(),
        "updated_at": now_iso(),
        "tenant_id": safe_claim(args.tenant or current.get("tenant_id") or "default", "tenant", required=False),
        "user_id": safe_claim(args.user or current.get("user_id") or "default", "user", required=False),
        "company_name": safe_claim(args.company_name if args.company_name is not None else current.get("company_name", ""), "company name", required=False),
        "secretary_name": safe_claim(args.secretary_name or current.get("secretary_name") or "Digital Office Secretary", "secretary name"),
        "tone_note": safe_claim(args.tone_note if args.tone_note is not None else current.get("tone_note", ""), "tone note", required=False),
        "choices": choices,
        "choice_summaries": [choice_summary(presets, field, choices[field]) for field in ONBOARDING_FIELDS],
        "source": getattr(args, "preference_source", "gui_first_run_onboarding"),
        "outputs": ONBOARDING_OUTPUTS,
    }
    write_json(onboarding_preferences_path(root), preferences)
    write_text(onboarding_preferences_markdown_path(root), render_onboarding_markdown(preferences))
    append_log(root, {"event": "preference_update", "tenant": preferences["tenant_id"], "user": preferences["user_id"], "source": preferences["source"]})
    print(json.dumps(preferences, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def add_preferences_arguments(parser: argparse.ArgumentParser) -> None:
    parser.add_argument("--tenant")
    parser.add_argument("--user")
    parser.add_argument("--company-name")
    parser.add_argument("--secretary-name")
    parser.add_argument("--assistant-style")
    parser.add_argument("--address-style")
    parser.add_argument("--language")
    parser.add_argument("--initiative-level")
    parser.add_argument("--pushback-style")
    parser.add_argument("--approval-strictness")
    parser.add_argument("--memory-mode")
    parser.add_argument("--work-mode")
    parser.add_argument("--tone-note")
    parser.add_argument("--confirmed", action="store_true")


def read_last_jsonl(path: Path) -> dict[str, Any] | None:
    if not path.exists():
        return None
    last = ""
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            if line.strip():
                last = line
    if not last:
        return None
    try:
        return json.loads(last)
    except json.JSONDecodeError:
        return None


def append_audit_event(
    root: Path,
    event_type: str,
    *,
    actor_id: str = "",
    actor_role: str = "",
    tenant_id: str = "",
    deployment_id: str = "",
    project_id: str = "",
    agent_id: str = "",
    resource_type: str = "",
    resource_id: str = "",
    workflow_run_id: str = "",
    task_id: str = "",
    approval_id: str = "",
    outcome: str = "recorded",
    reason: str = "",
    extra: dict[str, Any] | None = None,
) -> dict[str, Any]:
    path = root / "logs" / "audit-events.jsonl"
    event = {
        "version": "1.0.0",
        "kind": "digital-office-audit-event",
        "event_id": f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}",
        "time": now_iso(),
        "event": event_type,
        "actor": {"user_id": actor_id, "role": actor_role},
        "tenant_id": tenant_id,
        "deployment_id": deployment_id,
        "project_id": project_id,
        "agent_id": agent_id,
        "resource": {"type": resource_type, "id": resource_id},
        "links": {
            "workflow_run_id": workflow_run_id,
            "task_id": task_id,
            "approval_id": approval_id,
        },
        "outcome": outcome,
        "reason": reason,
        "extra": extra or {},
        "previous_event_hash": "",
    }
    audit_lock = path.with_name(f".{path.name}.lock")
    with JsonFileLock(audit_lock):
        previous = read_last_jsonl(path)
        event["previous_event_hash"] = (previous or {}).get("event_hash", "")
        event["event_hash"] = canonical_hash(event)
        append_jsonl_unlocked(path, event)
    append_log(root, {"event": "audit_event", "audit_event": event_type, "event_id": event["event_id"], "outcome": outcome})
    return event


def notification_path(root: Path, notification_id: str) -> Path:
    return root / "notifications" / f"{safe_component(notification_id, 'notification id')}.json"


def emit_notification(
    root: Path,
    *,
    user_id: str,
    title: str,
    body: str,
    topic: str,
    resource_type: str,
    resource_id: str,
    severity: str = "info",
) -> dict[str, Any]:
    notification_id = f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    record = {
        "version": "1.0.0",
        "kind": "digital-office-notification",
        "notification_id": notification_id,
        "user_id": safe_claim(user_id, "notification user", required=False),
        "title": safe_claim(title, "notification title"),
        "body": body,
        "topic": safe_component(topic, "notification topic"),
        "resource_type": safe_component(resource_type, "notification resource type"),
        "resource_id": safe_claim(resource_id, "notification resource id", required=False),
        "severity": safe_component(severity, "notification severity"),
        "status": "unread",
        "created_at": now_iso(),
    }
    write_json(notification_path(root, notification_id), record)
    append_log(root, {"event": "notification_created", "notification_id": notification_id, "topic": topic, "user_id": record["user_id"]})
    return record


def read_records(directory: Path, kind: str | None = None) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    if not directory.exists():
        return records
    for path in sorted(directory.glob("*.json")):
        try:
            data = read_json(path)
        except Exception:
            continue
        if kind and data.get("kind") != kind:
            continue
        records.append(data)
    records.sort(key=lambda item: str(item.get("updated_at") or item.get("created_at") or ""), reverse=True)
    return records


def find_by_idempotency(root: Path, directory: str, key: str | None) -> dict[str, Any] | None:
    if not key:
        return None
    for record in read_records(root / directory):
        if record.get("idempotency_key") == key:
            return record
    return None


def task_path(root: Path, task_id: str) -> Path:
    return root / "tasks" / f"{safe_component(task_id, 'task id')}.json"


def approval_path(root: Path, approval_id: str) -> Path:
    return root / "approvals" / f"{safe_component(approval_id, 'approval id')}.json"


def judgment_path(root: Path, case_id: str) -> Path:
    return root / "judgments" / f"{safe_component(case_id, 'judgment case id')}.json"


def rule_proposal_path(root: Path, proposal_id: str) -> Path:
    return root / "rule-proposals" / f"{safe_component(proposal_id, 'rule proposal id')}.json"


def run_dir(root: Path, run_id: str) -> Path:
    return root / "runs" / safe_component(run_id, "run id")


def run_ledger_path(root: Path, run_id: str) -> Path:
    return run_dir(root, run_id) / "ledger.jsonl"


def checkpoint_dir(root: Path, run_id: str) -> Path:
    return run_dir(root, run_id) / "checkpoints"


def checkpoint_path(root: Path, run_id: str, checkpoint_id: str) -> Path:
    return checkpoint_dir(root, run_id) / f"{safe_component(checkpoint_id, 'checkpoint id')}.json"


def handoff_dir(root: Path, run_id: str) -> Path:
    return run_dir(root, run_id) / "handoffs"


def handoff_path(root: Path, run_id: str, handoff_id: str) -> Path:
    return handoff_dir(root, run_id) / f"{safe_component(handoff_id, 'handoff id')}.json"


def load_judgment_case(root: Path, case_id: str) -> dict[str, Any]:
    path = judgment_path(root, case_id)
    if not path.exists():
        print(f"office-system: judgment case not found: {case_id}", file=sys.stderr)
        raise SystemExit(2)
    return read_json(path)


def load_rule_proposal(root: Path, proposal_id: str) -> dict[str, Any]:
    path = rule_proposal_path(root, proposal_id)
    if not path.exists():
        print(f"office-system: rule proposal not found: {proposal_id}", file=sys.stderr)
        raise SystemExit(2)
    return read_json(path)


def load_task(root: Path, task_id: str) -> dict[str, Any]:
    path = task_path(root, task_id)
    if not path.exists():
        print(f"office-system: task not found: {task_id}", file=sys.stderr)
        raise SystemExit(2)
    return read_json(path)


def load_approval(root: Path, approval_id: str) -> dict[str, Any]:
    path = approval_path(root, approval_id)
    if not path.exists():
        print(f"office-system: approval not found: {approval_id}", file=sys.stderr)
        raise SystemExit(2)
    return read_json(path)


def update_task_status(root: Path, task_id: str, status: str, *, message: str = "", actor_id: str = "", actor_role: str = "") -> dict[str, Any]:
    status = safe_component(status, "task status")
    if status not in TASK_STATUSES:
        print(f"office-system: invalid task status: {status}", file=sys.stderr)
        raise SystemExit(2)
    task = load_task(root, task_id)
    task["status"] = status
    task["updated_at"] = now_iso()
    task.setdefault("history", []).append({"time": task["updated_at"], "status": status, "message": message, "actor_id": actor_id, "actor_role": actor_role})
    write_json(task_path(root, task_id), task)
    return task


def project_path(root: Path, project: str) -> Path:
    return root / "projects" / safe_component(project, "project id")


def ensure_project(root: Path, project: str) -> Path:
    path = project_path(root, project)
    if not (path / "project.json").exists():
        print(f"office-system: project not found: {project}", file=sys.stderr)
        raise SystemExit(2)
    return path


def load_project_record(root: Path, project: str) -> dict[str, Any] | None:
    if not project:
        return None
    project_dir = ensure_project(root, project)
    return read_json(project_dir / "project.json")


def project_member(project_data: dict[str, Any] | None, user_id: str) -> dict[str, Any] | None:
    if not project_data:
        return None
    for member in project_data.get("human_members", []) or []:
        if member.get("user_id") == user_id and member.get("status", "active") == "active":
            return member
    return None


def role_allows_action(role: str, action: str) -> bool:
    allowed = ROLE_ACTIONS.get(role, set())
    return "*" in allowed or action in allowed


def compute_authorization_decision(
    root: Path,
    *,
    tenant_id: str,
    deployment_id: str,
    user_id: str,
    user_role: str,
    action: str,
    resource_type: str,
    resource_id: str,
    project_id: str = "",
    agent_id: str = "",
    workflow_run_id: str = "",
    reason: str = "",
    audit: bool = True,
) -> dict[str, Any]:
    tenant_id = safe_claim(tenant_id, "tenant id")
    deployment_id = safe_claim(deployment_id, "deployment id")
    user_id = safe_claim(user_id, "user id")
    user_role = safe_component(user_role, "user role")
    action = safe_claim(action, "action")
    resource_type = safe_component(resource_type, "resource type")
    resource_id = safe_claim(resource_id, "resource id")
    project_id = safe_component(project_id, "project id") if project_id else ""
    agent_id = registered_agent(root, agent_id) if agent_id else ""
    workflow_run_id = safe_claim(workflow_run_id, "workflow run id", required=False)

    reasons: list[str] = []
    project_data = load_project_record(root, project_id) if project_id else None
    if not role_allows_action(user_role, action):
        reasons.append(f"role {user_role} cannot perform {action}")

    members = (project_data or {}).get("human_members", []) if project_data else []
    if project_id and members and user_role not in TENANT_ADMIN_ROLES:
        member = project_member(project_data, user_id)
        if not member:
            reasons.append("user is not an active project member")
        elif member.get("role") and member.get("role") != user_role:
            reasons.append("user role claim does not match project membership")

    if action in {"workflow.start", "agent.delegate"} and project_data and agent_id:
        roster = project_data.get("agent_roster", []) or []
        global_intake_agent = action == "workflow.start" and agent_id == "secretary"
        if agent_id not in roster and not global_intake_agent:
            reasons.append("agent is not on the project roster")

    if action == "regulated_output.approve" and user_role != "professional_reviewer":
        reasons.append("regulated outputs require professional_reviewer approval")

    allowed = not reasons
    decision = {
        "version": "1.0.0",
        "kind": "digital-office-authorization-decision",
        "decision_id": f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}",
        "allowed": allowed,
        "outcome": "allow" if allowed else "deny",
        "reasons": reasons,
        "tenant_id": tenant_id,
        "deployment_id": deployment_id,
        "user_id": user_id,
        "user_role": user_role,
        "action": action,
        "resource_type": resource_type,
        "resource_id": resource_id,
        "project_id": project_id,
        "agent_id": agent_id,
        "workflow_run_id": workflow_run_id,
        "reason": reason,
        "created_at": now_iso(),
    }
    if audit:
        event = append_audit_event(
            root,
            "authorization_decision",
            actor_id=user_id,
            actor_role=user_role,
            tenant_id=tenant_id,
            deployment_id=deployment_id,
            project_id=project_id,
            agent_id=agent_id,
            resource_type=resource_type,
            resource_id=resource_id,
            workflow_run_id=workflow_run_id,
            outcome=decision["outcome"],
            reason="; ".join(reasons) if reasons else reason,
            extra={"action": action, "decision_id": decision["decision_id"]},
        )
        decision["audit_event_id"] = event["event_id"]
    return decision


PROJECT_CONTEXT_FIELDS = {
    "intent_summary", "goal", "deliverables", "acceptance_criteria", "constraints", "source_refs",
    "deadline", "stakeholders", "risk_level", "open_questions", "assumptions",
}


def fresh_project_context(brief: str = "", *, required: bool = False) -> dict[str, Any]:
    return {
        "version": 1,
        "required": required,
        "intent_summary": brief.strip(),
        "intent_confirmed_hash": "",
        "intent_confirmed_at": "",
        "intent_confirmed_by": "",
        "goal": brief.strip(),
        "deliverables": [],
        "acceptance_criteria": [],
        "constraints": [],
        "source_refs": [],
        "deadline": "",
        "stakeholders": [],
        "risk_level": "normal",
        "open_questions": [],
        "assumptions": [],
        "confirmed_hash": "",
        "confirmed_at": "",
        "confirmed_by": "",
        "updated_at": now_iso(),
    }


def normalized_project_context(value: Any) -> dict[str, Any]:
    source = value if isinstance(value, dict) else {}
    result = fresh_project_context(str(source.get("goal", "")), required=bool(source.get("required", False)))
    result["version"] = max(1, int(source.get("version", 1) or 1))
    result["intent_summary"] = str(source.get("intent_summary", source.get("goal", ""))).strip()[:4000]
    for field in ("deliverables", "acceptance_criteria", "constraints", "source_refs", "stakeholders", "assumptions"):
        raw = source.get(field, [])
        result[field] = [str(item).strip()[:1000] for item in raw[:50] if str(item).strip()] if isinstance(raw, list) else []
    questions = source.get("open_questions", [])
    if isinstance(questions, list):
        result["open_questions"] = [
            {"question": str(item.get("question", "")).strip()[:1000], "critical": bool(item.get("critical", False))}
            if isinstance(item, dict) else {"question": str(item).strip()[:1000], "critical": False}
            for item in questions[:30] if (str(item.get("question", "")).strip() if isinstance(item, dict) else str(item).strip())
        ]
    result["deadline"] = str(source.get("deadline", "")).strip()[:200]
    result["risk_level"] = str(source.get("risk_level", "normal")).strip() if str(source.get("risk_level", "normal")).strip() in {"low", "normal", "high", "regulated"} else "normal"
    for field in ("intent_confirmed_hash", "intent_confirmed_at", "intent_confirmed_by", "confirmed_hash", "confirmed_at", "confirmed_by", "updated_at"):
        result[field] = str(source.get(field, ""))
    return result


def project_context_hash(context: dict[str, Any]) -> str:
    material = {key: context.get(key) for key in sorted(PROJECT_CONTEXT_FIELDS)}
    return canonical_hash(material)


def project_intent_hash(context: dict[str, Any]) -> str:
    return canonical_hash({"intent_summary": context.get("intent_summary", ""), "goal": context.get("goal", "")})


def assess_project_context(root: Path, project_id: str, project: dict[str, Any]) -> dict[str, Any]:
    context = normalized_project_context(project.get("context_intake"))
    knowledge_count = len(load_entries(root, project_path(root, project_id) / "knowledge" / "entries"))
    points = {
        "goal": 25 if context["goal"] else 0,
        "deliverables": 20 if context["deliverables"] else 0,
        "acceptance_criteria": 20 if context["acceptance_criteria"] else 0,
        "constraints": 10 if context["constraints"] else 0,
        "sources": 10 if context["source_refs"] or knowledge_count else 0,
        "deadline": 5 if context["deadline"] else 0,
        "stakeholders": 5 if context["stakeholders"] else 0,
        "risk": 5 if context["risk_level"] else 0,
    }
    blockers = []
    intent_hash = project_intent_hash(context)
    intent_confirmed = bool(context["intent_summary"] and context["intent_confirmed_hash"] == intent_hash)
    if not intent_confirmed:
        blockers.append("intent_confirmation_required")
    for field in ("goal", "deliverables", "acceptance_criteria"):
        if not context[field]:
            blockers.append(f"missing_{field}")
    if any(item.get("critical") for item in context["open_questions"]):
        blockers.append("critical_questions_unresolved")
    context_hash = project_context_hash(context)
    threshold = 70
    confirmed = bool(context["confirmed_hash"] and context["confirmed_hash"] == context_hash and sum(points.values()) >= threshold and not blockers)
    suggestions = []
    suggestion_specs = [
        ("goal", "先确认项目最终要解决什么问题？", "目标越清楚，秘书越不容易把任务分错方向。"),
        ("deliverables", "最后希望拿到哪些具体成果？", "明确交付物可以让 Agent 在正确的位置停止。"),
        ("acceptance_criteria", "什么样的结果才算完成并且可用？", "可检查的标准能防止 Loop 反复改写却没有进展。"),
        ("constraints", "有哪些预算、格式、合规或不能触碰的边界？", "提前说明边界可以减少返工和高风险操作。"),
        ("sources", "有哪些现成资料、链接或历史文件可以作为依据？", "原始资料能降低猜测和上下文漂移。"),
    ]
    for field, prompt, why in suggestion_specs:
        empty = not (context["source_refs"] or knowledge_count) if field == "sources" else not context.get(field)
        if empty:
            suggestions.append({"field": field, "prompt": prompt, "why": why, "priority": "required" if field in {"goal", "deliverables", "acceptance_criteria"} else "recommended"})
    depth_questions = [
        {"field": "acceptance_criteria", "prompt": "如果项目只能带来一个可观察的变化，那个变化是什么，谁来判断它已经发生？", "why": "从真实结果而不是任务名称出发，可以校准目标和验收标准。", "priority": "socratic"},
        {"field": "constraints", "prompt": "哪一种结果即使看起来完成了，你也会认为它是失败的？为什么？", "why": "反例能暴露不能牺牲的质量、合规和业务边界。", "priority": "socratic"},
        {"field": "source_refs", "prompt": "哪些原始文件、历史决定或数据最能证明事实？哪些内容目前只是猜测？", "why": "区分证据和假设，是防止 Loop 漂移的关键。", "priority": "socratic"},
        {"field": "deliverables", "prompt": "最终成果会被谁在什么场景下使用，需要以什么形式交付？", "why": "从使用场景倒推交付物，比泛泛描述更可执行。", "priority": "socratic"},
        {"field": "open_questions", "prompt": "现在最不确定、但一旦判断错误就会让项目走偏的三件事是什么？", "why": "优先消除高影响不确定性，可以少走无效循环。", "priority": "socratic"},
    ]
    seen = {(item["field"], item["prompt"]) for item in suggestions}
    for item in depth_questions:
        if len(suggestions) >= 5:
            break
        if (item["field"], item["prompt"]) not in seen:
            suggestions.append(item)
    return {
        "required": context["required"],
        "readiness_score": sum(points.values()),
        "readiness_threshold": threshold,
        "ready": sum(points.values()) >= threshold and not blockers,
        "confirmed": confirmed,
        "intent": {
            "summary": context["intent_summary"],
            "hash": intent_hash,
            "confirmed": intent_confirmed,
            "confirmed_at": context["intent_confirmed_at"],
            "confirmed_by": context["intent_confirmed_by"],
        },
        "context_hash": context_hash,
        "context_version": context["version"],
        "knowledge_count": knowledge_count,
        "blockers": blockers,
        "question_policy": {"minimum_questions": 3, "recommended_questions": 5, "method": "first_principles_socratic"},
        "suggestions": suggestions[:5],
        "context": context,
    }


def copy_template_project(root: Path, project_id: str, name: str, agents: list[str], schedule: str, *, guided_intake: bool = False, brief: str = "") -> Path:
    template = root / "projects" / "_template"
    project_id = safe_component(project_id, "project id")
    target = project_path(root, project_id)
    if target.exists():
        print(f"office-system: project already exists: {project_id}", file=sys.stderr)
        raise SystemExit(2)
    shutil.copytree(template, target)
    data = read_json(target / "project.json")
    data.update(
        {
            "project_id": project_id,
            "name": name,
            "status": "active",
            "created_at": now_iso(),
        }
    )
    if agents:
        data["agent_roster"] = agents
    data["context_intake"] = fresh_project_context(brief, required=guided_intake)
    data.setdefault("methodology_promotion", {})["schedule"] = schedule
    write_json(target / "project.json", data)
    return target


def project_context_status(args: argparse.Namespace) -> int:
    root = system_root()
    project_id = safe_component(args.project, "project id")
    project_file = ensure_project(root, project_id) / "project.json"
    project = read_json(project_file)
    assessment = assess_project_context(root, project_id, project)
    print(json.dumps({"status": "ready" if assessment["confirmed"] else "needs_context", "project_id": project_id, **assessment}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def project_context_update(args: argparse.Namespace) -> int:
    root = system_root()
    project_id = safe_component(args.project, "project id")
    project_file = ensure_project(root, project_id) / "project.json"
    project = read_json(project_file)
    try:
        patch = json.loads(args.context_json)
    except json.JSONDecodeError as exc:
        print(f"office-system: invalid project context JSON: {exc}", file=sys.stderr)
        return 2
    if not isinstance(patch, dict) or any(key not in PROJECT_CONTEXT_FIELDS for key in patch):
        print("office-system: project context contains unsupported fields", file=sys.stderr)
        return 2
    current = normalized_project_context(project.get("context_intake"))
    before_hash = project_context_hash(current)
    before_intent_hash = project_intent_hash(current)
    for key, value in patch.items():
        current[key] = value
    updated = normalized_project_context(current)
    after_hash = project_context_hash(updated)
    if after_hash != before_hash:
        updated["version"] = current["version"] + 1
        updated["confirmed_hash"] = ""
        updated["confirmed_at"] = ""
        updated["confirmed_by"] = ""
        if project_intent_hash(updated) != before_intent_hash:
            updated["intent_confirmed_hash"] = ""
            updated["intent_confirmed_at"] = ""
            updated["intent_confirmed_by"] = ""
    updated["required"] = bool(current.get("required", True))
    updated["updated_at"] = now_iso()
    project["context_intake"] = updated
    project["updated_at"] = now_iso()
    write_json(project_file, project)
    assessment = assess_project_context(root, project_id, project)
    append_log(root, {"event": "project_context_updated", "project": project_id, "version": updated["version"], "updated_by": args.updated_by or ""})
    print(json.dumps({"status": "updated", "project_id": project_id, **assessment}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def project_intent_confirm(args: argparse.Namespace) -> int:
    if not args.confirmed:
        print("office-system: project-intent-confirm requires --confirmed", file=sys.stderr)
        return 2
    root = system_root()
    project_id = safe_component(args.project, "project id")
    project_file = ensure_project(root, project_id) / "project.json"
    project = read_json(project_file)
    context = normalized_project_context(project.get("context_intake"))
    actual_hash = project_intent_hash(context)
    if not context["intent_summary"]:
        print("office-system: intent summary is required before confirmation", file=sys.stderr)
        return 2
    if args.expected_hash and args.expected_hash != actual_hash:
        print(json.dumps({"status": "conflict", "reason": "intent_changed", "expected_hash": args.expected_hash, "actual_hash": actual_hash}, ensure_ascii=False, indent=2))
        return 4
    context["intent_confirmed_hash"] = actual_hash
    context["intent_confirmed_at"] = now_iso()
    context["intent_confirmed_by"] = safe_claim(args.confirmed_by, "confirmed by")
    context["confirmed_hash"] = ""
    context["confirmed_at"] = ""
    context["confirmed_by"] = ""
    project["context_intake"] = context
    project["updated_at"] = now_iso()
    write_json(project_file, project)
    assessment = assess_project_context(root, project_id, project)
    append_log(root, {"event": "project_intent_confirmed", "project": project_id, "intent_hash": actual_hash, "confirmed_by": context["intent_confirmed_by"]})
    print(json.dumps({"status": "intent_confirmed", "project_id": project_id, **assessment}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def project_context_confirm(args: argparse.Namespace) -> int:
    if not args.confirmed:
        print("office-system: project-context-confirm requires --confirmed", file=sys.stderr)
        return 2
    root = system_root()
    project_id = safe_component(args.project, "project id")
    project_file = ensure_project(root, project_id) / "project.json"
    project = read_json(project_file)
    assessment = assess_project_context(root, project_id, project)
    if not assessment["ready"]:
        print(json.dumps({"status": "needs_context", "project_id": project_id, **assessment}, ensure_ascii=False, indent=2, sort_keys=True))
        return 4
    context = normalized_project_context(project.get("context_intake"))
    context["confirmed_hash"] = assessment["context_hash"]
    context["confirmed_at"] = now_iso()
    context["confirmed_by"] = safe_claim(args.confirmed_by, "confirmed by")
    project["context_intake"] = context
    project["updated_at"] = now_iso()
    write_json(project_file, project)
    assessment = assess_project_context(root, project_id, project)
    append_log(root, {"event": "project_context_confirmed", "project": project_id, "version": context["version"], "confirmed_by": context["confirmed_by"]})
    print(json.dumps({"status": "confirmed", "project_id": project_id, **assessment}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def entry_root(root: Path, scope: str, project: str | None) -> Path:
    if scope == "company":
        return root / "knowledge" / "company" / "entries"
    if not project:
        print("office-system: --project is required for project knowledge", file=sys.stderr)
        raise SystemExit(2)
    return ensure_project(root, project) / "knowledge" / "entries"


def classify_file(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in TEXT_EXTS:
        return "text"
    if ext in DOCX_EXTS:
        return "word"
    if ext in PDF_EXTS:
        return "pdf"
    if ext in IMAGE_EXTS:
        return "image"
    return "binary"


def extract_text_file(source: Path, target: Path) -> tuple[str, str]:
    try:
        text = source.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = source.read_text(encoding="utf-8", errors="replace")
    write_text(target, text)
    return "local_extracted", "python_stdlib"


def extract_docx(source: Path, target: Path) -> tuple[str, str]:
    try:
        import docx  # type: ignore

        document = docx.Document(str(source))
        text_parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))
        write_text(target, "\n\n".join(text_parts).strip() + "\n")
        return "local_extracted", "python_docx"
    except Exception:
        pass

    paragraphs: list[str] = []
    try:
        with zipfile.ZipFile(source) as archive:
            xml = archive.read("word/document.xml")
        tree = ElementTree.fromstring(xml)
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for para in tree.iter(f"{namespace}p"):
            chunks = [node.text for node in para.iter(f"{namespace}t") if node.text]
            if chunks:
                paragraphs.append("".join(chunks))
    except Exception as exc:
        write_text(target.with_name("extraction-error.txt"), f"docx extraction failed: {exc}\n")
        return "pending_extraction", "python_stdlib_zip_xml"
    write_text(target, "\n\n".join(paragraphs).strip() + "\n")
    return "local_extracted", "python_stdlib_zip_xml"


def extract_pdf(source: Path, target: Path) -> tuple[str, str]:
    tool = shutil.which("pdftotext")
    if tool:
        proc = subprocess.run([tool, str(source), "-"], text=True, capture_output=True)
        if proc.returncode == 0:
            write_text(target, proc.stdout)
            return "local_extracted", "poppler_pdftotext"

    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(source))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            pages.append(f"\n\n--- page {index} ---\n\n{page.extract_text() or ''}")
        write_text(target, "\n".join(pages).strip() + "\n")
        return "local_extracted", "pypdf"
    except Exception as exc:
        write_text(target.with_name("README.md"), f"PDF text extraction pending or failed: {exc}\n")
        return "pending_local_model_install", "poppler_pdftotext_or_pypdf"


def extract_image_rapidocr(source: Path, target: Path) -> tuple[str, str] | None:
    try:
        from rapidocr_onnxruntime import RapidOCR  # type: ignore
    except Exception:
        return None
    try:
        engine = RapidOCR()
        result, _ = engine(str(source))
        lines = []
        for item in result or []:
            text = item[1] if len(item) > 1 else ""
            score = item[2] if len(item) > 2 else ""
            if text:
                lines.append(f"{text}\t{score}")
        write_text(target, "\n".join(lines).strip() + "\n")
        return "local_extracted", "rapidocr_onnxruntime"
    except Exception as exc:
        write_text(target.with_name("rapidocr-error.txt"), f"rapidocr failed: {exc}\n")
        return None


def extract_image_ocr(source: Path, target: Path) -> tuple[str, str]:
    rapid = extract_image_rapidocr(source, target)
    if rapid:
        return rapid

    tool = shutil.which("tesseract")
    if not tool:
        write_text(
            target.with_name("README.md"),
            "Image OCR pending. Install base-ocr-python for RapidOCR or base-ocr for Tesseract OCR.\n",
        )
        return "pending_local_model_install", "rapidocr_or_tesseract"
    proc = subprocess.run([tool, str(source), "stdout", "-l", "eng+chi_sim"], text=True, capture_output=True)
    if proc.returncode != 0:
        write_text(target.with_name("ocr-error.txt"), proc.stderr or "tesseract failed\n")
        return "pending_extraction", "tesseract_cli"
    write_text(target, proc.stdout)
    return "local_extracted", "tesseract_cli"


def run_extraction(source: Path, extracted_dir: Path, kind: str) -> tuple[str, str]:
    extracted_dir.mkdir(parents=True, exist_ok=True)
    target = extracted_dir / "text.md"
    if kind == "text":
        return extract_text_file(source, target)
    if kind == "word" and source.suffix.lower() == ".docx":
        return extract_docx(source, target)
    if kind == "pdf":
        return extract_pdf(source, target)
    if kind == "image":
        return extract_image_ocr(source, extracted_dir / "ocr.txt")
    write_text(extracted_dir / "README.md", f"No local extractor registered for {source.suffix}.\n")
    return "pending_extraction", "none"


def add_file_entry(args: argparse.Namespace) -> int:
    root = system_root()
    source = Path(args.file).expanduser().resolve()
    if not source.exists():
        print(f"office-system: file not found: {source}", file=sys.stderr)
        return 2
    title = args.title or source.stem
    entry_id = f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{slugify(title)}"
    base = entry_root(root, args.scope, args.project) / entry_id
    source_dir = base / "source"
    extracted_dir = base / "extracted"
    source_dir.mkdir(parents=True, exist_ok=True)
    copied = source_dir / source.name
    shutil.copy2(source, copied)

    kind = args.kind or classify_file(source)
    status, extractor = run_extraction(copied, extracted_dir, kind)
    review_status = "pending_review" if status == "local_extracted" else status
    entry = {
        "version": "1.0.0",
        "entry_id": entry_id,
        "scope": args.scope,
        "project_id": args.project,
        "title": title,
        "kind": kind,
        "mime_type": mimetypes.guess_type(source.name)[0],
        "source_file": str(copied.relative_to(root)),
        "extracted_dir": str(extracted_dir.relative_to(root)),
        "status": review_status,
        "extractor": extractor,
        "created_at": now_iso(),
        "agent_readable": args.approve,
        "notes": args.notes or "",
    }
    if args.approve:
        entry["status"] = "approved_for_agent_use"
    write_json(base / "entry.json", entry)
    append_log(root, {"event": "knowledge_add_file", "entry_id": entry_id, "scope": args.scope, "project": args.project})
    print(json.dumps(entry, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def add_text_entry(args: argparse.Namespace) -> int:
    root = system_root()
    body = args.body if args.body is not None else sys.stdin.read()
    title = args.title
    entry_id = f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{slugify(title)}"
    base = entry_root(root, args.scope, args.project) / entry_id
    write_text(base / "source" / "entry.md", body.rstrip() + "\n")
    write_text(base / "extracted" / "text.md", body.rstrip() + "\n")
    entry = {
        "version": "1.0.0",
        "entry_id": entry_id,
        "scope": args.scope,
        "project_id": args.project,
        "title": title,
        "kind": "text",
        "source_file": str((base / "source" / "entry.md").relative_to(root)),
        "extracted_dir": str((base / "extracted").relative_to(root)),
        "status": "approved_for_agent_use" if args.approve else "pending_review",
        "extractor": "gui_text_entry",
        "created_at": now_iso(),
        "agent_readable": args.approve,
    }
    write_json(base / "entry.json", entry)
    append_log(root, {"event": "knowledge_add_text", "entry_id": entry_id, "scope": args.scope, "project": args.project})
    print(json.dumps(entry, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def write_rule_to_store(root: Path, *, scope: str, title: str, body: str, project: str = "", agent: str = "") -> Path:
    filename = f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{slugify(title)}.md"
    if scope == "global":
        target = root / "rules" / "global" / filename
    elif scope == "agent":
        if not agent:
            print("office-system: --agent is required for agent rule", file=sys.stderr)
            raise SystemExit(2)
        agent = registered_agent(root, agent)
        target = root / "rules" / "agents" / f"{agent}.md"
        section = f"\n\n## {title}\n\n{body.rstrip()}\n"
        if target.exists():
            target.write_text(target.read_text(encoding="utf-8") + section, encoding="utf-8")
            return target
        write_text(target, f"# Agent Rules: {agent}\n{section}")
        return target
    else:
        if not project:
            print("office-system: --project is required for project rule", file=sys.stderr)
            raise SystemExit(2)
        target = ensure_project(root, project) / "rules" / filename
    text = f"# {title}\n\n{body.rstrip()}\n"
    write_text(target, text)
    return target


def add_rule(args: argparse.Namespace) -> int:
    root = system_root()
    body = args.body if args.body is not None else sys.stdin.read()
    target = write_rule_to_store(root, scope=args.scope, title=args.title, body=body, project=args.project or "", agent=args.agent or "")
    append_log(root, {"event": "rule_add", "scope": args.scope, "target": str(target.relative_to(root))})
    print(str(target))
    return 0


def collect_rule_text(root: Path, project: str = "", agent: str = "") -> str:
    parts: list[str] = []
    for path in sorted((root / "rules" / "global").glob("*.md")):
        parts.append(path.read_text(encoding="utf-8", errors="replace"))
    if project:
        for path in sorted((project_path(root, project) / "rules").glob("*.md")):
            parts.append(path.read_text(encoding="utf-8", errors="replace"))
    if agent:
        path = root / "rules" / "agents" / f"{safe_component(agent, 'agent id')}.md"
        if path.exists():
            parts.append(path.read_text(encoding="utf-8", errors="replace"))
    return "\n\n".join(parts).lower()


def rule_target_files(root: Path, scope: str, project: str = "", agent: str = "") -> list[Path]:
    if scope == "global":
        return sorted((root / "rules" / "global").glob("*.md"))
    if scope == "project" and project:
        return sorted((project_path(root, project) / "rules").glob("*.md"))
    if scope == "agent" and agent:
        path = root / "rules" / "agents" / f"{safe_component(agent, 'agent id')}.md"
        return [path] if path.exists() else []
    return []


def rule_conflict_domains(text: str) -> list[str]:
    text = text.lower()
    domains = {
        "external_delivery": ["customer", "external", "publish", "send", "客户", "外部", "发布", "发送"],
        "data_boundary": ["data", "privacy", "export", "secret", "数据", "隐私", "导出", "机密"],
        "evidence_standard": ["evidence", "citation", "source", "proof", "证据", "引用", "来源"],
        "memory_promotion": ["memory", "knowledge", "global rule", "记忆", "知识", "全局规则"],
        "regulated_work": ["legal", "medical", "tax", "investment", "法律", "医疗", "税务", "投资"],
        "agent_style": ["style", "tone", "voice", "风格", "语气"],
    }
    return [name for name, terms in domains.items() if any(term in text for term in terms)]


def has_negative_modal(text: str) -> bool:
    text = text.lower()
    return any(term in text for term in ["must not", "never", "forbid", "forbidden", "do not", "禁止", "不得", "不能", "不要"])


def has_positive_modal(text: str) -> bool:
    text = text.lower()
    return any(term in text for term in ["must", "always", "require", "required", "should", "必须", "总是", "需要", "应该"])


def rule_governance_report(
    root: Path,
    *,
    title: str,
    body: str,
    scope: str,
    project: str = "",
    agent: str = "",
    source: str = "",
    created_by: str = "",
    scope_confidence: str = "",
) -> dict[str, Any]:
    text = f"{title}\n{body}".strip()
    normalized = text.lower()
    new_domains = set(rule_conflict_domains(text))
    conflicts: list[dict[str, Any]] = []
    duplicates: list[dict[str, Any]] = []
    for path in rule_target_files(root, scope, project, agent):
        existing = path.read_text(encoding="utf-8", errors="replace")
        existing_lower = existing.lower()
        rel = str(path.relative_to(root))
        if slugify(title) in path.stem or title.lower() in existing_lower:
            duplicates.append({"path": rel, "reason": "same or very similar rule title already exists"})
        existing_domains = set(rule_conflict_domains(existing))
        shared_domains = sorted(new_domains & existing_domains)
        if shared_domains and has_negative_modal(normalized) != has_negative_modal(existing_lower) and (has_positive_modal(normalized) or has_positive_modal(existing_lower)):
            conflicts.append({"path": rel, "reason": "opposite modal language in same rule domain", "domains": shared_domains})
    requirements = ["human_confirmation", "scope_confirmation", "provenance_recorded"]
    if scope_confidence in {"", "low"}:
        requirements.append("explicit_scope_confirmation")
    if conflicts:
        requirements.append("conflict_override_with_reason")
    return {
        "version": "1.0.0",
        "kind": "digital-office-rule-governance-report",
        "status": "blocked_by_conflict" if conflicts else "ready_for_confirmation",
        "scope": scope,
        "scope_confidence": scope_confidence,
        "project_id": project,
        "agent_id": agent,
        "body_hash": canonical_hash(body),
        "title_hash": canonical_hash(title),
        "domains": sorted(new_domains),
        "duplicates": duplicates,
        "conflicts": conflicts,
        "promotion_requirements": requirements,
        "provenance": {
            "source": safe_claim(source, "rule source", required=False),
            "created_by": safe_claim(created_by, "created by", required=False),
            "captured_at": now_iso(),
        },
    }


def rule_elicitation_topics() -> list[dict[str, Any]]:
    return [
        {
            "id": "human_judgment",
            "keywords": ["approval", "human", "judgment", "confirm", "批准", "人工", "确认"],
            "prompt": "Ask which actions should stop for human judgment before an Agent continues.",
        },
        {
            "id": "evidence_standard",
            "keywords": ["evidence", "citation", "source", "proof", "证据", "引用", "来源"],
            "prompt": "Ask what evidence, citations, or source quality are required before delivery.",
        },
        {
            "id": "role_boundary",
            "keywords": ["role", "handoff", "owner", "agent boundary", "角色", "交接", "边界"],
            "prompt": "Ask which Agent owns the decision and where handoff to another Agent is required.",
        },
        {
            "id": "quality_bar",
            "keywords": ["quality", "acceptance", "test", "review", "质量", "验收", "测试"],
            "prompt": "Ask how the user judges this work as production-ready rather than merely drafted.",
        },
        {
            "id": "data_boundary",
            "keywords": ["data", "privacy", "secret", "export", "数据", "隐私", "机密", "导出"],
            "prompt": "Ask what data can be used, shared, stored, exported, or promoted to knowledge.",
        },
        {
            "id": "agent_specific_style",
            "keywords": ["style", "tone", "designer", "writer", "lawyer", "coder", "风格", "语气", "设计", "写作", "律师", "工程"],
            "prompt": "Ask whether the preference belongs to one specialist Agent instead of every Agent.",
        },
    ]


def rule_elicit(args: argparse.Namespace) -> int:
    root = system_root()
    project = safe_component(args.project, "project id") if args.project else ""
    agent = registered_agent(root, args.agent) if args.agent else ""
    existing = collect_rule_text(root, project, agent)
    context = (args.context or "").lower()
    prompts = []
    for topic in rule_elicitation_topics():
        covered = any(str(keyword).lower() in existing for keyword in topic["keywords"])
        relevant = not context or any(str(keyword).lower() in context for keyword in topic["keywords"])
        if not covered and relevant:
            prompts.append(
                {
                    "topic": topic["id"],
                    "prompt": topic["prompt"],
                    "capture_hint": "If the user gives a durable preference or constraint, call rule-suggest instead of silently remembering it.",
                }
            )
        if len(prompts) >= args.limit:
            break
    if not prompts:
        prompts.append(
            {
                "topic": "open_rule_gap",
                "prompt": "Ask whether this collaboration revealed any durable rule about approval, evidence, role boundaries, quality, or data handling.",
                "capture_hint": "Create a rule proposal only after the user states the rule in their own words.",
            }
        )
    payload = {
        "version": "1.0.0",
        "kind": "digital-office-rule-elicitation",
        "project_id": project,
        "agent_id": agent,
        "conversation_prompts": prompts,
        "next_action": "rule-suggest --title <title> --body <user-stated-rule>",
    }
    print(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def infer_rule_scope(root: Path, *, title: str, body: str, project: str = "", agent: str = "") -> dict[str, Any]:
    text = f"{title}\n{body}".lower()
    registry = effective_agents_registry(root)
    agents = registry.get("agents", {})
    aliases = registry.get("aliases", {})
    global_terms = [
        "all agents",
        "every agent",
        "company",
        "global",
        "security",
        "approval",
        "data sharing",
        "memory",
        "knowledge promotion",
        "所有agent",
        "全部agent",
        "全局",
        "公司",
        "安全",
        "审批",
        "记忆",
        "知识沉淀",
    ]
    project_terms = ["this project", "client", "project", "delivery", "这个项目", "客户", "本项目", "交付"]
    agent_terms = ["only this agent", "writer", "designer", "lawyer", "coder", "researcher", "只针对", "写作", "设计", "律师", "工程", "研究"]
    matched_agent = agent
    if not matched_agent:
        for name in sorted(agents.keys(), key=len, reverse=True):
            if name.lower() in text:
                matched_agent = name
                break
    if not matched_agent:
        for alias, name in aliases.items():
            if str(alias).lower() in text:
                matched_agent = str(name)
                break
    if matched_agent:
        return {
            "scope": "agent",
            "confidence": "high" if agent else "medium",
            "agent_id": registered_agent(root, matched_agent),
            "project_id": project,
            "reasons": ["agent was explicitly supplied or mentioned"],
            "scope_options": ["agent", "project", "global"],
        }
    if any(term in text for term in global_terms):
        return {
            "scope": "global",
            "confidence": "medium",
            "agent_id": "",
            "project_id": "",
            "reasons": ["rule affects all Agents, safety, approvals, memory, or company knowledge"],
            "scope_options": ["global", "project", "agent"],
        }
    if project and any(term in text for term in project_terms):
        return {
            "scope": "project",
            "confidence": "medium",
            "agent_id": "",
            "project_id": project,
            "reasons": ["rule appears tied to this project, client, or delivery context"],
            "scope_options": ["project", "agent", "global"],
        }
    if any(term in text for term in agent_terms):
        return {
            "scope": "agent",
            "confidence": "low",
            "agent_id": "",
            "project_id": project,
            "reasons": ["rule sounds role-specific but no concrete Agent was identified"],
            "scope_options": ["agent", "project", "global"],
        }
    return {
        "scope": "project" if project else "global",
        "confidence": "low",
        "agent_id": "",
        "project_id": project,
        "reasons": ["scope is not explicit; human confirmation is required"],
        "scope_options": ["project", "global", "agent"] if project else ["global", "agent"],
    }


def rule_suggest(args: argparse.Namespace) -> int:
    root = system_root()
    body = args.body if args.body is not None else sys.stdin.read()
    if not body.strip():
        print("office-system: rule-suggest requires --body or stdin", file=sys.stderr)
        return 2
    project = safe_component(args.project, "project id") if args.project else ""
    agent = registered_agent(root, args.agent) if args.agent else ""
    scope = infer_rule_scope(root, title=args.title, body=body, project=project, agent=agent)
    proposal_id = safe_component(args.proposal_id, "rule proposal id") if args.proposal_id else f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}-{slugify(args.title)}"
    if rule_proposal_path(root, proposal_id).exists():
        print(f"office-system: rule proposal already exists: {proposal_id}", file=sys.stderr)
        return 2
    governance = rule_governance_report(
        root,
        title=args.title,
        body=body.strip(),
        scope=scope["scope"],
        project=scope.get("project_id", ""),
        agent=scope.get("agent_id", ""),
        source=args.source or "",
        created_by=args.created_by or "",
        scope_confidence=scope["confidence"],
    )
    proposal = {
        "version": "1.0.0",
        "kind": "digital-office-rule-proposal",
        "proposal_id": proposal_id,
        "status": "pending_user_confirmation",
        "title": safe_claim(args.title, "rule title"),
        "body": body.strip(),
        "proposed_scope": scope["scope"],
        "scope_confidence": scope["confidence"],
        "scope_reasons": scope["reasons"],
        "scope_options": scope["scope_options"],
        "project_id": scope.get("project_id", ""),
        "agent_id": scope.get("agent_id", ""),
        "source": safe_claim(args.source, "rule source", required=False),
        "created_by": safe_claim(args.created_by, "created by", required=False),
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "requires_user_confirmation": True,
        "governance": governance,
        "applied_rule_path": "",
        "decisions": [],
    }
    write_json(rule_proposal_path(root, proposal_id), proposal)
    append_audit_event(root, "rule_proposal_created", actor_id=args.created_by or "", actor_role=args.role or "", project_id=project, agent_id=proposal.get("agent_id", ""), resource_type="rule_proposal", resource_id=proposal_id, outcome="pending_user_confirmation", reason=args.source or "")
    print(json.dumps(proposal, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def rule_proposal_list(args: argparse.Namespace) -> int:
    root = system_root()
    records = []
    for proposal in read_records(root / "rule-proposals", "digital-office-rule-proposal"):
        if args.status and proposal.get("status") != args.status:
            continue
        if args.scope and proposal.get("proposed_scope") != args.scope:
            continue
        if args.project and proposal.get("project_id") != args.project:
            continue
        if args.agent and proposal.get("agent_id") != args.agent:
            continue
        records.append(proposal)
        if len(records) >= args.limit:
            break
    print(json.dumps({"rule_proposals": records}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def rule_proposal_decision(args: argparse.Namespace) -> int:
    root = system_root()
    if args.decision == "approve" and not args.confirmed:
        print("office-system: approving a rule proposal requires --confirmed", file=sys.stderr)
        return 2
    proposal = load_rule_proposal(root, args.proposal_id)
    if proposal.get("status") not in {"pending_user_confirmation", "needs_tuning", "approved"}:
        print("office-system: rule proposal is not open", file=sys.stderr)
        return 2
    scope = args.scope or str(proposal.get("proposed_scope", "global"))
    if scope not in {"global", "project", "agent"}:
        print(f"office-system: invalid rule scope: {scope}", file=sys.stderr)
        return 2
    project = safe_component(args.project, "project id") if args.project else str(proposal.get("project_id", ""))
    agent = registered_agent(root, args.agent) if args.agent else str(proposal.get("agent_id", ""))
    auth = compute_authorization_decision(
        root,
        tenant_id=args.tenant,
        deployment_id=args.deployment,
        user_id=args.decided_by,
        user_role=args.role,
        action="rule.manage",
        resource_type="rule_proposal",
        resource_id=args.proposal_id,
        project_id=project if scope == "project" else "",
        agent_id=agent if scope == "agent" else "",
        reason=args.message or "",
    )
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    status = RULE_PROPOSAL_DECISION_TO_STATUS[args.decision]
    proposal["status"] = status
    proposal["updated_at"] = now_iso()
    proposal.setdefault("decisions", []).append({"time": proposal["updated_at"], "decision": args.decision, "decided_by": args.decided_by, "role": args.role, "message": args.message or "", "scope": scope, "project_id": project, "agent_id": agent})
    applied_path = ""
    if args.decision == "approve":
        governance = rule_governance_report(
            root,
            title=args.title or proposal["title"],
            body=args.body or proposal["body"],
            scope=scope,
            project=project,
            agent=agent,
            source=str(proposal.get("source", "")),
            created_by=str(proposal.get("created_by", "")),
            scope_confidence=str(proposal.get("scope_confidence", "")),
        )
        proposal["governance"] = governance
        if governance.get("conflicts") and not args.override_conflicts:
            proposal["status"] = "pending_user_confirmation"
            proposal.setdefault("blockers", [])
            if "rule_conflict_requires_override" not in proposal["blockers"]:
                proposal["blockers"].append("rule_conflict_requires_override")
            write_json(rule_proposal_path(root, args.proposal_id), proposal)
            print(json.dumps({"status": "blocked", "reason": "rule_conflict_requires_override", "governance": governance, "authorization": auth}, ensure_ascii=False, indent=2, sort_keys=True))
            return 2
        target = write_rule_to_store(root, scope=scope, title=args.title or proposal["title"], body=args.body or proposal["body"], project=project, agent=agent)
        applied_path = str(target.relative_to(root))
        proposal["status"] = "applied"
        proposal["applied_rule_path"] = applied_path
        proposal["applied_scope"] = scope
        proposal["applied_project_id"] = project
        proposal["applied_agent_id"] = agent
    write_json(rule_proposal_path(root, args.proposal_id), proposal)
    append_audit_event(root, "rule_proposal_decision", actor_id=args.decided_by, actor_role=args.role, tenant_id=args.tenant, deployment_id=args.deployment, project_id=project, agent_id=agent, resource_type="rule_proposal", resource_id=args.proposal_id, outcome=proposal["status"], reason=args.message or "", extra={"applied_rule_path": applied_path, "scope": scope})
    print(json.dumps({"rule_proposal": proposal, "authorization": auth}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def iter_entry_manifests(path: Path) -> list[Path]:
    if not path.exists():
        return []
    return sorted(path.glob("*/entry.json"))


def load_entries(root: Path, base: Path) -> list[dict[str, Any]]:
    entries = []
    for manifest in iter_entry_manifests(base):
        try:
            entry = read_json(manifest)
            entry["_manifest"] = str(manifest.relative_to(root))
            entries.append(entry)
        except Exception:
            continue
    return entries


def tokenize(text: str) -> list[str]:
    words = re.findall(r"[a-zA-Z0-9_]+|[\u4e00-\u9fff]", text.lower())
    return [word for word in words if word.strip()]


def chunk_text(text: str, size: int = 1200, overlap: int = 180) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + size)
        chunks.append(text[start:end].strip())
        if end >= len(text):
            break
        start = max(0, end - overlap)
    return chunks


def extracted_text_for_entry(root: Path, entry: dict[str, Any]) -> str:
    extracted = root / str(entry.get("extracted_dir", ""))
    if not extracted.exists():
        return ""
    parts: list[str] = []
    for path in sorted(extracted.glob("*.md")) + sorted(extracted.glob("*.txt")):
        try:
            parts.append(path.read_text(encoding="utf-8", errors="replace"))
        except Exception:
            continue
    return "\n\n".join(parts)


def rag_entry_base(root: Path, scope: str, project: str | None) -> Path:
    if scope == "company":
        return root / "knowledge" / "company" / "entries"
    if not project:
        print("office-system: --project is required for project RAG", file=sys.stderr)
        raise SystemExit(2)
    return ensure_project(root, project) / "knowledge" / "entries"


def rag_index_dir(root: Path, scope: str, project: str | None) -> Path:
    if scope == "company":
        return root / "knowledge" / "company" / "index"
    if not project:
        print("office-system: --project is required for project RAG", file=sys.stderr)
        raise SystemExit(2)
    return ensure_project(root, project) / "knowledge" / "index"


def try_embed(texts: list[str], model_name: str) -> tuple[list[list[float]] | None, str]:
    try:
        from sentence_transformers import SentenceTransformer  # type: ignore
    except Exception:
        return None, "sentence_transformers_not_installed"
    try:
        model = SentenceTransformer(model_name)
        vectors = model.encode(texts, normalize_embeddings=True).tolist()
        return vectors, model_name
    except Exception as exc:
        return None, f"embedding_failed:{exc}"


def rag_index(args: argparse.Namespace) -> int:
    root = system_root()
    entries = load_entries(root, rag_entry_base(root, args.scope, args.project))
    records: list[dict[str, Any]] = []
    for entry in entries:
        if entry.get("status") == "archived":
            continue
        if not args.include_pending and not entry.get("agent_readable") and entry.get("status") != "approved_for_agent_use":
            continue
        text = extracted_text_for_entry(root, entry)
        for index, chunk in enumerate(chunk_text(text, args.chunk_chars, args.chunk_overlap), start=1):
            records.append(
                {
                    "chunk_id": f"{entry.get('entry_id')}::{index}",
                    "entry_id": entry.get("entry_id"),
                    "title": entry.get("title"),
                    "scope": args.scope,
                    "project_id": args.project,
                    "source_file": entry.get("source_file"),
                    "status": entry.get("status"),
                    "text": chunk,
                    "tokens": tokenize(chunk),
                }
            )

    mode = args.mode
    embedding_status = "not_requested"
    if records and mode in {"auto", "embedding"}:
        vectors, embedding_status = try_embed([record["text"] for record in records], args.embedding_model)
        if vectors:
            for record, vector in zip(records, vectors):
                record["embedding"] = vector
            mode = "embedding"
        elif args.mode == "embedding":
            print(f"office-system: embedding index failed: {embedding_status}", file=sys.stderr)
            return 1
        else:
            mode = "lexical"
    elif mode == "auto":
        mode = "lexical"

    target_dir = rag_index_dir(root, args.scope, args.project)
    target_dir.mkdir(parents=True, exist_ok=True)
    index_file = target_dir / "index.jsonl"
    with index_file.open("w", encoding="utf-8") as handle:
        for record in records:
            handle.write(json.dumps(record, ensure_ascii=False, sort_keys=True) + "\n")
    meta = {
        "version": "1.0.0",
        "scope": args.scope,
        "project_id": args.project,
        "mode": mode,
        "embedding_model": args.embedding_model if mode == "embedding" else None,
        "embedding_status": embedding_status,
        "chunks": len(records),
        "created_at": now_iso(),
    }
    write_json(target_dir / "index.meta.json", meta)
    append_log(root, {"event": "rag_index", "scope": args.scope, "project": args.project, "chunks": len(records), "mode": mode})
    print(json.dumps(meta, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def lexical_score(query_tokens: list[str], record_tokens: list[str]) -> float:
    if not query_tokens or not record_tokens:
        return 0.0
    counts: dict[str, int] = {}
    for token in record_tokens:
        counts[token] = counts.get(token, 0) + 1
    score = 0.0
    for token in query_tokens:
        score += math.log1p(counts.get(token, 0))
    return score / math.sqrt(len(record_tokens))


def dot(left: list[float], right: list[float]) -> float:
    return sum(a * b for a, b in zip(left, right))


def rag_search(args: argparse.Namespace) -> int:
    root = system_root()
    index_dir = rag_index_dir(root, args.scope, args.project)
    index_file = index_dir / "index.jsonl"
    if not index_file.exists():
        print(f"office-system: RAG index not found: {index_file}", file=sys.stderr)
        return 2
    meta = read_json(index_dir / "index.meta.json") if (index_dir / "index.meta.json").exists() else {}
    records = []
    with index_file.open("r", encoding="utf-8") as handle:
        for line in handle:
            if line.strip():
                records.append(json.loads(line))

    scored: list[tuple[float, dict[str, Any]]] = []
    if meta.get("mode") == "embedding":
        vectors, status = try_embed([args.query], meta.get("embedding_model") or args.embedding_model)
        if vectors:
            query_vector = vectors[0]
            for record in records:
                vector = record.get("embedding")
                if vector:
                    scored.append((dot(query_vector, vector), record))
        else:
            print(f"office-system: embedding search unavailable, falling back to lexical: {status}", file=sys.stderr)

    if not scored:
        query_tokens = tokenize(args.query)
        scored = [(lexical_score(query_tokens, record.get("tokens", [])), record) for record in records]

    results = [
        {
            "score": score,
            "chunk_id": record.get("chunk_id"),
            "entry_id": record.get("entry_id"),
            "title": record.get("title"),
            "source_file": record.get("source_file"),
            "status": record.get("status"),
            "text": record.get("text"),
        }
        for score, record in sorted(scored, key=lambda item: item[0], reverse=True)[: args.limit]
        if score > 0
    ]
    print(json.dumps({"query": args.query, "scope": args.scope, "project_id": args.project, "results": results}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def context(args: argparse.Namespace) -> int:
    root = system_root()
    agent = args.agent
    project = args.project
    project_dir = ensure_project(root, project) if project else None
    project_data = read_json(project_dir / "project.json") if project_dir else None

    print("# Digital Office Agent Context")
    print()
    print("## Priority")
    print("1. Current user task and explicit project selection")
    print("2. System bootstrap and company global rules")
    print("3. Active project rules and project knowledge source files")
    print("4. Agent-specific rules")
    print("5. Company global knowledge and approved methodologies")
    print("6. Licensed industry references mounted by entitlement")
    print("7. KeyMemory project relay snapshots, preferences, approved summaries, and retrieval pointers")
    print()
    print("Fact authority: project knowledge > company global knowledge > licensed industry reference > KeyMemory.")
    print("Handoff authority: current task > KeyMemory project relay > project latest decisions > company methods.")
    print("If KeyMemory or licensed references conflict with source-backed project or company knowledge, source-backed knowledge wins.")
    print()

    preferences = current_onboarding_preferences(root)
    print("## User Preferences")
    if preferences:
        print(f"- source: {preferences.get('source', 'unknown')}")
        print(f"- secretary_name: {preferences.get('secretary_name')}")
        print(f"- company_name: {preferences.get('company_name') or 'not set'}")
        for field in ONBOARDING_FIELDS:
            print(f"- {field}: {preferences.get('choices', {}).get(field)}")
    else:
        print("- not configured; use the neutral default behavior and let the GUI show global settings.")
    print("- Preferences do not override safety, authorization, approval, knowledge authority, or production harness gates.")
    print()

    print("## Rule Layers")
    for path in sorted((root / "rules" / "global").glob("*.md")):
        print(f"- company/global: {path.relative_to(root)}")
    if project_dir:
        for path in sorted((project_dir / "rules").glob("*.md")):
            print(f"- project/{project}: {path.relative_to(root)}")
    if agent:
        agent_rule = root / "rules" / "agents" / f"{agent}.md"
        if agent_rule.exists():
            print(f"- agent/{agent}: {agent_rule.relative_to(root)}")
    print()

    if project_data:
        print("## Project")
        print(f"- id: {project_data.get('project_id')}")
        print(f"- name: {project_data.get('name')}")
        print(f"- status: {project_data.get('status')}")
        print(f"- agent_roster: {', '.join(project_data.get('agent_roster', []))}")
        print(f"- methodology_schedule: {project_data.get('methodology_promotion', {}).get('schedule')}")
        print()

    print("## Project Knowledge")
    if project_dir:
        entries = load_entries(root, project_dir / "knowledge" / "entries")
        if entries:
            for entry in entries:
                print(f"- {entry.get('entry_id')}: {entry.get('title')} [{entry.get('kind')}, {entry.get('status')}]")
                print(f"  source: {entry.get('source_file')}")
                print(f"  extracted: {entry.get('extracted_dir')}")
        else:
            print("- none")
    else:
        print("- no active project")
    print()

    print("## Company Knowledge")
    company_entries = load_entries(root, root / "knowledge" / "company" / "entries")
    if company_entries:
        for entry in company_entries:
            print(f"- {entry.get('entry_id')}: {entry.get('title')} [{entry.get('kind')}, {entry.get('status')}]")
            print(f"  source: {entry.get('source_file')}")
    else:
        print("- no entry manifests yet")
    print()
    print("## KeyMemory")
    print("- Use after source-backed project and company knowledge.")
    print("- Allowed: preferences, durable operating memories, approved methodology summaries, semantic pointers.")
    print("- Forbidden: raw PDF/Word/image storage, unapproved project drafts, ordinary plaintext secrets.")
    return 0


def create_project(args: argparse.Namespace) -> int:
    root = system_root()
    project_id = args.project or slugify(args.name)
    agents = [registered_agent(root, item.strip()) for item in args.agents.split(",") if item.strip()] if args.agents else []
    target = copy_template_project(root, project_id, args.name, agents, args.methodology_schedule, guided_intake=bool(args.guided_intake), brief=args.brief or "")
    append_log(root, {"event": "project_create", "project": project_id})
    project = read_json(target / "project.json")
    assessment = assess_project_context(root, project_id, project)
    print(json.dumps({"status": "created", "project": project, "project_id": project_id, "path": str(target), "context_readiness": assessment}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def run_record_path(root: Path, run_id: str) -> Path:
    return run_dir(root, run_id) / "run.json"


def load_run_record(root: Path, run_id: str) -> dict[str, Any]:
    path = run_record_path(root, run_id)
    if not path.exists():
        print(f"office-system: workflow run not found: {run_id}", file=sys.stderr)
        raise SystemExit(2)
    run, _ = normalize_loop_run(root, read_json(path))
    return run


def read_run_records(root: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for path in sorted((root / "runs").glob("*/run.json")):
        try:
            run, _ = normalize_loop_run(root, read_json(path))
            records.append(run)
        except Exception:
            continue
    records.sort(key=lambda item: str(item.get("updated_at") or item.get("created_at") or ""), reverse=True)
    return records


def append_run_ledger_event(
    root: Path,
    run_id: str,
    event_type: str,
    *,
    stage: str = "",
    action: str = "",
    agent_id: str = "",
    actor_id: str = "",
    actor_role: str = "",
    input_payload: Any | None = None,
    output_payload: Any | None = None,
    artifact_refs: list[str] | None = None,
    checkpoint_id: str = "",
    handoff_id: str = "",
    model: str = "",
    provider: str = "",
    parameters: dict[str, Any] | None = None,
) -> dict[str, Any]:
    run_id = safe_component(run_id, "run id")
    path = run_ledger_path(root, run_id)
    input_hash = canonical_hash(input_payload) if input_payload is not None else ""
    output_hash = canonical_hash(output_payload) if output_payload is not None else ""
    event = {
        "version": "1.0.0",
        "kind": "digital-office-run-ledger-event",
        "event_id": f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}",
        "time": now_iso(),
        "run_id": run_id,
        "stage": normalize_loop_stage(stage) if stage else "",
        "event": safe_component(event_type, "ledger event"),
        "action": safe_claim(action, "ledger action", required=False),
        "agent_id": registered_agent(root, agent_id) if agent_id else "",
        "actor": {
            "user_id": safe_claim(actor_id, "ledger actor", required=False),
            "role": safe_component(actor_role, "ledger actor role") if actor_role else "",
        },
        "model": safe_claim(model, "ledger model", required=False),
        "provider": safe_claim(provider, "ledger provider", required=False),
        "parameters_hash": canonical_hash(parameters or {}),
        "input_hash": input_hash,
        "output_hash": output_hash,
        "input": input_payload if input_payload is not None else {},
        "output": output_payload if output_payload is not None else {},
        "artifact_refs": [safe_claim(item, "artifact ref") for item in (artifact_refs or [])],
        "checkpoint_id": safe_component(checkpoint_id, "checkpoint id") if checkpoint_id else "",
        "handoff_id": safe_component(handoff_id, "handoff id") if handoff_id else "",
        "previous_event_hash": "",
    }
    ledger_lock = path.with_name(f".{path.name}.lock")
    with JsonFileLock(ledger_lock):
        previous = read_last_jsonl(path)
        event["previous_event_hash"] = (previous or {}).get("event_hash", "")
        event["event_hash"] = canonical_hash({key: value for key, value in event.items() if key != "event_hash"})
        append_jsonl_unlocked(path, event)
    append_log(root, {"event": "run_ledger_event", "run_id": run_id, "ledger_event": event_type, "event_id": event["event_id"]})
    return event


def list_run_ledger(root: Path, run_id: str, limit: int | None = None) -> list[dict[str, Any]]:
    return read_jsonl(run_ledger_path(root, run_id), limit=limit)


def verify_run_ledger(root: Path, run_id: str) -> dict[str, Any]:
    run_id = safe_component(run_id, "run id")
    path = run_ledger_path(root, run_id)
    with JsonFileLock(path.with_name(f".{path.name}.lock")):
        events = read_jsonl(path)
    issues: list[dict[str, Any]] = []
    previous_hash = ""
    seen_ids: set[str] = set()
    for index, event in enumerate(events):
        event_id = str(event.get("event_id", ""))
        if not event_id:
            issues.append({"index": index, "issue": "missing_event_id"})
        elif event_id in seen_ids:
            issues.append({"index": index, "event_id": event_id, "issue": "duplicate_event_id"})
        seen_ids.add(event_id)
        if event.get("run_id") != run_id:
            issues.append({"index": index, "event_id": event_id, "issue": "run_id_mismatch"})
        if str(event.get("previous_event_hash", "")) != previous_hash:
            issues.append({"index": index, "event_id": event_id, "issue": "previous_event_hash_mismatch"})
        expected_hash = canonical_hash({key: value for key, value in event.items() if key != "event_hash"})
        actual_hash = str(event.get("event_hash", ""))
        if not hmac.compare_digest(expected_hash, actual_hash):
            issues.append({"index": index, "event_id": event_id, "issue": "event_hash_mismatch"})
        previous_hash = actual_hash
    return {
        "status": "valid" if not issues else "invalid",
        "run_id": run_id,
        "event_count": len(events),
        "last_event_hash": previous_hash,
        "issues": issues,
    }


def find_run_by_idempotency(root: Path, key: str | None) -> dict[str, Any] | None:
    if not key:
        return None
    for record in read_run_records(root):
        if record.get("idempotency_key") == key:
            return record
    return None


def normalize_loop_stage(stage: str | None, *, default: str = "context") -> str:
    value = str(stage or default).strip()
    return LOOP_STAGE_ALIASES.get(value, value)


def fresh_loop_stage_state(root: Path, stage: str) -> dict[str, Any]:
    stage = normalize_loop_stage(stage)
    manifest = loop_manifest(root)
    definition = manifest["stages"][stage]
    return {
        "status": "pending",
        "required_artifacts": list(definition["required_artifacts"]),
        "gates": [{"gate": gate, "status": "pending"} for gate in definition["gates"]],
        "artifacts": [],
        "notes": [],
    }


def loop_stage_records(root: Path) -> dict[str, Any]:
    return {stage: fresh_loop_stage_state(root, stage) for stage in LOOP_STAGES}


def initial_loop_control(root: Path, overrides: dict[str, Any] | None = None) -> dict[str, Any]:
    defaults = dict(loop_manifest(root).get("controller", {}).get("default_budgets", {}))
    for key, value in (overrides or {}).items():
        if value is not None:
            defaults[key] = int(value)
    return {
        "cycle_index": 1,
        "budgets": defaults,
        "usage": {
            "tool_calls": 0,
            "model_calls": 0,
            "input_tokens": 0,
            "output_tokens": 0,
            "cost_microunits": 0,
        },
        "stage_retries": {},
        "last_progress_score": 0.0,
        "stagnant_cycles": 0,
        "decision_history": [],
        "cycle_history": [],
        "started_at": now_iso(),
    }


def normalize_loop_run(root: Path, run: dict[str, Any]) -> tuple[dict[str, Any], bool]:
    changed = False
    stages = run.get("stages", {}) or {}
    if any(stage not in stages for stage in LOOP_STAGES):
        legacy_candidates = {
            "context": ["context", "perceive"],
            "decide": ["decide", "plan", "reason"],
            "act": ["act", "execute"],
            "evaluate": ["evaluate", "reflect", "iterate"],
        }
        normalized: dict[str, Any] = {}
        for stage in LOOP_STAGES:
            source = next((stages.get(name) for name in legacy_candidates[stage] if isinstance(stages.get(name), dict)), None)
            normalized[stage] = dict(source) if source is not None else fresh_loop_stage_state(root, stage)
        run["stages"] = normalized
        changed = True
    current = normalize_loop_stage(str(run.get("current_stage", "context")))
    if current not in LOOP_STAGES:
        current = "context"
    if run.get("current_stage") != current:
        run["current_stage"] = current
        changed = True
    if not isinstance(run.get("control"), dict):
        run["control"] = initial_loop_control(root)
        changed = True
    if not run.get("context_id"):
        run["context_id"] = str(run.get("run_id", ""))
        changed = True
    if not run.get("task_id"):
        tasks = run.get("tasks", []) or []
        run["task_id"] = str(tasks[0]) if tasks else str(run.get("run_id", ""))
        changed = True
    run.setdefault("loop_contract_version", "1.x-compat" if changed else "2.0.0")
    return run, changed


def loop_budget_blockers(run: dict[str, Any]) -> list[str]:
    control = run.get("control", {}) or {}
    budgets = control.get("budgets", {}) or {}
    usage = control.get("usage", {}) or {}
    blockers: list[str] = []
    if int(control.get("cycle_index", 1)) > int(budgets.get("max_cycles", 0) or 0) > 0:
        blockers.append("max_cycles_exceeded")
    for usage_key, budget_key in (
        ("tool_calls", "max_tool_calls"),
        ("model_calls", "max_model_calls"),
        ("cost_microunits", "max_cost_microunits"),
    ):
        limit = int(budgets.get(budget_key, 0) or 0)
        if limit > 0 and int(usage.get(usage_key, 0) or 0) >= limit:
            blockers.append(f"{budget_key}_exhausted")
    duration_limit = int(budgets.get("max_duration_seconds", 0) or 0)
    if duration_limit > 0:
        try:
            started = dt.datetime.fromisoformat(str(control.get("started_at", run.get("created_at", ""))))
            elapsed = (dt.datetime.now(dt.timezone.utc) - started.astimezone(dt.timezone.utc)).total_seconds()
            if elapsed >= duration_limit:
                blockers.append("max_duration_seconds_exhausted")
        except (TypeError, ValueError):
            blockers.append("invalid_loop_start_time")
    return blockers


def reset_loop_from_stage(root: Path, run: dict[str, Any], target_stage: str) -> None:
    target_stage = normalize_loop_stage(target_stage)
    start_index = LOOP_STAGES.index(target_stage)
    for stage in LOOP_STAGES[start_index:]:
        run.setdefault("stages", {})[stage] = fresh_loop_stage_state(root, stage)
    run["current_stage"] = target_stage
    run["status"] = LOOP_STATUS_BY_STAGE[target_stage]


def route_task(root: Path, task: str, requested_agent: str, workflow: str | None) -> dict[str, Any]:
    router = root.parent / "scripts" / "agent-router"
    if not router.exists():
        print(f"office-system: agent router not found: {router}", file=sys.stderr)
        raise SystemExit(2)
    command = [str(router), "--route-json", "--agent", requested_agent or "auto"]
    if workflow:
        command.extend(["--workflow", workflow])
    command.append(task)
    proc = subprocess.run(
        command,
        text=True,
        capture_output=True,
        env={**os.environ, "HERMES_HOME": str(root.parent)},
    )
    if proc.returncode not in {0, 3}:
        print(proc.stderr or proc.stdout or "office-system: route failed", file=sys.stderr)
        raise SystemExit(proc.returncode)
    try:
        return json.loads(proc.stdout)
    except json.JSONDecodeError:
        print("office-system: router returned invalid JSON", file=sys.stderr)
        raise SystemExit(1)


def judgment_policy(root: Path) -> dict[str, Any]:
    path = root / "judgment.policy.json"
    if path.exists():
        return read_json(path)
    return {
        "version": "0.0.0",
        "kind": "digital-office-judgment-policy",
        "thresholds": {
            "pause_risk_score": 0.65,
            "route_low_confidence_values": ["low", "ambiguous"],
            "agent_pause_decisions": ["pause", "escalate", "abort"],
        },
        "default_options": [
            {"id": "approve_with_scope", "label": "Approve with explicit scope"},
            {"id": "request_more_evidence", "label": "Request more evidence"},
            {"id": "reject_or_stop", "label": "Reject or stop this action"},
        ],
        "hard_block_actions": [],
        "categories": [],
    }


def parse_json_value(value: str | None, path: str | None, label: str, *, required: bool = False) -> dict[str, Any]:
    if path:
        data = read_json(Path(path).expanduser())
    elif value:
        try:
            data = json.loads(value)
        except json.JSONDecodeError as exc:
            print(f"office-system: invalid {label} JSON: {exc}", file=sys.stderr)
            raise SystemExit(2)
    elif required:
        text = sys.stdin.read().strip()
        if not text:
            print(f"office-system: {label} JSON is required", file=sys.stderr)
            raise SystemExit(2)
        try:
            data = json.loads(text)
        except json.JSONDecodeError as exc:
            print(f"office-system: invalid {label} JSON: {exc}", file=sys.stderr)
            raise SystemExit(2)
    else:
        return {}
    if not isinstance(data, dict):
        print(f"office-system: {label} JSON must be an object", file=sys.stderr)
        raise SystemExit(2)
    return data


def keyword_hit(text: str, keyword: str) -> bool:
    return bool(keyword) and keyword.lower() in text.lower()


def normalize_actions(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [safe_component(str(item), "judgment action") for item in value if str(item).strip()]
    return [safe_component(item.strip(), "judgment action") for item in str(value).split(",") if item.strip()]


def risk_label(score: float) -> str:
    if score >= 0.85:
        return "critical"
    if score >= 0.65:
        return "high"
    if score >= 0.4:
        return "medium"
    return "low"


def evaluate_judgment(
    root: Path,
    *,
    task: str,
    stage: str = "context",
    agent_id: str = "",
    workflow_run_id: str = "",
    task_id: str = "",
    action: str = "",
    route: dict[str, Any] | None = None,
    signal: dict[str, Any] | None = None,
) -> dict[str, Any]:
    policy = judgment_policy(root)
    thresholds = policy.get("thresholds", {})
    pause_threshold = float(thresholds.get("pause_risk_score", 0.65))
    route_pause_values = {str(item) for item in thresholds.get("route_low_confidence_values", ["low", "ambiguous"])}
    agent_pause_values = {str(item) for item in thresholds.get("agent_pause_decisions", ["pause", "escalate", "abort"])}
    route = route or {}
    signal = signal or {}
    text = "\n".join(
        [
            task or "",
            str(signal.get("reason", "")),
            str(signal.get("summary", "")),
            str(route.get("routing_reason", "")),
            str(route.get("workflow_reason", "")),
        ]
    )
    triggers: list[dict[str, Any]] = []
    blocked_actions: set[str] = set()
    required_roles: list[str] = []
    evidence_refs: list[str] = []
    options = list(policy.get("default_options", []))
    risk_score = 0.0

    for category in policy.get("categories", []):
        if not isinstance(category, dict):
            continue
        matched = [str(keyword) for keyword in category.get("keywords", []) if keyword_hit(text, str(keyword))]
        if not matched:
            continue
        category_score = float(category.get("risk_score", 0.0))
        risk_score = max(risk_score, category_score)
        blocked_actions.update(str(item) for item in category.get("blocked_actions", []))
        role = str(category.get("required_human_role", "project_manager"))
        if role and role not in required_roles:
            required_roles.append(role)
        triggers.append(
            {
                "type": "policy_category",
                "category": str(category.get("id", "unknown")),
                "risk_score": category_score,
                "matched_keywords": matched[:8],
                "required_human_role": role,
            }
        )

    route_confidence = str(route.get("confidence", ""))
    if route.get("fallback") or route.get("clarification_required") or route_confidence in route_pause_values:
        risk_score = max(risk_score, 0.72)
        if "workflow_continue" not in blocked_actions:
            blocked_actions.add("workflow_continue")
        if "project_manager" not in required_roles:
            required_roles.append("project_manager")
        triggers.append(
            {
                "type": "route_uncertainty",
                "risk_score": 0.72,
                "confidence": route_confidence,
                "fallback": bool(route.get("fallback")),
                "clarification_required": bool(route.get("clarification_required")),
            }
        )

    for item in normalize_actions(action):
        if item in set(str(entry) for entry in policy.get("hard_block_actions", [])):
            risk_score = max(risk_score, 0.9)
            blocked_actions.add(item)
            if "project_manager" not in required_roles:
                required_roles.append("project_manager")
            triggers.append({"type": "hard_block_action", "action": item, "risk_score": 0.9})

    agent_decision = str(signal.get("decision", "")).strip().lower()
    if agent_decision in agent_pause_values:
        score = float(signal.get("risk_score", 0.8) or 0.8)
        risk_score = max(risk_score, score)
        blocked_actions.update(str(item) for item in signal.get("blocked_actions", []) if str(item).strip())
        evidence_refs.extend(str(item) for item in signal.get("evidence_refs", []) if str(item).strip())
        role = str(signal.get("required_human_role", "project_manager"))
        if role and role not in required_roles:
            required_roles.append(role)
        if isinstance(signal.get("options"), list) and signal["options"]:
            options = signal["options"]
        triggers.append(
            {
                "type": "agent_stop_signal",
                "decision": agent_decision,
                "risk_score": score,
                "reason": str(signal.get("reason", "")),
            }
        )

    decision = "pause" if risk_score >= pause_threshold or bool(blocked_actions) else "continue"
    return {
        "version": "1.0.0",
        "kind": "digital-office-judgment-evaluation",
        "decision": decision,
        "risk_score": round(risk_score, 3),
        "risk_label": risk_label(risk_score),
        "stage": normalize_loop_stage(stage),
        "agent_id": registered_agent(root, agent_id) if agent_id else str(route.get("agent", "")),
        "workflow_run_id": safe_claim(workflow_run_id, "workflow run id", required=False),
        "task_id": safe_claim(task_id, "task id", required=False),
        "task_sha256": hashlib.sha256((task or "").encode("utf-8")).hexdigest(),
        "route_confidence": route_confidence,
        "triggers": triggers,
        "blocked_actions": sorted(blocked_actions),
        "required_human_role": required_roles[0] if required_roles else "project_manager",
        "required_human_roles": required_roles,
        "evidence_refs": evidence_refs,
        "options": options,
        "recommended_option": "request_more_evidence" if risk_score >= pause_threshold else "approve_with_scope",
        "created_at": now_iso(),
    }


def open_judgment_cases(root: Path, run: dict[str, Any] | None = None) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    if run is not None:
        for case_id in run.get("judgment_cases", []) or []:
            try:
                case = load_judgment_case(root, str(case_id))
            except SystemExit:
                continue
            if case.get("status") in JUDGMENT_BLOCKING_STATUSES:
                records.append(case)
        return records
    for case in read_records(root / "judgments", "digital-office-judgment-case"):
        if case.get("status") in JUDGMENT_BLOCKING_STATUSES:
            records.append(case)
    return records


def stage_artifact_satisfies(required: str, artifacts: list[Any]) -> bool:
    for item in artifacts:
        if isinstance(item, dict):
            value = str(item.get("artifact_id") or item.get("artifact") or item.get("path") or "")
        else:
            value = str(item)
        if value == required or value.startswith(f"{required}:") or value.startswith(f"{required}="):
            return True
    return False


def stage_gate_status(stage_state: dict[str, Any], gate_id: str) -> str:
    for gate in stage_state.get("gates", []) or []:
        if str(gate.get("gate")) == gate_id:
            return str(gate.get("status", "pending"))
    return "missing"


def loop_completion_blockers(root: Path, run: dict[str, Any], *, require_all_stages: bool = True) -> list[str]:
    blockers: list[str] = []
    open_cases = open_judgment_cases(root, run)
    if open_cases:
        blockers.append("human_judgment_pending")
    for handoff in read_records(handoff_dir(root, str(run.get("run_id", ""))), "digital-office-typed-handoff"):
        if handoff.get("status") in {"pending_acceptance", "needs_context"}:
            blockers.append(f"handoff_{handoff.get('handoff_id', 'unknown')}_not_accepted")
    if not require_all_stages:
        return blockers
    stages = run.get("stages", {}) or {}
    for stage in LOOP_STAGES:
        state = stages.get(stage, {})
        if state.get("status") not in {"completed", "skipped"}:
            blockers.append(f"stage_{stage}_not_completed")
            continue
        if state.get("status") == "skipped":
            continue
        for artifact in state.get("required_artifacts", []) or []:
            if not stage_artifact_satisfies(str(artifact), state.get("artifacts", []) or []):
                blockers.append(f"stage_{stage}_missing_artifact_{artifact}")
        for gate in state.get("gates", []) or []:
            gate_name = str(gate.get("gate", ""))
            gate_status = str(gate.get("status", "pending"))
            if gate_status not in PASSING_GATE_STATUSES:
                blockers.append(f"stage_{stage}_gate_{gate_name}_not_passed")
    return blockers


def status_for_current_stage(run: dict[str, Any]) -> str:
    return LOOP_STATUS_BY_STAGE.get(normalize_loop_stage(str(run.get("current_stage", "context"))), "created")


def create_judgment_case(
    root: Path,
    evaluation: dict[str, Any],
    *,
    task: str,
    reason: str = "",
    created_by: str = "",
    created_by_role: str = "",
    case_id: str | None = None,
) -> dict[str, Any]:
    case_id = safe_component(case_id, "judgment case id") if case_id else f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    if judgment_path(root, case_id).exists():
        print(f"office-system: judgment case already exists: {case_id}", file=sys.stderr)
        raise SystemExit(2)
    run_id = evaluation.get("workflow_run_id", "")
    task_id = evaluation.get("task_id", "")
    case = {
        "version": "1.0.0",
        "kind": "digital-office-judgment-case",
        "case_id": case_id,
        "status": "pending",
        "decision": evaluation.get("decision", "pause"),
        "stage": normalize_loop_stage(str(evaluation.get("stage", "context"))),
        "agent_id": evaluation.get("agent_id", ""),
        "workflow_run_id": run_id,
        "task_id": task_id,
        "risk_score": evaluation.get("risk_score", 0.0),
        "risk_label": evaluation.get("risk_label", "low"),
        "required_human_role": evaluation.get("required_human_role", "project_manager"),
        "blocked_actions": evaluation.get("blocked_actions", []),
        "triggers": evaluation.get("triggers", []),
        "evidence_refs": evaluation.get("evidence_refs", []),
        "options": evaluation.get("options", []),
        "recommended_option": evaluation.get("recommended_option", ""),
        "reason": reason or "; ".join(str(item.get("type", "")) for item in evaluation.get("triggers", [])) or "Agent requested human judgment",
        "task_sha256": hashlib.sha256((task or "").encode("utf-8")).hexdigest(),
        "context_snapshot_hash": hashlib.sha256(json.dumps(evaluation, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest(),
        "created_by": safe_claim(created_by, "created by", required=False),
        "created_by_role": safe_component(created_by_role, "created by role") if created_by_role else "",
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "decisions": [],
    }
    write_json(judgment_path(root, case_id), case)
    if run_id:
        run = load_run_record(root, str(run_id))
        cases = run.setdefault("judgment_cases", [])
        if case_id not in cases:
            cases.append(case_id)
        run["status"] = "waiting_human_judgment"
        run.setdefault("blockers", [])
        if "human_judgment_pending" not in run["blockers"]:
            run["blockers"].append("human_judgment_pending")
        run["updated_at"] = now_iso()
        run.setdefault("events", []).append({"time": run["updated_at"], "event": "judgment_case_opened", "case_id": case_id, "risk_score": case["risk_score"]})
        write_json(run_record_path(root, str(run_id)), run)
    if task_id:
        update_task_status(root, str(task_id), "waiting_human_judgment", message=f"waiting for judgment {case_id}", actor_id=created_by, actor_role=created_by_role)
    append_audit_event(
        root,
        "judgment_case_opened",
        actor_id=created_by,
        actor_role=created_by_role,
        resource_type="judgment",
        resource_id=case_id,
        workflow_run_id=str(run_id or ""),
        task_id=str(task_id or ""),
        agent_id=str(case.get("agent_id", "")),
        outcome="pending",
        reason=case["reason"],
        extra={"risk_score": case["risk_score"], "blocked_actions": case["blocked_actions"]},
    )
    if run_id:
        append_run_ledger_event(
            root,
            str(run_id),
            "human_judgment_opened",
            stage=str(case.get("stage", "")),
            action="judgment.open",
            agent_id=str(case.get("agent_id", "")),
            actor_id=created_by,
            actor_role=created_by_role,
            input_payload={"evaluation": evaluation, "task_sha256": case["task_sha256"]},
            output_payload={"case_id": case_id, "status": case["status"], "blocked_actions": case["blocked_actions"]},
        )
    return case


def task_title(body: str) -> str:
    compact = re.sub(r"\s+", " ", body).strip()
    return compact[:80] or "Untitled task"


def parse_json_arg(value: str | None, path: str | None, label: str) -> dict[str, Any]:
    if path:
        return read_json(Path(path).expanduser())
    if value:
        try:
            data = json.loads(value)
        except json.JSONDecodeError as exc:
            print(f"office-system: invalid {label} JSON: {exc}", file=sys.stderr)
            raise SystemExit(2)
        if not isinstance(data, dict):
            print(f"office-system: {label} JSON must be an object", file=sys.stderr)
            raise SystemExit(2)
        return data
    text = sys.stdin.read().strip()
    if not text:
        print(f"office-system: {label} JSON is required", file=sys.stderr)
        raise SystemExit(2)
    try:
        data = json.loads(text)
    except json.JSONDecodeError as exc:
        print(f"office-system: invalid {label} JSON: {exc}", file=sys.stderr)
        raise SystemExit(2)
    if not isinstance(data, dict):
        print(f"office-system: {label} JSON must be an object", file=sys.stderr)
        raise SystemExit(2)
    return data


def canvas_node(node_id: str, node_type: str, title: str, **extra: Any) -> dict[str, Any]:
    node_id = safe_component(node_id, "node id")
    if node_type not in CANVAS_NODE_TYPES:
        print(f"office-system: invalid canvas node type: {node_type}", file=sys.stderr)
        raise SystemExit(2)
    node = {
        "node_id": node_id,
        "type": node_type,
        "title": safe_claim(title, "node title"),
        "inputs": sorted(NODE_IO_TYPES.get(node_type, {}).get("inputs", set())),
        "outputs": sorted(NODE_IO_TYPES.get(node_type, {}).get("outputs", set())),
        "status": "pending",
    }
    node.update({key: value for key, value in extra.items() if value not in (None, "")})
    return node


def default_canvas_from_route(root: Path, *, route: dict[str, Any], body: str, agent_id: str, workflow: str) -> dict[str, Any]:
    steps = route.get("steps") or ([agent_id] if agent_id else [])
    nodes = [canvas_node("start", "start", "Start")]
    edges: list[dict[str, str]] = []
    previous = "start"
    for index, step_agent in enumerate(steps, start=1):
        registered = registered_agent(root, str(step_agent))
        node_id = f"agent-{index}-{registered}"
        nodes.append(
            canvas_node(
                node_id,
                "agent_task",
                f"{registered} task",
                agent_id=registered,
                instruction=body,
                execution_mode="direct" if workflow == "direct_agent" else "workflow",
            )
        )
        edges.append({"from": previous, "to": node_id})
        previous = node_id
    output_id = "final-output"
    nodes.append(canvas_node(output_id, "output_artifact", "Final output", required=True))
    edges.append({"from": previous, "to": output_id})
    return {
        "mode": "simple",
        "nodes": nodes,
        "edges": edges,
        "entry_node_id": "start",
        "final_node_id": output_id,
        "workflow": workflow,
    }


def new_canvas_revision(
    root: Path,
    *,
    revision_id: str,
    status: str,
    created_by: str,
    source: str,
    parent_revision_id: str = "",
    canvas: dict[str, Any],
    change_summary: str = "",
) -> dict[str, Any]:
    validation = validate_canvas(root, canvas)
    return {
        "revision_id": safe_component(revision_id, "revision id"),
        "status": safe_component(status, "revision status"),
        "created_at": now_iso(),
        "created_by": safe_claim(created_by, "revision creator", required=False),
        "source": safe_component(source, "revision source"),
        "parent_revision_id": safe_claim(parent_revision_id, "parent revision id", required=False),
        "change_summary": safe_claim(change_summary, "change summary", required=False),
        "canvas": canvas,
        "validation": validation,
    }


def initial_canvas_revision(root: Path, *, route: dict[str, Any], body: str, agent_id: str, workflow: str, created_by: str) -> dict[str, Any]:
    revision_id = f"rev-{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:6]}"
    canvas = default_canvas_from_route(root, route=route, body=body, agent_id=agent_id, workflow=workflow)
    return new_canvas_revision(
        root,
        revision_id=revision_id,
        status="active",
        created_by=created_by,
        source="workflow_start",
        canvas=canvas,
        change_summary="Initial workflow canvas generated from router decision.",
    )


def revision_list(run: dict[str, Any]) -> list[dict[str, Any]]:
    return run.setdefault("revisions", [])


def find_revision(run: dict[str, Any], revision_id: str) -> dict[str, Any] | None:
    revision_id = safe_component(revision_id, "revision id")
    for revision in revision_list(run):
        if revision.get("revision_id") == revision_id:
            return revision
    return None


def active_revision(run: dict[str, Any]) -> dict[str, Any] | None:
    active_id = run.get("active_revision_id")
    if active_id:
        found = find_revision(run, active_id)
        if found:
            return found
    for revision in revision_list(run):
        if revision.get("status") == "active":
            return revision
    return None


def normalize_canvas_payload(root: Path, payload: dict[str, Any], base_canvas: dict[str, Any] | None = None) -> dict[str, Any]:
    canvas = json.loads(json.dumps(base_canvas or {"mode": "simple", "nodes": [], "edges": []}, ensure_ascii=False))
    if "mode" in payload:
        mode = safe_component(str(payload["mode"]), "canvas mode")
        if mode not in {"simple", "professional", "admin"}:
            print(f"office-system: invalid canvas mode: {mode}", file=sys.stderr)
            raise SystemExit(2)
        canvas["mode"] = mode
    if "nodes" in payload:
        if not isinstance(payload["nodes"], list):
            print("office-system: canvas nodes must be a list", file=sys.stderr)
            raise SystemExit(2)
        nodes = []
        for item in payload["nodes"]:
            if not isinstance(item, dict):
                print("office-system: canvas node must be an object", file=sys.stderr)
                raise SystemExit(2)
            node_type = safe_component(str(item.get("type", "")), "canvas node type")
            node_id = safe_component(str(item.get("node_id", "")), "node id")
            title = safe_claim(str(item.get("title") or node_id), "node title")
            extra = {key: value for key, value in item.items() if key not in {"node_id", "type", "title", "inputs", "outputs"}}
            nodes.append(canvas_node(node_id, node_type, title, **extra))
        canvas["nodes"] = nodes
    if "edges" in payload:
        if not isinstance(payload["edges"], list):
            print("office-system: canvas edges must be a list", file=sys.stderr)
            raise SystemExit(2)
        edges = []
        for item in payload["edges"]:
            if not isinstance(item, dict):
                print("office-system: canvas edge must be an object", file=sys.stderr)
                raise SystemExit(2)
            edges.append({"from": safe_component(str(item.get("from", "")), "edge from"), "to": safe_component(str(item.get("to", "")), "edge to")})
        canvas["edges"] = edges
    for operation in payload.get("operations", []) or []:
        if not isinstance(operation, dict):
            print("office-system: canvas operation must be an object", file=sys.stderr)
            raise SystemExit(2)
        apply_canvas_operation(root, canvas, operation)
    return canvas


def apply_canvas_operation(root: Path, canvas: dict[str, Any], operation: dict[str, Any]) -> None:
    op = safe_component(str(operation.get("op", "")), "canvas operation")
    nodes = canvas.setdefault("nodes", [])
    edges = canvas.setdefault("edges", [])
    if op == "add_node":
        item = operation.get("node")
        if not isinstance(item, dict):
            print("office-system: add_node requires node object", file=sys.stderr)
            raise SystemExit(2)
        node = normalize_canvas_payload(root, {"nodes": [item]})["nodes"][0]
        if any(existing.get("node_id") == node["node_id"] for existing in nodes):
            print(f"office-system: canvas node already exists: {node['node_id']}", file=sys.stderr)
            raise SystemExit(2)
        nodes.append(node)
    elif op == "remove_node":
        node_id = safe_component(str(operation.get("node_id", "")), "node id")
        canvas["nodes"] = [node for node in nodes if node.get("node_id") != node_id]
        canvas["edges"] = [edge for edge in edges if edge.get("from") != node_id and edge.get("to") != node_id]
    elif op == "replace_node":
        item = operation.get("node")
        if not isinstance(item, dict):
            print("office-system: replace_node requires node object", file=sys.stderr)
            raise SystemExit(2)
        node = normalize_canvas_payload(root, {"nodes": [item]})["nodes"][0]
        replaced = False
        for index, existing in enumerate(nodes):
            if existing.get("node_id") == node["node_id"]:
                nodes[index] = node
                replaced = True
                break
        if not replaced:
            print(f"office-system: canvas node not found: {node['node_id']}", file=sys.stderr)
            raise SystemExit(2)
    elif op == "add_edge":
        edge = {"from": safe_component(str(operation.get("from", "")), "edge from"), "to": safe_component(str(operation.get("to", "")), "edge to")}
        if edge not in edges:
            edges.append(edge)
    elif op == "remove_edge":
        edge = {"from": safe_component(str(operation.get("from", "")), "edge from"), "to": safe_component(str(operation.get("to", "")), "edge to")}
        canvas["edges"] = [item for item in edges if item != edge]
    else:
        print(f"office-system: unsupported canvas operation: {op}", file=sys.stderr)
        raise SystemExit(2)


def validate_canvas(root: Path, canvas: dict[str, Any]) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    nodes = canvas.get("nodes") or []
    edges = canvas.get("edges") or []
    node_by_id: dict[str, dict[str, Any]] = {}
    if not isinstance(nodes, list) or not isinstance(edges, list):
        return {"status": "invalid", "errors": ["canvas nodes and edges must be lists"], "warnings": warnings}
    for node in nodes:
        node_id = str(node.get("node_id", ""))
        node_type = str(node.get("type", ""))
        if not node_id:
            errors.append("node missing node_id")
            continue
        if node_id in node_by_id:
            errors.append(f"duplicate node id: {node_id}")
            continue
        node_by_id[node_id] = node
        if node_type not in CANVAS_NODE_TYPES:
            errors.append(f"invalid node type for {node_id}: {node_type}")
        if canvas.get("mode", "simple") == "simple" and node_type not in SIMPLE_CANVAS_NODE_TYPES | {"start"}:
            warnings.append(f"{node_id} uses advanced component {node_type}")
        if node_type == "agent_task":
            agent = str(node.get("agent_id", ""))
            if not agent:
                errors.append(f"agent_task {node_id} missing agent_id")
            else:
                try:
                    registered_agent(root, agent)
                except SystemExit:
                    errors.append(f"agent_task {node_id} references unknown agent {agent}")
            if not str(node.get("instruction", "")).strip():
                errors.append(f"agent_task {node_id} missing instruction")
        if node_type == "file_ref" and not node.get("file_id"):
            errors.append(f"file_ref {node_id} missing file_id")
        if node_type == "folder_ref" and not node.get("folder_id"):
            errors.append(f"folder_ref {node_id} missing folder_id")
        if node_type == "knowledge_scope" and not (node.get("scope_id") or node.get("space_id")):
            errors.append(f"knowledge_scope {node_id} missing scope_id or space_id")

    if not any(node.get("type") == "start" for node in nodes):
        errors.append("canvas must include a start node")
    if not any(node.get("type") == "output_artifact" for node in nodes):
        errors.append("canvas must include a final output_artifact node")

    incoming: dict[str, int] = {node_id: 0 for node_id in node_by_id}
    outgoing: dict[str, int] = {node_id: 0 for node_id in node_by_id}
    adjacency: dict[str, list[str]] = {node_id: [] for node_id in node_by_id}
    for edge in edges:
        source = str(edge.get("from", ""))
        target = str(edge.get("to", ""))
        if source not in node_by_id:
            errors.append(f"edge references missing source node: {source}")
            continue
        if target not in node_by_id:
            errors.append(f"edge references missing target node: {target}")
            continue
        source_outputs = set(node_by_id[source].get("outputs") or NODE_IO_TYPES.get(str(node_by_id[source].get("type")), {}).get("outputs", set()))
        target_inputs = set(node_by_id[target].get("inputs") or NODE_IO_TYPES.get(str(node_by_id[target].get("type")), {}).get("inputs", set()))
        if target_inputs and source_outputs and not source_outputs.intersection(target_inputs):
            errors.append(f"edge {source}->{target} has incompatible output/input types")
        incoming[target] += 1
        outgoing[source] += 1
        adjacency[source].append(target)

    for node_id, node in node_by_id.items():
        node_type = node.get("type")
        if node_type != "start" and incoming.get(node_id, 0) == 0:
            errors.append(f"node is unreachable: {node_id}")
        if node_type != "output_artifact" and outgoing.get(node_id, 0) == 0:
            errors.append(f"node has no outgoing edge: {node_id}")

    visiting: set[str] = set()
    visited: set[str] = set()

    def visit(node_id: str) -> None:
        if node_id in visiting:
            errors.append(f"cycle detected at node: {node_id}")
            return
        if node_id in visited:
            return
        visiting.add(node_id)
        for child in adjacency.get(node_id, []):
            visit(child)
        visiting.remove(node_id)
        visited.add(node_id)

    for node_id in node_by_id:
        visit(node_id)

    for node in nodes:
        if node.get("type") == "agent_task" and (node.get("risk") in {"high", "regulated"} or node.get("requires_approval")):
            seen: set[str] = set()
            stack = list(adjacency.get(str(node.get("node_id")), []))
            has_approval = False
            while stack:
                current = stack.pop()
                if current in seen:
                    continue
                seen.add(current)
                if node_by_id.get(current, {}).get("type") == "approval_gate":
                    has_approval = True
                    break
                stack.extend(adjacency.get(current, []))
            if not has_approval:
                errors.append(f"high-risk agent node lacks downstream approval_gate: {node.get('node_id')}")

    return {"status": "valid" if not errors else "invalid", "errors": errors, "warnings": warnings}


def create_task_record(
    root: Path,
    *,
    task_id: str,
    title: str,
    body: str,
    status: str,
    priority: str,
    project_id: str,
    agent_id: str,
    workflow_run_id: str,
    assigned_user: str,
    requested_by: str,
    route: dict[str, Any],
    idempotency_key: str = "",
) -> dict[str, Any]:
    if status not in TASK_STATUSES:
        print(f"office-system: invalid task status: {status}", file=sys.stderr)
        raise SystemExit(2)
    task = {
        "version": "1.0.0",
        "kind": "digital-office-task",
        "task_id": task_id,
        "title": title,
        "body": body,
        "status": status,
        "priority": priority,
        "project_id": project_id,
        "agent_id": agent_id,
        "assigned_agent": agent_id,
        "assigned_user": assigned_user,
        "requested_by": requested_by,
        "workflow_run_id": workflow_run_id,
        "route": route,
        "idempotency_key": idempotency_key,
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "history": [{"time": now_iso(), "status": status, "message": "task created"}],
        "artifacts": [],
    }
    write_json(task_path(root, task_id), task)
    return task


def auth_decision(args: argparse.Namespace) -> int:
    root = system_root()
    decision = compute_authorization_decision(
        root,
        tenant_id=args.tenant,
        deployment_id=args.deployment,
        user_id=args.user,
        user_role=args.role,
        action=args.action,
        resource_type=args.resource_type,
        resource_id=args.resource_id,
        project_id=args.project or "",
        agent_id=args.agent or "",
        workflow_run_id=args.workflow_run or "",
        reason=args.reason or "",
    )
    print(json.dumps(decision, ensure_ascii=False, indent=2, sort_keys=True))
    return 0 if decision["allowed"] else 2


def workflow_start(args: argparse.Namespace) -> int:
    root = system_root()
    body = args.task if args.task is not None else sys.stdin.read()
    body = body.strip()
    if not body:
        print("office-system: workflow-start requires --task or stdin body", file=sys.stderr)
        return 2
    existing = find_run_by_idempotency(root, args.idempotency_key)
    if existing:
        print(json.dumps({"idempotent": True, "run": existing}, ensure_ascii=False, indent=2, sort_keys=True))
        return 0

    project = safe_component(args.project, "project id") if args.project else ""
    if project:
        project_file = ensure_project(root, project) / "project.json"
        project_record = read_json(project_file)
        context_assessment = assess_project_context(root, project, project_record)
        if context_assessment["required"] and not context_assessment["confirmed"]:
            print(json.dumps({
                "status": "needs_context",
                "reason": "project_context_not_confirmed",
                "project_id": project,
                "context_readiness": context_assessment,
                "next_actions": ["project-context-update", "project-context-confirm"],
            }, ensure_ascii=False, indent=2, sort_keys=True))
            return 4
    route = route_task(root, body, args.agent, args.workflow)
    agent = registered_agent(root, str(route.get("agent", ""))) if route.get("agent") else ""
    judgment_eval = evaluate_judgment(
        root,
        task=body,
        stage="context",
        agent_id=agent,
        action="workflow.start",
        route=route,
    )
    judgment_required = judgment_eval["decision"] == "pause"
    run_id = safe_component(args.run_id, "run id") if args.run_id else f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    task_id = safe_component(args.task_id, "task id") if args.task_id else f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    if run_record_path(root, run_id).exists():
        print(f"office-system: workflow run already exists: {run_id}", file=sys.stderr)
        return 2
    if task_path(root, task_id).exists():
        print(f"office-system: task already exists: {task_id}", file=sys.stderr)
        return 2

    auth = compute_authorization_decision(
        root,
        tenant_id=args.tenant,
        deployment_id=args.deployment,
        user_id=args.user,
        user_role=args.role,
        action="workflow.start",
        resource_type="workflow_run",
        resource_id=run_id,
        project_id=project,
        agent_id=agent,
        workflow_run_id=run_id,
        reason=args.reason or "",
    )
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2

    needs_clarification = bool(route.get("fallback") or route.get("clarification_required"))
    run_status = "waiting_human_judgment" if judgment_required else ("blocked" if needs_clarification else "created")
    task_status = "waiting_human_judgment" if judgment_required else ("blocked" if needs_clarification else "queued")
    initial_revision = initial_canvas_revision(root, route=route, body=body, agent_id=agent, workflow=route.get("workflow", ""), created_by=args.user)
    run = {
        "version": "1.0.0",
        "kind": "digital-office-workflow-run",
        "run_id": run_id,
        "context_id": run_id,
        "task_id": task_id,
        "run_type": "workflow_run",
        "status": run_status,
        "current_stage": "context",
        "project_id": project,
        "agent_id": agent,
        "workflow": route.get("workflow", ""),
        "requested_by": safe_claim(args.user, "requested by"),
        "requested_by_role": safe_component(args.role, "requested by role"),
        "tenant_id": safe_claim(args.tenant, "tenant id"),
        "deployment_id": safe_claim(args.deployment, "deployment id"),
        "task": body,
        "task_sha256": hashlib.sha256(body.encode("utf-8")).hexdigest(),
        "tasks": [task_id],
        "approvals": [],
        "judgment_cases": [],
        "judgment_evaluation": judgment_eval,
        "route": route,
        "active_revision_id": initial_revision["revision_id"],
        "revisions": [initial_revision],
        "canvas": initial_revision["canvas"],
        "authorization": auth,
        "blockers": ["human_judgment_pending"] if judgment_required else (["clarification_required"] if needs_clarification else []),
        "idempotency_key": args.idempotency_key or "",
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "loop_contract_version": "2.0.0",
        "control": initial_loop_control(root),
        "stages": loop_stage_records(root),
        "events": [{"time": now_iso(), "event": "workflow_started", "status": run_status}],
    }
    task = create_task_record(
        root,
        task_id=task_id,
        title=args.title or task_title(body),
        body=body,
        status=task_status,
        priority=args.priority,
        project_id=project,
        agent_id=agent,
        workflow_run_id=run_id,
        assigned_user=args.user,
        requested_by=args.user,
        route=route,
        idempotency_key=args.idempotency_key or "",
    )
    write_json(run_record_path(root, run_id), run)
    append_run_ledger_event(
        root,
        run_id,
        "workflow_started",
        stage="context",
        action="workflow.start",
        agent_id=agent,
        actor_id=args.user,
        actor_role=args.role,
        input_payload={
            "task": body,
            "title": args.title or task["title"],
            "project_id": project,
            "requested_agent": args.agent or "",
            "requested_workflow": args.workflow or "",
            "idempotency_key": args.idempotency_key or "",
        },
        output_payload={
            "run_id": run_id,
            "task_id": task_id,
            "status": run_status,
            "task_status": task_status,
            "workflow": run["workflow"],
            "route": route,
            "judgment_required": judgment_required,
            "clarification_required": needs_clarification,
        },
        artifact_refs=[
            str(run_record_path(root, run_id).relative_to(root)),
            str(task_path(root, task_id).relative_to(root)),
        ],
    )
    judgment_case = None
    if judgment_required:
        judgment_eval["workflow_run_id"] = run_id
        judgment_eval["task_id"] = task_id
        judgment_case = create_judgment_case(
            root,
            judgment_eval,
            task=body,
            reason="workflow-start requires human judgment before dispatch",
            created_by=args.user,
            created_by_role=args.role,
        )
    event = append_audit_event(
        root,
        "workflow_started",
        actor_id=args.user,
        actor_role=args.role,
        tenant_id=args.tenant,
        deployment_id=args.deployment,
        project_id=project,
        agent_id=agent,
        resource_type="workflow_run",
        resource_id=run_id,
        workflow_run_id=run_id,
        task_id=task_id,
        outcome="waiting_human_judgment" if judgment_required else ("blocked" if needs_clarification else "created"),
        reason="human judgment required" if judgment_required else ("clarification required" if needs_clarification else args.reason or ""),
        extra={"workflow": run["workflow"], "route_confidence": route.get("confidence")},
    )
    emit_notification(
        root,
        user_id=args.user,
        title="Workflow needs human judgment" if judgment_required else ("Workflow started" if not needs_clarification else "Workflow needs clarification"),
        body=task["title"],
        topic="workflow",
        resource_type="workflow_run",
        resource_id=run_id,
        severity="warning" if judgment_required or needs_clarification else "info",
    )
    print(
        json.dumps(
            {
                "run_id": run_id,
                "task_id": task_id,
                "status": run_status,
                "task_status": task_status,
                "route": route,
                "judgment": judgment_eval,
                "judgment_case": judgment_case,
                "authorization": auth,
                "audit_event_id": event["event_id"],
                "next_actions": ["judgment-list", "judgment-decision", "workflow-status"] if judgment_required else ["workflow-status", "task-status", "approval-create"],
            },
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
        )
    )
    return 0


def workflow_status(args: argparse.Namespace) -> int:
    root = system_root()
    print(json.dumps(load_run_record(root, args.run_id), ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def workflow_list(args: argparse.Namespace) -> int:
    root = system_root()
    project = safe_component(args.project, "project id") if args.project else ""
    records = []
    for run in read_run_records(root):
        if args.status and run.get("status") != args.status:
            continue
        if project and run.get("project_id") != project:
            continue
        if args.user and run.get("requested_by") != args.user:
            continue
        records.append(run)
        if len(records) >= args.limit:
            break
    print(json.dumps({"runs": records}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def linked_tasks(root: Path, run: dict[str, Any]) -> list[str]:
    return [safe_component(task_id, "task id") for task_id in run.get("tasks", []) if task_id]


def sync_run_status_from_tasks(root: Path, run_id: str) -> dict[str, Any] | None:
    if not run_id:
        return None
    run = load_run_record(root, run_id)
    task_ids = linked_tasks(root, run)
    if not task_ids:
        return run
    tasks = [load_task(root, task_id) for task_id in task_ids]
    statuses = {task.get("status") for task in tasks}
    next_status = ""
    event_name = ""
    if statuses == {"completed"}:
        blockers = loop_completion_blockers(root, run, require_all_stages=True)
        if blockers:
            next_status = "blocked"
            event_name = "workflow_blocked_by_completion_contract"
            run.setdefault("blockers", [])
            for blocker in blockers:
                if blocker not in run["blockers"]:
                    run["blockers"].append(blocker)
        else:
            next_status = "completed"
            event_name = "workflow_completed"
    elif "failed" in statuses:
        next_status = "blocked"
        event_name = "workflow_blocked_by_task_failure"
        run.setdefault("blockers", [])
        if "task_failed" not in run["blockers"]:
            run["blockers"].append("task_failed")
    elif statuses == {"cancelled"}:
        next_status = "cancelled"
        event_name = "workflow_cancelled_by_tasks"
    elif "waiting_human_judgment" in statuses:
        next_status = "waiting_human_judgment"
        event_name = "workflow_waiting_for_human_judgment"
    elif "waiting_approval" in statuses:
        next_status = "waiting_user_confirmation"
        event_name = "workflow_waiting_for_task_approval"
    if not next_status or run.get("status") == next_status:
        return run
    run["status"] = next_status
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": event_name, "task_statuses": sorted(str(item) for item in statuses)})
    write_json(run_record_path(root, run_id), run)
    return run


def workflow_cancel(args: argparse.Namespace) -> int:
    root = system_root()
    if not args.confirmed:
        print("office-system: workflow-cancel requires --confirmed", file=sys.stderr)
        return 2
    run = load_run_record(root, args.run_id)
    auth = compute_authorization_decision(
        root,
        tenant_id=run.get("tenant_id", "local"),
        deployment_id=run.get("deployment_id", "local"),
        user_id=args.requested_by,
        user_role=args.role,
        action="workflow.cancel",
        resource_type="workflow_run",
        resource_id=args.run_id,
        project_id=run.get("project_id", ""),
        agent_id=run.get("agent_id", ""),
        workflow_run_id=args.run_id,
    )
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    if run.get("status") in {"completed", "cancelled"}:
        print("office-system: completed or cancelled workflow cannot be cancelled again", file=sys.stderr)
        return 2
    run["status"] = "cancelled"
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "workflow_cancelled", "reason": args.reason or ""})
    for task_id in linked_tasks(root, run):
        update_task_status(root, task_id, "cancelled", message=args.reason or "workflow cancelled", actor_id=args.requested_by or "", actor_role=args.role or "")
    write_json(run_record_path(root, args.run_id), run)
    append_audit_event(
        root,
        "workflow_cancelled",
        actor_id=args.requested_by or "",
        actor_role=args.role or "",
        project_id=run.get("project_id", ""),
        agent_id=run.get("agent_id", ""),
        resource_type="workflow_run",
        resource_id=args.run_id,
        workflow_run_id=args.run_id,
        outcome="cancelled",
        reason=args.reason or "",
    )
    print(json.dumps({"run_id": args.run_id, "status": "cancelled"}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def workflow_resume(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    auth = compute_authorization_decision(
        root,
        tenant_id=run.get("tenant_id", "local"),
        deployment_id=run.get("deployment_id", "local"),
        user_id=args.requested_by,
        user_role=args.role,
        action="workflow.resume",
        resource_type="workflow_run",
        resource_id=args.run_id,
        project_id=run.get("project_id", ""),
        agent_id=run.get("agent_id", ""),
        workflow_run_id=args.run_id,
    )
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    if run.get("status") == "cancelled":
        print("office-system: cancelled workflow must be retried instead of resumed", file=sys.stderr)
        return 2
    if run.get("status") == "completed":
        print("office-system: completed workflow cannot be resumed", file=sys.stderr)
        return 2
    open_cases = open_judgment_cases(root, run)
    if open_cases:
        print(json.dumps({"status": "blocked", "reason": "human_judgment_pending", "open_judgments": open_cases}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    run["status"] = LOOP_STATUS_BY_STAGE.get(normalize_loop_stage(str(run.get("current_stage", "context"))), "created")
    run["updated_at"] = now_iso()
    run.setdefault("blockers", [])
    run["blockers"] = [item for item in run["blockers"] if item != "clarification_required"]
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "workflow_resumed", "reason": args.reason or ""})
    for task_id in linked_tasks(root, run):
        task = load_task(root, task_id)
        if task.get("status") in {"blocked", "waiting_approval"}:
            update_task_status(root, task_id, "queued", message=args.reason or "workflow resumed", actor_id=args.requested_by or "", actor_role=args.role or "")
    write_json(run_record_path(root, args.run_id), run)
    append_audit_event(root, "workflow_resumed", actor_id=args.requested_by or "", actor_role=args.role or "", project_id=run.get("project_id", ""), agent_id=run.get("agent_id", ""), resource_type="workflow_run", resource_id=args.run_id, workflow_run_id=args.run_id, outcome="resumed", reason=args.reason or "")
    print(json.dumps({"run_id": args.run_id, "status": run["status"]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def workflow_retry(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    auth = compute_authorization_decision(
        root,
        tenant_id=run.get("tenant_id", "local"),
        deployment_id=run.get("deployment_id", "local"),
        user_id=args.requested_by,
        user_role=args.role,
        action="workflow.retry",
        resource_type="workflow_run",
        resource_id=args.run_id,
        project_id=run.get("project_id", ""),
        agent_id=run.get("agent_id", ""),
        workflow_run_id=args.run_id,
    )
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    stage = normalize_loop_stage(args.stage or str(run.get("current_stage", "context")))
    if stage not in LOOP_STAGES:
        print(f"office-system: invalid retry stage: {stage}", file=sys.stderr)
        return 2
    open_cases = open_judgment_cases(root, run)
    if open_cases:
        print(json.dumps({"status": "blocked", "reason": "human_judgment_pending", "open_judgments": open_cases}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    budget_blockers = loop_budget_blockers(run)
    if budget_blockers:
        print(json.dumps({"status": "budget_exhausted", "blockers": budget_blockers}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    control = run.setdefault("control", initial_loop_control(root))
    retries = control.setdefault("stage_retries", {})
    retries[stage] = int(retries.get(stage, 0)) + 1
    max_retries = int(control.get("budgets", {}).get("max_stage_retries", 0) or 0)
    if max_retries >= 0 and retries[stage] > max_retries:
        print(json.dumps({"status": "blocked", "reason": "stage_retry_budget_exhausted", "stage": stage, "attempts": retries[stage]}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    retry = {
        "time": now_iso(),
        "stage": stage,
        "requested_by": safe_claim(args.requested_by, "requested by", required=False),
        "reason": args.reason or "",
    }
    run.setdefault("retries", []).append(retry)
    reset_loop_from_stage(root, run, stage)
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "workflow_retry", "stage": stage, "reason": args.reason or ""})
    for task_id in linked_tasks(root, run):
        task = load_task(root, task_id)
        if task.get("status") in {"failed", "blocked", "cancelled"}:
            update_task_status(root, task_id, "queued", message=args.reason or "workflow retry", actor_id=args.requested_by or "", actor_role=args.role or "")
    write_json(run_record_path(root, args.run_id), run)
    append_audit_event(root, "workflow_retry", actor_id=args.requested_by or "", actor_role=args.role or "", project_id=run.get("project_id", ""), agent_id=run.get("agent_id", ""), resource_type="workflow_run", resource_id=args.run_id, workflow_run_id=args.run_id, outcome="retry_scheduled", reason=args.reason or "", extra={"stage": stage})
    print(json.dumps({"run_id": args.run_id, "stage": stage, "status": run["status"]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def agent_invoke(args: argparse.Namespace) -> int:
    root = system_root()
    if not args.project:
        print(json.dumps({"status": "denied", "error": "agent-invoke requires --project for governed dispatch"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    project = safe_component(args.project, "project id")
    ensure_project(root, project)
    registry = effective_agents_registry(root)
    agent = safe_component(args.agent, "agent id")
    if agent not in registry.get("agents", {}):
        print(json.dumps({"status": "denied", "error": "unknown agent", "agent_id": agent}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    body = args.task if args.task is not None else sys.stdin.read()
    body = body.strip()
    if not body:
        print(json.dumps({"status": "denied", "error": "agent-invoke requires --task or stdin body"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    run_id = safe_component(args.run_id, "run id") if args.run_id else f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    task_id = safe_component(args.task_id, "task id") if args.task_id else f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    if run_record_path(root, run_id).exists() or task_path(root, task_id).exists():
        print(json.dumps({"status": "denied", "error": "run or task already exists", "run_id": run_id, "task_id": task_id}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    auth = compute_authorization_decision(root, tenant_id=args.tenant, deployment_id=args.deployment, user_id=args.user, user_role=args.role, action="agent.delegate", resource_type="agent", resource_id=agent, project_id=project, agent_id=agent, workflow_run_id=run_id, reason=args.reason or "direct GUI @Agent invocation")
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    route = {"agent": agent, "workflow": "direct_agent", "steps": [agent], "confidence": "direct", "routing_reason": "explicit_agent_invocation", "display_name": registry["agents"][agent].get("display_name", agent)}
    judgment_eval = evaluate_judgment(root, task=body, stage="act", agent_id=agent, workflow_run_id=run_id, action="agent.delegate", route=route)
    judgment_required = judgment_eval["decision"] == "pause"
    initial_revision = initial_canvas_revision(root, route=route, body=body, agent_id=agent, workflow="direct_agent", created_by=args.user)
    run = {
        "version": "1.0.0",
        "kind": "digital-office-workflow-run",
        "run_id": run_id,
        "context_id": run_id,
        "task_id": task_id,
        "run_type": "workflow_run",
        "invocation_mode": "direct_agent",
        "status": "waiting_human_judgment" if judgment_required else "created",
        "current_stage": "context",
        "project_id": project,
        "agent_id": agent,
        "workflow": "direct_agent",
        "requested_by": safe_claim(args.user, "requested by"),
        "requested_by_role": safe_component(args.role, "requested by role"),
        "tenant_id": safe_claim(args.tenant, "tenant id"),
        "deployment_id": safe_claim(args.deployment, "deployment id"),
        "task": body,
        "task_sha256": hashlib.sha256(body.encode("utf-8")).hexdigest(),
        "tasks": [task_id],
        "approvals": [],
        "judgment_cases": [],
        "judgment_evaluation": judgment_eval,
        "route": route,
        "active_revision_id": initial_revision["revision_id"],
        "revisions": [initial_revision],
        "canvas": initial_revision["canvas"],
        "authorization": auth,
        "blockers": ["human_judgment_pending"] if judgment_required else [],
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "loop_contract_version": "2.0.0",
        "control": initial_loop_control(root),
        "stages": loop_stage_records(root),
        "events": [{"time": now_iso(), "event": "agent_invoked", "agent_id": agent, "status": "waiting_human_judgment" if judgment_required else "created"}],
    }
    task = create_task_record(root, task_id=task_id, title=args.title or task_title(body), body=body, status="waiting_human_judgment" if judgment_required else "queued", priority=args.priority, project_id=project, agent_id=agent, workflow_run_id=run_id, assigned_user=args.user, requested_by=args.user, route=route, idempotency_key="")
    task["invocation_mode"] = "direct_agent"
    write_json(task_path(root, task_id), task)
    write_json(run_record_path(root, run_id), run)
    append_run_ledger_event(
        root,
        run_id,
        "agent_invoked",
        stage="context",
        action="agent.delegate",
        agent_id=agent,
        actor_id=args.user,
        actor_role=args.role,
        input_payload={
            "task": body,
            "title": args.title or task["title"],
            "project_id": project,
            "agent_id": agent,
            "reason": args.reason or "direct GUI @Agent invocation",
        },
        output_payload={
            "run_id": run_id,
            "task_id": task_id,
            "status": run["status"],
            "task_status": task["status"],
            "route": route,
            "judgment_required": judgment_required,
        },
        artifact_refs=[
            str(run_record_path(root, run_id).relative_to(root)),
            str(task_path(root, task_id).relative_to(root)),
        ],
    )
    judgment_case = None
    if judgment_required:
        judgment_eval["task_id"] = task_id
        judgment_case = create_judgment_case(root, judgment_eval, task=body, reason="direct Agent invocation requires human judgment before dispatch", created_by=args.user, created_by_role=args.role)
    event = append_audit_event(root, "agent_invoked", actor_id=args.user, actor_role=args.role, tenant_id=args.tenant, deployment_id=args.deployment, project_id=project, agent_id=agent, resource_type="agent", resource_id=agent, workflow_run_id=run_id, task_id=task_id, outcome="created", reason=args.reason or "direct GUI @Agent invocation")
    emit_notification(root, user_id=args.user, title="Agent task needs human judgment" if judgment_required else "Agent task queued", body=task["title"], topic="agent", resource_type="workflow_run", resource_id=run_id, severity="warning" if judgment_required else "info")
    print(json.dumps({"status": run["status"], "invocation_mode": "direct_agent", "agent_id": agent, "project_id": project, "requested_by": args.user, "workflow_run_id": run_id, "task_id": task_id, "active_revision_id": initial_revision["revision_id"], "authorization": auth, "audit_event_id": event["event_id"], "judgment": judgment_eval, "judgment_case": judgment_case}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def authorize_workflow_mutation(root: Path, run: dict[str, Any], user: str, role: str, action: str) -> dict[str, Any]:
    return compute_authorization_decision(root, tenant_id=run.get("tenant_id", "local"), deployment_id=run.get("deployment_id", "local"), user_id=user, user_role=role, action=action, resource_type="workflow_run", resource_id=run.get("run_id", ""), project_id=run.get("project_id", ""), agent_id=run.get("agent_id", ""), workflow_run_id=run.get("run_id", ""))


def workflow_draft_create(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    auth = authorize_workflow_mutation(root, run, args.created_by, args.role, "workflow.edit")
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    parent = active_revision(run)
    if not parent:
        parent = initial_canvas_revision(root, route=run.get("route", {}), body=run.get("task", ""), agent_id=run.get("agent_id", ""), workflow=run.get("workflow", ""), created_by=args.created_by)
        revision_list(run).append(parent)
        run["active_revision_id"] = parent["revision_id"]
        run["canvas"] = parent["canvas"]
    revision_id = safe_component(args.revision_id, "revision id") if args.revision_id else f"draft-{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:6]}"
    if find_revision(run, revision_id):
        print(f"office-system: workflow revision already exists: {revision_id}", file=sys.stderr)
        return 2
    draft = new_canvas_revision(root, revision_id=revision_id, status="draft", created_by=args.created_by, source="gui_canvas_editor", parent_revision_id=parent["revision_id"], canvas=parent["canvas"], change_summary=args.summary or "Draft workflow canvas revision.")
    revision_list(run).append(draft)
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "workflow_draft_created", "revision_id": revision_id, "parent_revision_id": parent["revision_id"]})
    write_json(run_record_path(root, args.run_id), run)
    append_audit_event(root, "workflow_draft_created", actor_id=args.created_by, actor_role=args.role, project_id=run.get("project_id", ""), agent_id=run.get("agent_id", ""), resource_type="workflow_run", resource_id=args.run_id, workflow_run_id=args.run_id, outcome="draft", reason=args.summary or "")
    print(json.dumps({"run_id": args.run_id, "revision": draft}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def workflow_draft_patch(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    revision = find_revision(run, args.revision_id)
    if not revision or revision.get("status") != "draft":
        print("office-system: workflow draft revision not found or not draft", file=sys.stderr)
        return 2
    auth = authorize_workflow_mutation(root, run, args.updated_by, args.role, "workflow.edit")
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    patch = parse_json_arg(args.patch_json, args.patch_file, "workflow patch")
    old_canvas = revision.get("canvas", {})
    old_completed = {node.get("node_id"): node for node in old_canvas.get("nodes", []) if node.get("status") == "completed"}
    canvas = normalize_canvas_payload(root, patch, old_canvas)
    new_by_id = {node.get("node_id"): node for node in canvas.get("nodes", [])}
    for node_id, old_node in old_completed.items():
        if node_id not in new_by_id or new_by_id[node_id] != old_node:
            print(f"office-system: completed node cannot be modified in draft revision: {node_id}", file=sys.stderr)
            return 2
    revision["canvas"] = canvas
    revision["updated_at"] = now_iso()
    revision["updated_by"] = safe_claim(args.updated_by, "updated by", required=False)
    revision["change_summary"] = args.summary or revision.get("change_summary", "")
    revision["validation"] = validate_canvas(root, canvas)
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "workflow_draft_patched", "revision_id": args.revision_id, "validation_status": revision["validation"]["status"]})
    write_json(run_record_path(root, args.run_id), run)
    append_audit_event(root, "workflow_draft_patched", actor_id=args.updated_by, actor_role=args.role, project_id=run.get("project_id", ""), agent_id=run.get("agent_id", ""), resource_type="workflow_run", resource_id=args.run_id, workflow_run_id=args.run_id, outcome=revision["validation"]["status"], reason=args.summary or "")
    print(json.dumps({"run_id": args.run_id, "revision": revision}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def workflow_draft_validate(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    revision = find_revision(run, args.revision_id)
    if not revision:
        print("office-system: workflow revision not found", file=sys.stderr)
        return 2
    revision["validation"] = validate_canvas(root, revision.get("canvas", {}))
    revision["validated_at"] = now_iso()
    write_json(run_record_path(root, args.run_id), run)
    print(json.dumps({"run_id": args.run_id, "revision_id": args.revision_id, "validation": revision["validation"]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def workflow_draft_activate(args: argparse.Namespace) -> int:
    root = system_root()
    if not args.confirmed:
        print("office-system: workflow-draft-activate requires --confirmed", file=sys.stderr)
        return 2
    run = load_run_record(root, args.run_id)
    revision = find_revision(run, args.revision_id)
    if not revision or revision.get("status") != "draft":
        print("office-system: workflow draft revision not found or not draft", file=sys.stderr)
        return 2
    auth = authorize_workflow_mutation(root, run, args.activated_by, args.role, "workflow.edit")
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    validation = validate_canvas(root, revision.get("canvas", {}))
    revision["validation"] = validation
    if validation["status"] != "valid":
        write_json(run_record_path(root, args.run_id), run)
        print(json.dumps({"status": "invalid", "validation": validation}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    for item in revision_list(run):
        if item.get("status") == "active":
            item["status"] = "superseded"
            item["superseded_at"] = now_iso()
    revision["status"] = "active"
    revision["activated_at"] = now_iso()
    revision["activated_by"] = safe_claim(args.activated_by, "activated by", required=False)
    run["active_revision_id"] = args.revision_id
    run["canvas"] = revision.get("canvas", {})
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "workflow_draft_activated", "revision_id": args.revision_id})
    write_json(run_record_path(root, args.run_id), run)
    append_audit_event(root, "workflow_draft_activated", actor_id=args.activated_by, actor_role=args.role, project_id=run.get("project_id", ""), agent_id=run.get("agent_id", ""), resource_type="workflow_run", resource_id=args.run_id, workflow_run_id=args.run_id, outcome="active", reason=args.reason or "")
    print(json.dumps({"run_id": args.run_id, "active_revision_id": args.revision_id, "validation": validation}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def workflow_control(args: argparse.Namespace) -> int:
    root = system_root()
    action = safe_component(args.action, "workflow control action")
    if action not in WORKFLOW_CONTROL_ACTIONS:
        print(f"office-system: invalid workflow control action: {action}", file=sys.stderr)
        return 2
    if action == "stop" and not args.confirmed:
        print("office-system: workflow stop requires --confirmed", file=sys.stderr)
        return 2
    run = load_run_record(root, args.run_id)
    auth = authorize_workflow_mutation(root, run, args.requested_by, args.role, "workflow.control")
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    if run.get("status") in {"completed", "cancelled", "stopped"}:
        print("office-system: completed, cancelled, or stopped workflow cannot be controlled", file=sys.stderr)
        return 2
    if action in {"run", "resume"}:
        run["status"] = LOOP_STATUS_BY_STAGE.get(normalize_loop_stage(str(run.get("current_stage", "context"))), "created")
        run["pause_requested"] = False
        outcome = "running"
    elif action == "pause":
        run["pause_requested"] = True
        run["status"] = "paused_after_current_node"
        outcome = "pause_requested"
    else:
        run["status"] = "stopped"
        run["pause_requested"] = False
        outcome = "stopped"
        for task_id in linked_tasks(root, run):
            task = load_task(root, task_id)
            if task.get("status") not in {"completed", "failed", "cancelled"}:
                update_task_status(root, task_id, "cancelled", message=args.reason or "workflow stopped", actor_id=args.requested_by, actor_role=args.role)
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": f"workflow_control_{action}", "reason": args.reason or ""})
    write_json(run_record_path(root, args.run_id), run)
    append_audit_event(root, f"workflow_control_{action}", actor_id=args.requested_by, actor_role=args.role, project_id=run.get("project_id", ""), agent_id=run.get("agent_id", ""), resource_type="workflow_run", resource_id=args.run_id, workflow_run_id=args.run_id, outcome=outcome, reason=args.reason or "")
    print(json.dumps({"run_id": args.run_id, "action": action, "status": run["status"], "outcome": outcome}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def workflow_node_context(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    active = active_revision(run)
    if not active:
        print("office-system: workflow has no active revision", file=sys.stderr)
        return 2
    if args.revision_id and args.revision_id != active.get("revision_id"):
        print(json.dumps({"status": "stale_revision", "active_revision_id": active.get("revision_id"), "requested_revision_id": args.revision_id}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    node_id = safe_component(args.node_id, "node id")
    nodes = active.get("canvas", {}).get("nodes", [])
    node = next((item for item in nodes if item.get("node_id") == node_id), None)
    if not node:
        print(f"office-system: workflow node not found: {node_id}", file=sys.stderr)
        return 2
    edges = active.get("canvas", {}).get("edges", [])
    upstream_ids = [edge["from"] for edge in edges if edge.get("to") == node_id]
    downstream_ids = [edge["to"] for edge in edges if edge.get("from") == node_id]
    upstream_nodes = [item for item in nodes if item.get("node_id") in upstream_ids]
    print(json.dumps({"kind": "digital-office-workflow-node-context", "run_id": args.run_id, "project_id": run.get("project_id", ""), "active_revision_id": active.get("revision_id"), "node_id": node_id, "node": node, "upstream_node_ids": upstream_ids, "downstream_node_ids": downstream_ids, "inputs": upstream_nodes, "run_status": run.get("status"), "pause_requested": bool(run.get("pause_requested"))}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def task_list(args: argparse.Namespace) -> int:
    root = system_root()
    project = safe_component(args.project, "project id") if args.project else ""
    records = []
    for task in read_records(root / "tasks", "digital-office-task"):
        if args.status and task.get("status") != args.status:
            continue
        if project and task.get("project_id") != project:
            continue
        if args.assigned_agent and task.get("assigned_agent") != args.assigned_agent:
            continue
        if args.assigned_user and task.get("assigned_user") != args.assigned_user:
            continue
        records.append(task)
        if len(records) >= args.limit:
            break
    print(json.dumps({"tasks": records}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def task_status(args: argparse.Namespace) -> int:
    root = system_root()
    print(json.dumps(load_task(root, args.task_id), ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def task_update(args: argparse.Namespace) -> int:
    root = system_root()
    task = load_task(root, args.task_id)
    run = load_run_record(root, task["workflow_run_id"]) if task.get("workflow_run_id") else {}
    auth = compute_authorization_decision(
        root,
        tenant_id=run.get("tenant_id", "local"),
        deployment_id=run.get("deployment_id", "local"),
        user_id=args.updated_by,
        user_role=args.role,
        action="task.manage",
        resource_type="task",
        resource_id=args.task_id,
        project_id=task.get("project_id", ""),
        agent_id=task.get("assigned_agent", ""),
        workflow_run_id=task.get("workflow_run_id", ""),
    )
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    if args.status == "completed" and run and open_judgment_cases(root, run):
        print("office-system: task cannot be completed while human judgment is pending", file=sys.stderr)
        return 2
    if args.status:
        task = update_task_status(root, args.task_id, args.status, message=args.summary or "", actor_id=args.updated_by or "", actor_role=args.role or "")
    if args.assigned_agent:
        task["assigned_agent"] = registered_agent(root, args.assigned_agent)
    if args.assigned_user:
        task["assigned_user"] = safe_claim(args.assigned_user, "assigned user")
    for artifact in args.artifact or []:
        task.setdefault("artifacts", []).append(safe_claim(artifact, "artifact"))
    task["updated_at"] = now_iso()
    write_json(task_path(root, args.task_id), task)
    if task.get("workflow_run_id"):
        sync_run_status_from_tasks(root, task["workflow_run_id"])
    append_audit_event(root, "task_updated", actor_id=args.updated_by or "", actor_role=args.role or "", project_id=task.get("project_id", ""), agent_id=task.get("assigned_agent", ""), resource_type="task", resource_id=args.task_id, workflow_run_id=task.get("workflow_run_id", ""), task_id=args.task_id, outcome=task.get("status", "updated"), reason=args.summary or "")
    print(json.dumps(task, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def judgment_evaluate(args: argparse.Namespace) -> int:
    root = system_root()
    task = args.task if args.task is not None else sys.stdin.read()
    route = parse_json_value(args.route_json, args.route_file, "route", required=False)
    signal = parse_json_value(args.signal_json, args.signal_file, "agent signal", required=False)
    evaluation = evaluate_judgment(
        root,
        task=task.strip(),
        stage=args.stage,
        agent_id=args.agent or "",
        workflow_run_id=args.workflow_run or "",
        task_id=args.task_id or "",
        action=args.action or "",
        route=route,
        signal=signal,
    )
    case = None
    if args.create_case and evaluation["decision"] == "pause":
        case = create_judgment_case(
            root,
            evaluation,
            task=task.strip(),
            reason=args.reason or "judgment-evaluate created a blocking case",
            created_by=args.created_by or "",
            created_by_role=args.role or "",
            case_id=args.case_id,
        )
    print(json.dumps({"evaluation": evaluation, "case": case}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def judgment_list(args: argparse.Namespace) -> int:
    root = system_root()
    records = []
    for case in read_records(root / "judgments", "digital-office-judgment-case"):
        if args.status and case.get("status") != args.status:
            continue
        if args.workflow_run and case.get("workflow_run_id") != args.workflow_run:
            continue
        if args.required_human_role and case.get("required_human_role") != args.required_human_role:
            continue
        records.append(case)
        if len(records) >= args.limit:
            break
    print(json.dumps({"judgments": records}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def judgment_decision(args: argparse.Namespace) -> int:
    root = system_root()
    if args.decision in {"approve", "reject", "cancel"} and not args.confirmed:
        print("office-system: judgment decision requires --confirmed", file=sys.stderr)
        return 2
    case = load_judgment_case(root, args.case_id)
    if case.get("status") not in JUDGMENT_BLOCKING_STATUSES:
        print("office-system: judgment case is not pending", file=sys.stderr)
        return 2
    required_role = str(case.get("required_human_role", "project_manager"))
    if args.role not in TENANT_ADMIN_ROLES and args.role != required_role:
        print(f"office-system: judgment requires role {required_role}", file=sys.stderr)
        return 2
    action = "regulated_output.approve" if required_role == "professional_reviewer" else "approval.decide"
    auth = compute_authorization_decision(
        root,
        tenant_id=args.tenant,
        deployment_id=args.deployment,
        user_id=args.decided_by,
        user_role=args.role,
        action=action,
        resource_type="judgment",
        resource_id=args.case_id,
        project_id=args.project or "",
        agent_id=str(case.get("agent_id", "")),
        workflow_run_id=str(case.get("workflow_run_id", "")),
        reason=args.message or "",
    )
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    status = JUDGMENT_DECISION_TO_STATUS[args.decision]
    case["status"] = status
    case["updated_at"] = now_iso()
    if args.scope_note:
        case["scope_note"] = safe_claim(args.scope_note, "scope note")
    case.setdefault("decisions", []).append(
        {
            "time": case["updated_at"],
            "decision": args.decision,
            "status": status,
            "decided_by": safe_claim(args.decided_by, "decided by"),
            "role": safe_component(args.role, "role"),
            "message": args.message or "",
            "scope_note": args.scope_note or "",
        }
    )
    write_json(judgment_path(root, args.case_id), case)

    run_id = str(case.get("workflow_run_id", ""))
    task_id = str(case.get("task_id", ""))
    if run_id:
        run = load_run_record(root, run_id)
        run["updated_at"] = now_iso()
        run.setdefault("events", []).append({"time": run["updated_at"], "event": "judgment_decision", "case_id": args.case_id, "status": status})
        if status == "approved":
            remaining = open_judgment_cases(root, run)
            if not remaining:
                run["blockers"] = [item for item in run.get("blockers", []) if item != "human_judgment_pending"]
                run["status"] = status_for_current_stage(run)
        elif status == "needs_evidence":
            run["status"] = "waiting_human_judgment"
            run.setdefault("blockers", [])
            if "human_judgment_pending" not in run["blockers"]:
                run["blockers"].append("human_judgment_pending")
        else:
            run["status"] = "blocked" if status == "rejected" else run.get("status", "blocked")
            run.setdefault("blockers", [])
            blocker = f"judgment_{status}"
            if blocker not in run["blockers"]:
                run["blockers"].append(blocker)
        write_json(run_record_path(root, run_id), run)
    if task_id:
        if status == "approved":
            task = load_task(root, task_id)
            if task.get("status") == "waiting_human_judgment":
                update_task_status(root, task_id, "queued", message=args.message or f"judgment {status}", actor_id=args.decided_by, actor_role=args.role)
        elif status in {"rejected", "cancelled"}:
            update_task_status(root, task_id, "blocked", message=args.message or f"judgment {status}", actor_id=args.decided_by, actor_role=args.role)
        elif status == "needs_evidence":
            update_task_status(root, task_id, "waiting_human_judgment", message=args.message or "more evidence requested", actor_id=args.decided_by, actor_role=args.role)
    event = append_audit_event(
        root,
        "judgment_decision",
        actor_id=args.decided_by,
        actor_role=args.role,
        tenant_id=args.tenant,
        deployment_id=args.deployment,
        project_id=args.project or "",
        agent_id=str(case.get("agent_id", "")),
        resource_type="judgment",
        resource_id=args.case_id,
        workflow_run_id=run_id,
        task_id=task_id,
        outcome=status,
        reason=args.message or "",
        extra={"decision": args.decision, "required_human_role": required_role},
    )
    if run_id:
        append_run_ledger_event(
            root,
            run_id,
            "human_judgment_decision",
            stage=str(case.get("stage", "")),
            action=f"judgment.{args.decision}",
            agent_id=str(case.get("agent_id", "")),
            actor_id=args.decided_by,
            actor_role=args.role,
            input_payload={"case_id": args.case_id, "decision": args.decision, "message": args.message or "", "scope_note": args.scope_note or ""},
            output_payload={"case_status": status, "run_status": load_run_record(root, run_id).get("status", "")},
        )
    emit_notification(root, user_id="", title=f"Judgment {status}", body=case.get("reason", ""), topic="judgment", resource_type="judgment", resource_id=args.case_id, severity="info" if status == "approved" else "warning")
    print(json.dumps({"judgment": case, "authorization": auth, "audit_event_id": event["event_id"]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def judgment_resume(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    auth = authorize_workflow_mutation(root, run, args.requested_by, args.role, "workflow.resume")
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    remaining = open_judgment_cases(root, run)
    if remaining:
        print(json.dumps({"status": "blocked", "reason": "human_judgment_pending", "open_judgments": remaining}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    run["blockers"] = [item for item in run.get("blockers", []) if item != "human_judgment_pending"]
    run["status"] = status_for_current_stage(run)
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "judgment_resume", "requested_by": args.requested_by, "reason": args.reason or ""})
    write_json(run_record_path(root, args.run_id), run)
    for task_id in linked_tasks(root, run):
        task = load_task(root, task_id)
        if task.get("status") == "waiting_human_judgment":
            update_task_status(root, task_id, "queued", message=args.reason or "judgment resume", actor_id=args.requested_by, actor_role=args.role)
    append_audit_event(root, "judgment_resume", actor_id=args.requested_by, actor_role=args.role, project_id=run.get("project_id", ""), agent_id=run.get("agent_id", ""), resource_type="workflow_run", resource_id=args.run_id, workflow_run_id=args.run_id, outcome=run["status"], reason=args.reason or "")
    print(json.dumps({"run_id": args.run_id, "status": run["status"], "authorization": auth}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def approval_create(args: argparse.Namespace) -> int:
    root = system_root()
    existing = find_by_idempotency(root, "approvals", args.idempotency_key)
    if existing:
        print(json.dumps({"idempotent": True, "approval": existing}, ensure_ascii=False, indent=2, sort_keys=True))
        return 0
    project = safe_component(args.project, "project id") if args.project else ""
    agent = registered_agent(root, args.agent) if args.agent else ""
    approval_id = safe_component(args.approval_id, "approval id") if args.approval_id else f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    if approval_path(root, approval_id).exists():
        print(f"office-system: approval already exists: {approval_id}", file=sys.stderr)
        return 2
    auth = compute_authorization_decision(
        root,
        tenant_id=args.tenant,
        deployment_id=args.deployment,
        user_id=args.requested_by,
        user_role=args.requested_by_role,
        action="approval.create",
        resource_type=args.resource_type,
        resource_id=args.resource_id,
        project_id=project,
        agent_id=agent,
        workflow_run_id=args.workflow_run or "",
    )
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    record = {
        "version": "1.0.0",
        "kind": "digital-office-approval",
        "approval_id": approval_id,
        "status": "pending",
        "title": args.title,
        "body": args.body or "",
        "action": safe_claim(args.action, "approval action"),
        "resource_type": safe_component(args.resource_type, "resource type"),
        "resource_id": safe_claim(args.resource_id, "resource id"),
        "tenant_id": safe_claim(args.tenant, "tenant id"),
        "deployment_id": safe_claim(args.deployment, "deployment id"),
        "project_id": project,
        "agent_id": agent,
        "workflow_run_id": safe_claim(args.workflow_run, "workflow run id", required=False),
        "task_id": safe_claim(args.task_id, "task id", required=False),
        "requested_by": safe_claim(args.requested_by, "requested by"),
        "requested_by_role": safe_component(args.requested_by_role, "requested by role"),
        "approver_role": safe_component(args.approver_role, "approver role"),
        "risk": args.risk or "",
        "expires_at": safe_claim(args.expires_at, "expires at", required=False),
        "idempotency_key": args.idempotency_key or "",
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "decisions": [],
    }
    write_json(approval_path(root, approval_id), record)
    if record["task_id"]:
        update_task_status(root, record["task_id"], "waiting_approval", message=f"waiting for approval {approval_id}", actor_id=args.requested_by, actor_role=args.requested_by_role)
    if record["workflow_run_id"]:
        run = load_run_record(root, record["workflow_run_id"])
        run.setdefault("approvals", []).append(approval_id)
        run["status"] = "waiting_user_confirmation"
        run["updated_at"] = now_iso()
        run.setdefault("events", []).append({"time": run["updated_at"], "event": "approval_requested", "approval_id": approval_id})
        write_json(run_record_path(root, record["workflow_run_id"]), run)
    event = append_audit_event(root, "approval_created", actor_id=args.requested_by, actor_role=args.requested_by_role, tenant_id=args.tenant, deployment_id=args.deployment, project_id=project, agent_id=agent, resource_type="approval", resource_id=approval_id, workflow_run_id=record["workflow_run_id"], task_id=record["task_id"], approval_id=approval_id, outcome="pending", reason=args.title)
    emit_notification(root, user_id="", title="Approval requested", body=args.title, topic="approval", resource_type="approval", resource_id=approval_id, severity="warning")
    print(json.dumps({"approval": record, "audit_event_id": event["event_id"]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def approval_list(args: argparse.Namespace) -> int:
    root = system_root()
    project = safe_component(args.project, "project id") if args.project else ""
    records = []
    for approval in read_records(root / "approvals", "digital-office-approval"):
        if args.status and approval.get("status") != args.status:
            continue
        if project and approval.get("project_id") != project:
            continue
        if args.approver_role and approval.get("approver_role") != args.approver_role:
            continue
        records.append(approval)
        if len(records) >= args.limit:
            break
    print(json.dumps({"approvals": records}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def approval_decision(args: argparse.Namespace) -> int:
    root = system_root()
    if args.decision in {"approve", "reject"} and not args.confirmed:
        print("office-system: approval decision requires --confirmed", file=sys.stderr)
        return 2
    approval_path_value = approval_path(root, args.approval_id)
    # Acquire exclusive lock on the approval file to prevent race conditions
    lock_file = approval_path_value.with_name(f".{approval_path_value.name}.lock")
    with JsonFileLock(lock_file):
        # Re-read approval under lock to ensure atomicity
        approval = read_json(approval_path_value)
        if approval.get("kind") != "digital-office-approval":
            print("office-system: invalid approval record", file=sys.stderr)
            return 2
        auth = compute_authorization_decision(
            root,
            tenant_id=approval.get("tenant_id", "local"),
            deployment_id=approval.get("deployment_id", "local"),
            user_id=args.decided_by,
            user_role=args.role,
            action="approval.decide",
            resource_type="approval",
            resource_id=args.approval_id,
            project_id=approval.get("project_id", ""),
            agent_id=approval.get("agent_id", ""),
            workflow_run_id=approval.get("workflow_run_id", ""),
        )
        if not auth["allowed"]:
            print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
            return 2
        if approval.get("status") != "pending":
            print("office-system: approval is not pending", file=sys.stderr)
            return 2
        if approval.get("approver_role") == "professional_reviewer" and args.role != "professional_reviewer":
            print("office-system: this approval requires professional_reviewer", file=sys.stderr)
            return 2
        if args.role not in TENANT_ADMIN_ROLES and args.role != approval.get("approver_role"):
            print("office-system: actor role does not match required approver role", file=sys.stderr)
            return 2
        status = APPROVAL_DECISION_TO_STATUS[args.decision]
        approval["status"] = status
        approval["updated_at"] = now_iso()
        approval["decisions"].append({"time": approval["updated_at"], "decision": args.decision, "decided_by": args.decided_by, "role": args.role, "message": args.message or ""})
        write_json_unlocked(approval_path_value, approval)
        if approval.get("task_id"):
            update_task_status(root, approval["task_id"], "queued" if status == "approved" else "blocked", message=args.message or f"approval {status}", actor_id=args.decided_by, actor_role=args.role)
        if approval.get("workflow_run_id"):
            run = load_run_record(root, approval["workflow_run_id"])
            if status == "approved":
                if open_judgment_cases(root, run):
                    run["status"] = "waiting_human_judgment"
                    run.setdefault("blockers", [])
                    if "human_judgment_pending" not in run["blockers"]:
                        run["blockers"].append("human_judgment_pending")
                else:
                    run["status"] = LOOP_STATUS_BY_STAGE.get(normalize_loop_stage(str(run.get("current_stage", "context"))), "created")
            else:
                run["status"] = "blocked"
                run.setdefault("blockers", []).append(f"approval_{status}")
            run["updated_at"] = now_iso()
            run.setdefault("events", []).append({"time": run["updated_at"], "event": "approval_decision", "approval_id": args.approval_id, "status": status})
            write_json(run_record_path(root, approval["workflow_run_id"]), run)
    event = append_audit_event(root, "approval_decision", actor_id=args.decided_by, actor_role=args.role, tenant_id=approval.get("tenant_id", ""), deployment_id=approval.get("deployment_id", ""), project_id=approval.get("project_id", ""), agent_id=approval.get("agent_id", ""), resource_type="approval", resource_id=args.approval_id, workflow_run_id=approval.get("workflow_run_id", ""), task_id=approval.get("task_id", ""), approval_id=args.approval_id, outcome=status, reason=args.message or "")
    emit_notification(root, user_id=approval.get("requested_by", ""), title=f"Approval {status}", body=approval.get("title", ""), topic="approval", resource_type="approval", resource_id=args.approval_id, severity="info" if status == "approved" else "warning")
    print(json.dumps({"approval": approval, "audit_event_id": event["event_id"]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def audit_events(args: argparse.Namespace) -> int:
    root = system_root()
    path = root / "logs" / "audit-events.jsonl"
    records = []
    if path.exists():
        with path.open("r", encoding="utf-8") as handle:
            for line in handle:
                if not line.strip():
                    continue
                try:
                    event = json.loads(line)
                except json.JSONDecodeError:
                    continue
                if args.event and event.get("event") != args.event:
                    continue
                if args.resource_type and event.get("resource", {}).get("type") != args.resource_type:
                    continue
                if args.resource_id and event.get("resource", {}).get("id") != args.resource_id:
                    continue
                records.append(event)
    records = records[-args.limit :]
    print(json.dumps({"events": records}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def notification_list(args: argparse.Namespace) -> int:
    root = system_root()
    records = []
    for notification in read_records(root / "notifications", "digital-office-notification"):
        if args.user and notification.get("user_id") not in {"", args.user}:
            continue
        if args.unread_only and notification.get("status") != "unread":
            continue
        records.append(notification)
        if len(records) >= args.limit:
            break
    print(json.dumps({"notifications": records}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def notification_mark_read(args: argparse.Namespace) -> int:
    root = system_root()
    notification = read_json(notification_path(root, args.notification_id))
    notification["status"] = "read"
    notification["read_at"] = now_iso()
    write_json(notification_path(root, args.notification_id), notification)
    append_log(root, {"event": "notification_read", "notification_id": args.notification_id, "user_id": args.user or ""})
    print(json.dumps(notification, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def methodology_draft(args: argparse.Namespace) -> int:
    root = system_root()
    project_dir = ensure_project(root, args.project)
    project = read_json(project_dir / "project.json")
    entries = load_entries(root, project_dir / "knowledge" / "entries")
    stamp = dt.datetime.now().strftime("%Y%m%d%H%M%S")
    target = project_dir / "reports" / "methodology-review" / "drafts" / f"{stamp}-methodology-summary.md"
    lines = [
        f"# Methodology Summary Draft: {project.get('name')}",
        "",
        f"- Project ID: {project.get('project_id')}",
        f"- Period: {args.period}",
        f"- Created at: {now_iso()}",
        "- State: draft_created",
        "",
        "## User Review Required",
        "",
        "This draft must be shown in the GUI. The user may edit it before approval. It must not be promoted to company knowledge until the user confirms.",
        "",
        "## Evidence Inventory",
    ]
    if entries:
        for entry in entries:
            lines.append(f"- {entry.get('entry_id')}: {entry.get('title')} [{entry.get('kind')}, {entry.get('status')}]")
    else:
        lines.append("- No project knowledge entries yet.")
    lines.extend(
        [
            "",
            "## Reusable Methodology",
            "",
            "TODO: Summarize what worked, what failed, reusable workflows, decision rules, templates, and risks.",
            "",
            "## Applicability",
            "",
            "TODO: Explain which future projects should reuse this methodology and which should not.",
            "",
            "## User Edits",
            "",
            "TODO: GUI user may edit this section before approval.",
        ]
    )
    write_text(target, "\n".join(lines) + "\n")
    append_log(root, {"event": "methodology_draft", "project": args.project, "draft": str(target.relative_to(root))})
    print(str(target))
    return 0


def methodology_approve(args: argparse.Namespace) -> int:
    root = system_root()
    project_dir = ensure_project(root, args.project)
    draft = Path(args.draft)
    draft_base = (project_dir / "reports" / "methodology-review" / "drafts").resolve()
    if not draft.is_absolute():
        draft = draft_base / safe_component(args.draft.removesuffix(".md"), "draft id")
        if draft.suffix != ".md":
            draft = draft.with_suffix(".md")
    else:
        draft = draft.resolve()
        if not draft.is_relative_to(draft_base):
            print(f"office-system: methodology draft must be under {draft_base}", file=sys.stderr)
            return 2
    if not draft.exists():
        print(f"office-system: draft not found: {draft}", file=sys.stderr)
        return 2
    if not args.confirmed:
        print("office-system: approval requires --confirmed after GUI user review/edit", file=sys.stderr)
        return 2
    target = root / "knowledge" / "company" / "methodologies" / "approved" / draft.name
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(draft, target)
    append_log(root, {"event": "methodology_approve", "project": args.project, "target": str(target.relative_to(root))})
    print(str(target))
    return 0


def relay_add(args: argparse.Namespace) -> int:
    root = system_root()
    project_dir = ensure_project(root, args.project)
    agent = registered_agent(root, args.agent)
    refs = args.source_ref or []
    next_actions = args.next_action or []
    body = args.body if args.body is not None else sys.stdin.read()
    relay = {
        "version": "1.0.0",
        "relay_id": f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}",
        "project_id": args.project,
        "subproject_id": args.subproject,
        "agent": agent,
        "title": args.title,
        "summary": body.strip(),
        "status": args.status,
        "source_refs": refs,
        "next_actions": next_actions,
        "created_at": now_iso(),
        "authority": "relay_only",
        "sync_status": "pending_keymemory_sync",
    }
    outbox = project_dir / "relay" / "keymemory-outbox.jsonl"
    outbox.parent.mkdir(parents=True, exist_ok=True)
    with outbox.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(relay, ensure_ascii=False, sort_keys=True) + "\n")
    append_log(root, {"event": "relay_add", "project": args.project, "agent": agent, "relay_id": relay["relay_id"]})
    print(json.dumps(relay, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def agent_request_config_path(root: Path) -> Path:
    configured = os.environ.get("DIGITAL_OFFICE_AGENT_REQUEST_CONFIG")
    if configured:
        return Path(configured).expanduser()
    return root / "agent-requests" / "config.json"


def load_agent_request_config(root: Path) -> dict[str, Any]:
    path = agent_request_config_path(root)
    if path.exists():
        return read_json(path)
    example = root / "agent-requests" / "config.example.json"
    if example.exists():
        return read_json(example)
    return {
        "tenant_id": "",
        "server_url": "",
        "status_url": "",
        "auth_env": "DIGITAL_OFFICE_AGENT_REQUEST_TOKEN",
        "customer_visible_statuses": STATUS_LABELS,
    }


def agent_request_labels(config: dict[str, Any]) -> dict[str, str]:
    labels = dict(config.get("customer_visible_statuses", {}))
    labels.update(config.get("internal_statuses", {}))
    return labels


def send_json(url: str, payload: dict[str, Any], auth_env: str, timeout: int) -> tuple[int, str]:
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(url, data=data, method="POST")
    request.add_header("Content-Type", "application/json")
    token = os.environ.get(auth_env, "")
    if token:
        request.add_header("Authorization", f"Bearer {token}")
    with urllib.request.urlopen(request, timeout=timeout) as response:
        return response.status, response.read().decode("utf-8", errors="replace")


def agent_request_submit(args: argparse.Namespace) -> int:
    root = system_root()
    config = load_agent_request_config(root)
    if config.get("require_user_submit_confirmation", True) and not args.confirmed:
        print("office-system: agent request submission requires --confirmed after user review", file=sys.stderr)
        return 2

    body = args.body if args.body is not None else sys.stdin.read()
    request_id = f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    labels = agent_request_labels(config)
    payload = {
        "version": "1.0.0",
        "kind": "agent-plugin-request",
        "request_id": request_id,
        "tenant_id": config.get("tenant_id", ""),
        "project_id": args.project,
        "requested_by": args.requested_by,
        "title": args.title,
        "priority": args.priority,
        "body": body.strip(),
        "status": "received",
        "status_label": labels.get("received", STATUS_LABELS["received"]),
        "created_at": now_iso(),
        "source": "digital_office_secretary_agent",
        "skill_changes_requested": False,
        "notes": "This request asks the provider to design and publish an Agent plugin package. Customer-site skill add/remove/recompose operations are forbidden.",
    }
    payload = redact_obj(payload, config)

    outbox = root / "agent-requests" / "outbox" / f"{request_id}.json"
    write_json(outbox, {**payload, "sync_status": "pending_send"})

    server_url = config.get("server_url")
    sync_status = "pending_server_config"
    receipt: dict[str, Any] | None = None
    if server_url:
        try:
            status, response_body = send_json(server_url, payload, config.get("auth_env", "DIGITAL_OFFICE_AGENT_REQUEST_TOKEN"), args.timeout)
            sync_status = "sent"
            receipt = {
                "time": now_iso(),
                "request_id": request_id,
                "server_url": server_url,
                "status": status,
                "response": response_body[:2000],
            }
            write_json(root / "agent-requests" / "receipts" / f"{request_id}.json", receipt)
        except urllib.error.URLError as exc:
            sync_status = "send_failed"
            receipt = {
                "time": now_iso(),
                "request_id": request_id,
                "server_url": server_url,
                "error": str(exc),
            }
            write_json(root / "agent-requests" / "receipts" / f"{request_id}-failed.json", receipt)

    state = {
        **payload,
        "sync_status": sync_status,
        "outbox": str(outbox.relative_to(root)),
        "receipt": receipt,
    }
    write_json(root / "agent-requests" / "status" / f"{request_id}.json", state)
    append_log(root, {"event": "agent_request_submit", "request_id": request_id, "status": "received", "sync_status": sync_status})
    print(json.dumps(state, ensure_ascii=False, indent=2, sort_keys=True))
    return 0 if sync_status != "send_failed" else 1


def agent_request_set_status(args: argparse.Namespace) -> int:
    root = system_root()
    config = load_agent_request_config(root)
    labels = agent_request_labels(config)
    request_id = safe_component(args.request_id, "request id")
    status_file = root / "agent-requests" / "status" / f"{request_id}.json"
    state = read_json(status_file) if status_file.exists() else {
        "version": "1.0.0",
        "kind": "agent-plugin-request-status",
        "request_id": request_id,
        "tenant_id": config.get("tenant_id", ""),
    }
    state["status"] = args.status
    state["status_label"] = labels.get(args.status, args.status)
    state["status_updated_at"] = now_iso()
    if args.message:
        state["status_message"] = args.message
    if args.package:
        state["agent_plugin_package"] = args.package
    write_json(status_file, state)
    append_log(root, {"event": "agent_request_set_status", "request_id": request_id, "status": args.status})
    print(json.dumps(state, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def agent_request_status(args: argparse.Namespace) -> int:
    root = system_root()
    request_id = safe_component(args.request_id, "request id")
    status_file = root / "agent-requests" / "status" / f"{request_id}.json"
    if not status_file.exists():
        print(f"office-system: request status not found: {request_id}", file=sys.stderr)
        return 2
    print(json.dumps(read_json(status_file), ensure_ascii=False, indent=2, sort_keys=True))
    return 0


SKILL_CHANGE_PATTERNS = [
    re.compile(r"(?i)\b(add|remove|delete|install|uninstall|replace|compose|recompose)\b.{0,40}\bskill\b"),
    re.compile(r"(?i)\bskill\b.{0,40}\b(add|remove|delete|install|uninstall|replace|compose|recompose)\b"),
    re.compile(r"(增加|新增|删除|移除|安装|卸载|替换|重新组合|重组).{0,20}skill", re.IGNORECASE),
    re.compile(r"skill.{0,20}(增加|新增|删除|移除|安装|卸载|替换|重新组合|重组)", re.IGNORECASE),
]


def contains_skill_change(text: str) -> bool:
    return any(pattern.search(text) for pattern in SKILL_CHANGE_PATTERNS)


def agent_improvement_draft(args: argparse.Namespace) -> int:
    root = system_root()
    agent = registered_agent(root, args.agent)
    body = args.body if args.body is not None else sys.stdin.read()
    if contains_skill_change(body):
        print("office-system: existing Agent improvements may change SOUL/workflow only; skill add/remove/install/recompose is forbidden", file=sys.stderr)
        return 2
    stamp = dt.datetime.now().strftime("%Y%m%d%H%M%S")
    draft_id = f"{stamp}-{args.kind}-{slugify(args.title)}"
    target = root / "agent-improvements" / agent / "drafts" / f"{draft_id}.md"
    lines = [
        f"# Agent Improvement Draft: {args.title}",
        "",
        f"- Agent: {agent}",
        f"- Kind: {args.kind}",
        f"- Project: {args.project or ''}",
        f"- Created at: {now_iso()}",
        "- State: draft_created",
        "- Skill changes: forbidden",
        "",
        "## User-Approved Change",
        "",
        body.strip(),
        "",
        "## Activation Rule",
        "",
        "This draft may improve the existing Agent SOUL document or workflow overlay only. It must not add, remove, install, or recompose skills.",
    ]
    write_text(target, "\n".join(lines) + "\n")
    append_log(root, {"event": "agent_improvement_draft", "agent": agent, "kind": args.kind, "draft": str(target.relative_to(root))})
    print(str(target))
    return 0


def agent_improvement_approve(args: argparse.Namespace) -> int:
    root = system_root()
    agent = registered_agent(root, args.agent)
    if not args.confirmed:
        print("office-system: approving an Agent improvement requires --confirmed after user review", file=sys.stderr)
        return 2
    draft = Path(args.draft).expanduser()
    draft_base = (root / "agent-improvements" / agent / "drafts").resolve()
    if not draft.is_absolute():
        draft_name = safe_component(args.draft.removesuffix(".md"), "draft id") + ".md"
        draft = draft_base / draft_name
    else:
        draft = draft.resolve()
        if not draft.is_relative_to(draft_base):
            print(f"office-system: draft must be under {draft_base}", file=sys.stderr)
            return 2
    if not draft.exists():
        print(f"office-system: draft not found: {draft}", file=sys.stderr)
        return 2
    text = draft.read_text(encoding="utf-8", errors="replace")
    if contains_skill_change(text):
        print("office-system: draft contains a forbidden skill change request", file=sys.stderr)
        return 2
    target = root / "agent-improvements" / agent / "approved" / draft.name
    shutil.copy2(draft, target)
    append_log(root, {"event": "agent_improvement_approve", "agent": agent, "target": str(target.relative_to(root))})
    print(str(target))
    return 0


def resolve_plugin_manifest(package: str) -> tuple[Path, Path, dict[str, Any]]:
    path = Path(package).expanduser()
    if path.is_file():
        return path.parent, path, read_json(path)
    candidates = [
        path / "agent-plugin.json",
        path / "plugin.manifest.json",
        path / "agent-plugin.manifest.json",
    ]
    for candidate in candidates:
        if candidate.exists():
            return path, candidate, read_json(candidate)
    print(f"office-system: no Agent plugin manifest found in {path}", file=sys.stderr)
    raise SystemExit(2)


def agent_plugin_report(args: argparse.Namespace) -> int:
    root = system_root()
    package_dir, manifest_path, manifest = resolve_plugin_manifest(args.package)
    agent_id = manifest.get("agent_id") or manifest.get("registry_entry", {}).get("id")
    if not agent_id:
        print("office-system: Agent plugin manifest must contain agent_id", file=sys.stderr)
        return 2
    agent_id = safe_component(agent_id, "agent id")
    project_id = safe_component(args.project, "project id") if args.project else ""
    if project_id:
        ensure_project(root, project_id)
    registry = effective_agents_registry(root, include_inactive=True)
    current_agents = sorted(registry.get("agents", {}).keys())
    workflows = manifest.get("workflows", {})
    route_tests = manifest.get("route_tests", [])
    stamp = dt.datetime.now().strftime("%Y%m%d%H%M%S")
    report_id = f"{stamp}-{slugify(agent_id)}"
    target = root / "agent-plugins" / "reports" / f"{report_id}-integration-report.md"
    lines = [
        f"# New Agent Integration Report: {manifest.get('display_name', agent_id)}",
        "",
        "- State: pending_user_confirmation",
        f"- Agent ID: {agent_id}",
        f"- Package: {package_dir}",
        f"- Manifest: {manifest_path}",
        f"- Created at: {now_iso()}",
        f"- Report ID: {report_id}",
        "",
        "## Current System",
        "",
        f"- Existing agents: {', '.join(current_agents)}",
        f"- Target project: {project_id or 'global / not project-specific'}",
        "",
        "## How This Agent Will Be Added",
        "",
        "1. Copy the bundled profile template into the local profiles directory if needed.",
        "2. Add the packaged registry entry to `agent-system/agents.registry.json`.",
        "3. Add packaged workflow definitions and route tests.",
        "4. Optionally add the Agent to the selected project roster.",
        "5. Run router health checks before the Agent becomes active.",
        "",
        "## Workflow Impact",
        "",
    ]
    if workflows:
        for name, workflow in workflows.items():
            lines.append(f"- {name}: {workflow.get('label', '')}")
    else:
        lines.append("- No extra workflows declared by the package.")
    lines.extend(["", "## Route Tests", ""])
    if route_tests:
        for test in route_tests:
            lines.append(f"- `{test.get('prompt')}` -> `{test.get('expect')}`")
    else:
        lines.append("- No route tests declared by the package.")
    lines.extend(
        [
            "",
            "## User Confirmation Required",
            "",
            "The new Agent is not registered yet. Show this report in the GUI and wait for user confirmation.",
            "",
            "Available GUI actions:",
            "",
            "1. Confirm: register and deploy this Agent into the selected workflow.",
            "2. Tune Through Conversation: keep discussing and update the integration report before deployment.",
            "3. Pause: do not process now; keep the task suspended until the user returns.",
        ]
    )
    write_text(target, "\n".join(lines) + "\n")
    state = {
        "report_id": report_id,
        "request_id": safe_component(args.request_id, "request id") if args.request_id else None,
        "agent_id": agent_id,
        "status": "pending_user_confirmation",
        "status_label": STATUS_LABELS["pending_user_confirmation"],
        "package": str(package_dir),
        "manifest": str(manifest_path),
        "report": str(target.relative_to(root)),
        "project_id": project_id,
        "created_at": now_iso(),
        "allowed_actions": ["confirm", "tune", "pause"],
    }
    write_json(root / "agent-plugins" / "status" / f"{report_id}.json", state)
    if args.request_id:
        request_id = safe_component(args.request_id, "request id")
        status_file = root / "agent-requests" / "status" / f"{request_id}.json"
        request_state = read_json(status_file) if status_file.exists() else {"request_id": args.request_id}
        request_state["status"] = "pending_user_confirmation"
        request_state["status_label"] = STATUS_LABELS["pending_user_confirmation"]
        request_state["agent_plugin_report"] = str(target.relative_to(root))
        request_state["agent_plugin_report_id"] = report_id
        request_state["status_updated_at"] = now_iso()
        write_json(status_file, request_state)
    append_log(root, {"event": "agent_plugin_report", "agent": agent_id, "report": str(target.relative_to(root)), "report_id": report_id})
    print(str(target))
    return 0


def agent_plugin_decision(args: argparse.Namespace) -> int:
    root = system_root()
    report_id = safe_component(args.report_id, "report id")
    status_file = root / "agent-plugins" / "status" / f"{report_id}.json"
    if not status_file.exists():
        print(f"office-system: plugin report status not found: {report_id}", file=sys.stderr)
        return 2
    state = read_json(status_file)
    if args.decision == "confirm":
        state["status"] = "confirmed_for_activation"
        state["status_label"] = STATUS_LABELS["confirmed_for_activation"]
        state["next_action"] = f"run agent-plugin-activate --report-id {report_id} --confirmed"
        state.pop("pause_reason", None)
    elif args.decision == "tune":
        state["status"] = "needs_tuning"
        state["status_label"] = STATUS_LABELS["needs_tuning"]
        state.pop("pause_reason", None)
        if args.message:
            state.setdefault("tuning_notes", []).append({"time": now_iso(), "message": args.message})
        state["next_action"] = "secretary continues requirement conversation and regenerates the integration report"
    elif args.decision == "pause":
        state["status"] = "paused_by_user"
        state["status_label"] = STATUS_LABELS["paused_by_user"]
        if args.message:
            state["pause_reason"] = args.message
        state["next_action"] = "do nothing until the user resumes"
    state["decision_updated_at"] = now_iso()
    write_json(status_file, state)
    append_log(root, {"event": "agent_plugin_decision", "report_id": report_id, "decision": args.decision})
    print(json.dumps(state, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def agent_plugin_activate(args: argparse.Namespace) -> int:
    root = system_root()
    request_id = safe_component(args.request_id, "request id") if args.request_id else None
    project_id = safe_component(args.project, "project id") if args.project else ""
    if not args.confirmed:
        print("office-system: Agent plugin activation requires --confirmed after user review of the integration report", file=sys.stderr)
        return 2
    if not args.report_id:
        print("office-system: Agent plugin activation requires --report-id after the user confirms an integration report", file=sys.stderr)
        return 2
    report_id = safe_component(args.report_id, "report id")
    report_status_file = root / "agent-plugins" / "status" / f"{report_id}.json"
    if not report_status_file.exists():
        print(f"office-system: plugin report status not found: {report_id}", file=sys.stderr)
        return 2
    report_state = read_json(report_status_file)
    if report_state.get("status") != "confirmed_for_activation":
        print("office-system: Agent plugin activation requires a confirmed integration report", file=sys.stderr)
        return 2
    report_request = report_state.get("request_id") or None
    if request_id != report_request:
        print("office-system: activation request id must match the confirmed integration report", file=sys.stderr)
        return 2
    report_project = report_state.get("project_id") or ""
    if project_id != report_project:
        print("office-system: activation project must match the confirmed integration report", file=sys.stderr)
        return 2
    if args.report:
        report_path = Path(args.report).expanduser()
        if not report_path.is_absolute():
            report_path = root / args.report
        if not report_path.exists():
            print(f"office-system: integration report not found: {args.report}", file=sys.stderr)
            return 2
    package_dir, manifest_path, manifest = resolve_plugin_manifest(args.package)
    registry_entry = manifest.get("registry_entry")
    if not isinstance(registry_entry, dict):
        print("office-system: Agent plugin manifest must include registry_entry", file=sys.stderr)
        return 2
    agent_id = manifest.get("agent_id") or registry_entry.get("id")
    if not agent_id:
        print("office-system: Agent plugin manifest must contain agent_id", file=sys.stderr)
        return 2
    agent_id = safe_component(agent_id, "agent id")
    if report_state.get("agent_id") and report_state.get("agent_id") != agent_id:
        print("office-system: Agent plugin package does not match the confirmed integration report", file=sys.stderr)
        return 2

    registry_path_value = root / "agents.registry.json"
    registry = read_json(registry_path_value)
    if agent_id in registry.get("agents", {}) and not args.replace_agent:
        print(f"office-system: agent already registered: {agent_id}; pass --replace-agent to replace", file=sys.stderr)
        return 2

    stamp = dt.datetime.now().strftime("%Y%m%d%H%M%S")
    target_profile: Path | None = None
    profile_backup: Path | None = None
    profile_name = registry_entry.get("profile")
    if profile_name:
        profile_name = safe_component(profile_name, "profile name")
        source_profile = package_dir / "profiles" / profile_name
        target_profile = root.parent / "profiles" / profile_name
        if source_profile.exists():
            if target_profile.exists() and not args.replace_profile:
                print(f"office-system: profile already exists: {target_profile}; pass --replace-profile to replace", file=sys.stderr)
                return 2
            if target_profile.exists():
                profile_backup = target_profile.with_name(f"{target_profile.name}.bak.{stamp}")
                shutil.copytree(target_profile, profile_backup)
                shutil.rmtree(target_profile)
            shutil.copytree(source_profile, target_profile)

    backup = registry_path_value.with_suffix(f".json.bak.{stamp}")
    shutil.copy2(registry_path_value, backup)
    entry = dict(registry_entry)
    entry.pop("id", None)
    registry.setdefault("agents", {})[agent_id] = entry
    registry.setdefault("aliases", {}).update(manifest.get("aliases", {}))
    registry.setdefault("workflows", {}).update(manifest.get("workflows", {}))
    registry.setdefault("route_tests", []).extend(manifest.get("route_tests", []))
    write_json(registry_path_value, registry)

    project_file: Path | None = None
    project_backup: dict[str, Any] | None = None
    if project_id:
        project_dir = ensure_project(root, project_id)
        project_file = project_dir / "project.json"
        project = read_json(project_file)
        project_backup = json.loads(json.dumps(project))
        roster = project.setdefault("agent_roster", [])
        if agent_id not in roster:
            roster.append(agent_id)
        write_json(project_file, project)

    router = root.parent / "scripts" / "agent-router"
    health_proc = subprocess.run([str(router), "--health"], text=True, capture_output=True)
    if health_proc.returncode != 0:
        shutil.copy2(backup, registry_path_value)
        if project_file and project_backup is not None:
            write_json(project_file, project_backup)
        if target_profile and target_profile.exists():
            shutil.rmtree(target_profile)
        if profile_backup and profile_backup.exists():
            shutil.copytree(profile_backup, target_profile)
        print("office-system: router health failed after Agent registration; registry rolled back", file=sys.stderr)
        if health_proc.stdout:
            print(health_proc.stdout, file=sys.stderr)
        if health_proc.stderr:
            print(health_proc.stderr, file=sys.stderr)
        return 1

    state = {
        "agent_id": agent_id,
        "status": "activated",
        "status_label": STATUS_LABELS["downloaded_deployed"],
        "request_id": request_id,
        "report_id": report_id,
        "package": str(package_dir),
        "manifest": str(manifest_path),
        "registry_backup": str(backup),
        "project_id": project_id,
        "router_health": "passed",
    }
    if request_id:
        status_file = root / "agent-requests" / "status" / f"{request_id}.json"
        request_state = read_json(status_file) if status_file.exists() else {"request_id": request_id}
        request_state["status"] = "downloaded_deployed"
        request_state["status_label"] = STATUS_LABELS["downloaded_deployed"]
        request_state["deployed_agent_id"] = agent_id
        request_state["status_updated_at"] = now_iso()
        write_json(status_file, request_state)
    report_state["status"] = "activated"
    report_state["status_label"] = STATUS_LABELS["downloaded_deployed"]
    report_state["activated_agent_id"] = agent_id
    report_state["activated_at"] = now_iso()
    write_json(report_status_file, report_state)
    append_log(root, {"event": "agent_plugin_activate", "agent": agent_id, "project": project_id})
    print(json.dumps(state, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def consent_path(root: Path) -> Path:
    configured = os.environ.get("DIGITAL_OFFICE_CONSENT_FILE")
    if configured:
        return Path(configured).expanduser()
    return root / "data-sharing" / "consent.json"


def load_consent(root: Path) -> dict[str, Any]:
    path = consent_path(root)
    if path.exists():
        return read_json(path)
    example = root / "data-sharing" / "consent.example.json"
    if example.exists():
        return read_json(example)
    return {"enabled": False, "allowed_scopes": {}}


REDACTION_PATTERNS = {
    "email": re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"),
    "phone_cn": re.compile(r"\b1[3-9]\d{9}\b"),
    "id_card_cn": re.compile(r"\b\d{17}[\dXx]\b"),
    "api_key": re.compile(r"\b(?:sk|gho|ghp|xoxb|hf)_[A-Za-z0-9_\-]{16,}\b"),
    "bearer_token": re.compile(r"Bearer\s+[A-Za-z0-9._\-]{16,}", re.IGNORECASE),
    "secret_like": re.compile(r"(?i)(api[_-]?key|secret|token|password)\s*[:=]\s*[^\s,;]{6,}"),
}


def redact_text(text: str, consent: dict[str, Any]) -> str:
    redaction = consent.get("redaction", {})
    if not redaction.get("enabled", True):
        return text
    replacement = redaction.get("replace_with", "[REDACTED]")
    for name in redaction.get("patterns", []):
        pattern = REDACTION_PATTERNS.get(name)
        if pattern:
            text = pattern.sub(replacement, text)
    return text


def redact_obj(value: Any, consent: dict[str, Any]) -> Any:
    if isinstance(value, str):
        return redact_text(value, consent)
    if isinstance(value, list):
        return [redact_obj(item, consent) for item in value]
    if isinstance(value, dict):
        return {key: redact_obj(item, consent) for key, item in value.items()}
    return value


def read_jsonl(path: Path, limit: int | None = None) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if not path.exists():
        return rows
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            if line.strip():
                try:
                    rows.append(json.loads(line))
                except json.JSONDecodeError:
                    continue
    return rows[-limit:] if limit else rows


def count_jsonl_records(path: Path) -> int:
    if not path.exists():
        return 0
    count = 0
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            if line.strip():
                count += 1
    return count


def approved_methodology_summaries(root: Path, consent: dict[str, Any]) -> list[dict[str, Any]]:
    allowed = (
        consent.get("industry_experience_sharing", {}).get("enabled", False)
        and consent.get("allowed_scopes", {}).get("approved_methodologies", False)
    )
    if not allowed:
        return []
    rows: list[dict[str, Any]] = []
    for path in sorted((root / "knowledge" / "company" / "methodologies" / "approved").glob("*.md")):
        text = path.read_text(encoding="utf-8", errors="replace")
        rows.append(
            {
                "path": str(path.relative_to(root)),
                "sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
                "text": redact_text(text, consent),
            }
        )
    return rows


def relay_summaries(root: Path, consent: dict[str, Any]) -> list[dict[str, Any]]:
    allowed = (
        consent.get("industry_experience_sharing", {}).get("enabled", False)
        and consent.get("allowed_scopes", {}).get("project_relay_summaries", False)
    )
    if not allowed:
        return []
    rows: list[dict[str, Any]] = []
    for outbox in sorted((root / "projects").glob("*/relay/keymemory-outbox.jsonl")):
        for row in read_jsonl(outbox):
            sanitized = redact_obj(row, consent)
            project_id = str(sanitized.get("project_id", ""))
            subproject_id = str(sanitized.get("subproject_id", ""))
            source_refs = sanitized.get("source_refs", []) or []
            sanitized["project_hash"] = hashlib.sha256(project_id.encode("utf-8")).hexdigest() if project_id else ""
            sanitized["subproject_hash"] = hashlib.sha256(subproject_id.encode("utf-8")).hexdigest() if subproject_id else ""
            sanitized["source_ref_hashes"] = [
                hashlib.sha256(str(ref).encode("utf-8")).hexdigest() for ref in source_refs
            ]
            sanitized.pop("project_id", None)
            sanitized.pop("subproject_id", None)
            sanitized.pop("source_refs", None)
            rows.append(sanitized)
    return rows


def telemetry_payload(root: Path, consent: dict[str, Any]) -> dict[str, Any]:
    allowed = consent.get("allowed_scopes", {})
    health_data = {
        "agents_registry": (root / "agents.registry.json").exists(),
        "knowledge_registry": (root / "knowledge.registry.json").exists(),
        "rules_registry": (root / "rules" / "rules.registry.json").exists(),
        "judgment_policy": (root / "judgment.policy.json").exists(),
        "multimodal_pipeline": (root / "multimodal.pipeline.json").exists(),
        "rag_pipeline": (root / "rag.pipeline.json").exists(),
        "tesseract": bool(shutil.which("tesseract")),
        "pdftotext": bool(shutil.which("pdftotext")),
        "sentence_transformers": bool(importlib.util.find_spec("sentence_transformers")),
    }
    payload: dict[str, Any] = {
        "version": "1.0.0",
        "tenant_id": consent.get("tenant_id"),
        "exported_at": now_iso(),
        "mode": consent.get("mode"),
        "experience_telemetry": consent.get("experience_telemetry", {}),
        "industry_experience_sharing": consent.get("industry_experience_sharing", {}),
        "contains_raw_documents": false_value(),
        "contains_raw_images": false_value(),
        "contains_raw_keymemory_records": false_value(),
        "health": health_data if allowed.get("system_health", True) else {},
        "model_capability_status": health_data if allowed.get("model_capability_status", True) else {},
        "route_events": read_jsonl(root / "logs" / "router-events.jsonl", limit=500) if allowed.get("agent_routes", True) else [],
        "office_events": read_jsonl(root / "logs" / "office-system.jsonl", limit=500) if allowed.get("workflow_metrics", True) else [],
        "approved_methodologies": approved_methodology_summaries(root, consent),
        "project_relay_summaries": relay_summaries(root, consent),
    }
    return redact_obj(payload, consent)


def false_value() -> bool:
    return False


def telemetry_status(args: argparse.Namespace) -> int:
    root = system_root()
    consent = load_consent(root)
    status = {
        "consent_file": str(consent_path(root)),
        "enabled": consent.get("enabled", False),
        "server_configured": bool(consent.get("server_url")),
        "experience_telemetry_enabled": consent.get("experience_telemetry", {}).get("enabled", False),
        "industry_experience_sharing_enabled": consent.get("industry_experience_sharing", {}).get("enabled", False),
        "allowed_scopes": consent.get("allowed_scopes", {}),
    }
    print(json.dumps(status, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def telemetry_export(args: argparse.Namespace) -> int:
    root = system_root()
    consent = load_consent(root)
    if not consent.get("enabled", False):
        print("office-system: telemetry disabled by consent file", file=sys.stderr)
        return 2
    payload = telemetry_payload(root, consent)
    stamp = dt.datetime.now().strftime("%Y%m%d%H%M%S")
    target = root / "data-sharing" / "exports" / f"{stamp}-telemetry-bundle.json"
    write_json(target, payload)
    append_log(root, {"event": "telemetry_export", "target": str(target.relative_to(root))})
    print(str(target))
    return 0


def telemetry_send(args: argparse.Namespace) -> int:
    root = system_root()
    consent = load_consent(root)
    if not consent.get("enabled", False):
        print("office-system: telemetry disabled by consent file", file=sys.stderr)
        return 2
    if consent.get("review", {}).get("require_admin_review_before_send", True) and not args.confirmed:
        print("office-system: telemetry-send requires --confirmed after admin review", file=sys.stderr)
        return 2
    server_url = consent.get("server_url")
    if not server_url:
        print("office-system: telemetry server_url is not configured", file=sys.stderr)
        return 2
    token = os.environ.get(consent.get("auth_env", "DIGITAL_OFFICE_TELEMETRY_TOKEN"), "")
    bundle = Path(args.bundle).expanduser()
    if not bundle.is_absolute():
        bundle = root / "data-sharing" / "exports" / args.bundle
    if not bundle.exists():
        print(f"office-system: bundle not found: {bundle}", file=sys.stderr)
        return 2
    data = bundle.read_bytes()
    request = urllib.request.Request(server_url, data=data, method="POST")
    request.add_header("Content-Type", "application/json")
    if token:
        request.add_header("Authorization", f"Bearer {token}")
    try:
        with urllib.request.urlopen(request, timeout=args.timeout) as response:
            body = response.read().decode("utf-8", errors="replace")
            receipt = {
                "time": now_iso(),
                "bundle": str(bundle.relative_to(root)) if bundle.is_relative_to(root) else str(bundle),
                "server_url": server_url,
                "status": response.status,
                "response": body[:2000],
            }
    except urllib.error.URLError as exc:
        print(f"office-system: telemetry send failed: {exc}", file=sys.stderr)
        return 1
    target = root / "data-sharing" / "receipts" / f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-receipt.json"
    write_json(target, receipt)
    append_log(root, {"event": "telemetry_send", "receipt": str(target.relative_to(root))})
    print(json.dumps(receipt, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def identity_context(args: argparse.Namespace) -> int:
    root = system_root()
    project = safe_component(args.project, "project id") if args.project else ""
    agent = registered_agent(root, args.agent) if args.agent else ""
    if project:
        ensure_project(root, project)
    claims = {
        "version": "1.0.0",
        "kind": "digital-office-identity-context",
        "tenant_id": safe_claim(args.tenant, "tenant id"),
        "deployment_id": safe_claim(args.deployment, "deployment id"),
        "user_id": safe_claim(args.user, "user id"),
        "user_role": safe_claim(args.role, "user role"),
        "project_id": project,
        "agent_id": agent,
        "workflow_run_id": safe_claim(args.workflow_run, "workflow run id", required=False),
        "session_id": safe_claim(args.session, "session id", required=False),
        "created_at": now_iso(),
    }
    append_log(root, {"event": "identity_context", "tenant_id": claims["tenant_id"], "deployment_id": claims["deployment_id"], "user_id": claims["user_id"], "project": project, "agent": agent})
    print(json.dumps(claims, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


CUSTOMER_OWNED_MOUNT_TARGETS = {"company_knowledge", "project_knowledge", "agent_specialist_context"}
PROVIDER_SOLD_MOUNT_TARGETS = {"licensed_company_reference", "licensed_project_reference", "licensed_agent_reference"}


def knowledge_source_mount(args: argparse.Namespace) -> int:
    root = system_root()
    source_class = args.source_class
    mount_target = args.mount_target
    project = safe_component(args.project, "project id") if args.project else ""
    agent = registered_agent(root, args.agent) if args.agent else ""
    if project:
        ensure_project(root, project)
    if mount_target in {"project_knowledge", "licensed_project_reference"} and not project:
        print("office-system: project-scoped knowledge mounts require --project", file=sys.stderr)
        return 2
    if mount_target in {"agent_specialist_context", "licensed_agent_reference"} and not agent:
        print("office-system: agent-scoped knowledge mounts require --agent", file=sys.stderr)
        return 2
    if source_class == "customer_owned_external_kb" and mount_target not in CUSTOMER_OWNED_MOUNT_TARGETS:
        print("office-system: customer-owned external KB cannot be mounted as licensed industry reference", file=sys.stderr)
        return 2
    if source_class == "provider_sold_industry_kb" and mount_target not in PROVIDER_SOLD_MOUNT_TARGETS:
        print("office-system: provider-sold industry KB must be mounted as licensed reference, not company/project source storage", file=sys.stderr)
        return 2
    mount_id = args.mount_id or f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{slugify(args.source_id)}"
    mount_id = safe_component(mount_id, "mount id")
    provider_sold = source_class == "provider_sold_industry_kb"
    record = {
        "version": "1.0.0",
        "kind": "digital-office-knowledge-source-mount",
        "mount_id": mount_id,
        "source_class": source_class,
        "source_id": safe_claim(args.source_id, "source id"),
        "display_name": safe_claim(args.display_name or args.source_id, "display name"),
        "provider": safe_claim(args.provider, "provider", required=False),
        "tenant_id": safe_claim(args.tenant, "tenant id"),
        "deployment_id": safe_claim(args.deployment, "deployment id"),
        "created_by": safe_claim(args.created_by, "created by"),
        "mount_target": mount_target,
        "project_id": project,
        "agent_id": agent,
        "entitlement_id": safe_claim(args.entitlement, "entitlement id", required=provider_sold),
        "license_sku": safe_claim(args.license_sku, "license sku", required=False),
        "allowed_users": [safe_claim(item, "allowed user") for item in (args.allowed_user or [])],
        "allowed_roles": [safe_claim(item, "allowed role") for item in (args.allowed_role or [])],
        "inside_digital_office_only": provider_sold,
        "download_allowed": not provider_sold,
        "export_allowed": not provider_sold,
        "plain_source_files_on_customer_host": not provider_sold,
        "retrieval_mode": "controlled_remote_retrieval" if provider_sold else args.sync_mode,
        "created_at": now_iso(),
        "status": "active",
    }
    target = root / "knowledge" / "mounts" / f"{mount_id}.json"
    if target.exists() and not args.replace:
        print(f"office-system: knowledge mount already exists: {mount_id}; pass --replace to replace", file=sys.stderr)
        return 2
    write_json(target, record)
    append_log(root, {"event": "knowledge_source_mounted", "mount_id": mount_id, "source_class": source_class, "mount_target": mount_target, "tenant_id": record["tenant_id"], "deployment_id": record["deployment_id"], "created_by": record["created_by"]})
    print(json.dumps(record, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def load_mount_record(root: Path, mount_id: str) -> dict[str, Any]:
    mount_id = safe_component(mount_id, "mount id")
    path = root / "knowledge" / "mounts" / f"{mount_id}.json"
    if not path.exists():
        print(f"office-system: knowledge mount not found: {mount_id}", file=sys.stderr)
        raise SystemExit(2)
    return read_json(path)


def validate_mount_access(args: argparse.Namespace, mount: dict[str, Any], project: str, agent: str) -> int:
    if mount.get("status") != "active":
        print("office-system: knowledge mount is not active", file=sys.stderr)
        return 2
    if mount.get("source_class") != args.source_class:
        print("office-system: knowledge access source class does not match mount", file=sys.stderr)
        return 2
    if mount.get("source_id") != args.source_id:
        print("office-system: knowledge access source id does not match mount", file=sys.stderr)
        return 2
    if mount.get("project_id") and mount.get("project_id") != project:
        print("office-system: knowledge access project does not match mount scope", file=sys.stderr)
        return 2
    if mount.get("agent_id") and mount.get("agent_id") != agent:
        print("office-system: knowledge access Agent does not match mount scope", file=sys.stderr)
        return 2

    allowed_users = mount.get("allowed_users") or []
    allowed_roles = mount.get("allowed_roles") or []
    if args.decision == "allow":
        if allowed_users and args.user not in allowed_users:
            print("office-system: user is not allowed by this knowledge mount", file=sys.stderr)
            return 2
        if allowed_roles and args.role not in allowed_roles:
            print("office-system: role is not allowed by this knowledge mount", file=sys.stderr)
            return 2
        if args.source_class == "provider_sold_industry_kb":
            if not args.entitlement:
                print("office-system: provider-sold industry knowledge access requires --entitlement", file=sys.stderr)
                return 2
            if not args.knowledge_pack:
                print("office-system: provider-sold industry knowledge access requires --knowledge-pack", file=sys.stderr)
                return 2
            if args.knowledge_pack != mount.get("source_id"):
                print("office-system: knowledge pack does not match knowledge mount", file=sys.stderr)
                return 2
            if mount.get("entitlement_id") and mount.get("entitlement_id") != args.entitlement:
                print("office-system: entitlement does not match knowledge mount", file=sys.stderr)
                return 2
    return 0


def knowledge_access_log(args: argparse.Namespace) -> int:
    root = system_root()
    project = safe_component(args.project, "project id") if args.project else ""
    agent = registered_agent(root, args.agent) if args.agent else ""
    if project:
        ensure_project(root, project)
    mount_id = safe_component(args.mount_id, "mount id")
    mount = load_mount_record(root, mount_id)
    validation = validate_mount_access(args, mount, project, agent)
    if validation != 0:
        return validation
    query_hash = hashlib.sha256((args.query or "").encode("utf-8")).hexdigest() if args.query else ""
    event = {
        "time": now_iso(),
        "event": "knowledge_access",
        "tenant_id": safe_claim(args.tenant, "tenant id"),
        "deployment_id": safe_claim(args.deployment, "deployment id"),
        "user_id": safe_claim(args.user, "user id"),
        "user_role": safe_claim(args.role, "user role"),
        "project_id": project,
        "agent_id": agent,
        "workflow_run_id": safe_claim(args.workflow_run, "workflow run id", required=False),
        "source_class": args.source_class,
        "source_id": safe_claim(args.source_id, "source id"),
        "mount_id": mount_id,
        "knowledge_pack_id": safe_claim(args.knowledge_pack, "knowledge pack id", required=False),
        "entitlement_id": safe_claim(args.entitlement, "entitlement id", required=False),
        "query_hash": query_hash,
        "result_source_ids": [safe_claim(item, "result source id") for item in (args.result_source_id or [])],
        "snippet_count": args.snippet_count,
        "decision": args.decision,
        "deny_reason": safe_claim(args.deny_reason, "deny reason", required=False),
    }
    log = root / "logs" / "knowledge-access.jsonl"
    log.parent.mkdir(parents=True, exist_ok=True)
    with log.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(event, ensure_ascii=False, sort_keys=True) + "\n")
    append_log(root, {"event": "knowledge_access_log", "mount_id": event["mount_id"], "source_class": args.source_class, "decision": args.decision, "tenant_id": event["tenant_id"], "deployment_id": event["deployment_id"], "user_id": event["user_id"]})
    print(json.dumps(event, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def knowledge_spaces_root(root: Path) -> Path:
    return root / "knowledge" / "spaces"


def knowledge_space_id_for(space_type: str, *, owner: str = "", project: str = "", team: str = "", workflow_run: str = "") -> str:
    space_type = safe_component(space_type, "knowledge space type")
    if space_type not in KNOWLEDGE_SPACE_TYPES:
        print(f"office-system: invalid knowledge space type: {space_type}", file=sys.stderr)
        raise SystemExit(2)
    if space_type == "personal":
        owner = safe_claim(owner, "space owner")
        return safe_component(f"personal-{slugify(owner)}", "knowledge space id")
    if space_type == "project":
        project = safe_component(project, "project id")
        return safe_component(f"project-{project}", "knowledge space id")
    if space_type == "team":
        team = safe_component(team, "team id")
        return safe_component(f"team-{team}", "knowledge space id")
    if space_type == "workflow_artifacts":
        workflow_run = safe_component(workflow_run, "workflow run id")
        return safe_component(f"workflow-{workflow_run}", "knowledge space id")
    if space_type == "shared_with_me":
        owner = safe_claim(owner, "shared-with-me user")
        return safe_component(f"shared-{slugify(owner)}", "knowledge space id")
    return "company"


def knowledge_space_dir(root: Path, space_id: str) -> Path:
    return knowledge_spaces_root(root) / safe_component(space_id, "knowledge space id")


def knowledge_folder_path(root: Path, space_id: str, folder_id: str) -> Path:
    return knowledge_space_dir(root, space_id) / "folders" / f"{safe_component(folder_id, 'folder id')}.json"


def knowledge_item_path(root: Path, space_id: str, item_id: str) -> Path:
    return knowledge_space_dir(root, space_id) / "items" / f"{safe_component(item_id, 'item id')}.json"


def knowledge_share_path(root: Path, space_id: str, share_id: str) -> Path:
    return knowledge_space_dir(root, space_id) / "shares" / f"{safe_component(share_id, 'share id')}.json"


def ensure_knowledge_space(
    root: Path,
    *,
    space_type: str,
    owner: str = "",
    project: str = "",
    team: str = "",
    workflow_run: str = "",
    created_by: str = "",
) -> dict[str, Any]:
    space_id = knowledge_space_id_for(space_type, owner=owner, project=project, team=team, workflow_run=workflow_run)
    if space_type == "project":
        ensure_project(root, project)
    if space_type == "workflow_artifacts" and not run_record_path(root, workflow_run).exists():
        print(f"office-system: workflow run not found: {workflow_run}", file=sys.stderr)
        raise SystemExit(2)
    if space_type == "shared_with_me":
        return {
            "version": "1.0.0",
            "kind": "digital-office-knowledge-space",
            "space_id": space_id,
            "space_type": "shared_with_me",
            "owner_user_id": safe_claim(owner, "shared-with-me user"),
            "virtual": True,
        }
    directory = knowledge_space_dir(root, space_id)
    record_path = directory / "space.json"
    if record_path.exists():
        return read_json(record_path)
    record = {
        "version": "1.0.0",
        "kind": "digital-office-knowledge-space",
        "space_id": space_id,
        "space_type": space_type,
        "owner_user_id": safe_claim(owner, "space owner", required=space_type == "personal"),
        "project_id": safe_component(project, "project id") if project else "",
        "team_id": safe_component(team, "team id") if team else "",
        "workflow_run_id": safe_component(workflow_run, "workflow run id") if workflow_run else "",
        "default_visibility": "private" if space_type == "personal" else space_type,
        "created_by": safe_claim(created_by, "created by", required=False),
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    write_json(record_path, record)
    root_folder = {
        "version": "1.0.0",
        "kind": "digital-office-knowledge-folder",
        "space_id": space_id,
        "folder_id": "root",
        "parent_folder_id": "",
        "title": "Root",
        "created_by": record["created_by"],
        "created_at": record["created_at"],
        "updated_at": record["updated_at"],
    }
    write_json(knowledge_folder_path(root, space_id, "root"), root_folder)
    append_log(root, {"event": "knowledge_space_created", "space_id": space_id, "space_type": space_type, "created_by": record["created_by"]})
    return record


def load_knowledge_space(root: Path, space_id: str) -> dict[str, Any]:
    path = knowledge_space_dir(root, space_id) / "space.json"
    if not path.exists():
        print(f"office-system: knowledge space not found: {space_id}", file=sys.stderr)
        raise SystemExit(2)
    return read_json(path)


def load_knowledge_folders(root: Path, space_id: str) -> list[dict[str, Any]]:
    return read_records(knowledge_space_dir(root, space_id) / "folders", "digital-office-knowledge-folder")


def load_knowledge_items(root: Path, space_id: str) -> list[dict[str, Any]]:
    return read_records(knowledge_space_dir(root, space_id) / "items", "digital-office-knowledge-item")


def load_knowledge_shares(root: Path, space_id: str) -> list[dict[str, Any]]:
    return read_records(knowledge_space_dir(root, space_id) / "shares", "digital-office-knowledge-share")


def space_from_args(root: Path, args: argparse.Namespace, *, create: bool = True) -> dict[str, Any]:
    space_type = args.space_type
    owner = getattr(args, "owner", "") or getattr(args, "user", "") or getattr(args, "created_by", "") or getattr(args, "shared_by", "")
    project = getattr(args, "project", "") or ""
    team = getattr(args, "team", "") or ""
    workflow_run = getattr(args, "workflow_run", "") or ""
    if create:
        return ensure_knowledge_space(root, space_type=space_type, owner=owner, project=project, team=team, workflow_run=workflow_run, created_by=getattr(args, "created_by", "") or getattr(args, "shared_by", "") or getattr(args, "user", ""))
    space_id = knowledge_space_id_for(space_type, owner=owner, project=project, team=team, workflow_run=workflow_run)
    if space_type == "shared_with_me":
        return ensure_knowledge_space(root, space_type=space_type, owner=owner, project=project, team=team, workflow_run=workflow_run, created_by="")
    return load_knowledge_space(root, space_id)


def target_matches_share(share: dict[str, Any], *, user: str, role: str, agent: str, project: str, workflow_run: str) -> bool:
    target_type = share.get("target_type")
    target_id = share.get("target_id")
    return (
        (target_type == "user" and target_id == user)
        or (target_type == "role" and target_id == role)
        or (target_type == "agent" and target_id == agent)
        or (target_type == "project" and target_id == project)
        or (target_type == "workflow" and target_id == workflow_run)
    )


def permission_rank(permission: str) -> int:
    return {"read": 1, "write": 2, "manage": 3}.get(permission, 0)


def folder_lineage(root: Path, space_id: str, folder_id: str) -> list[str]:
    folders = {folder.get("folder_id"): folder for folder in load_knowledge_folders(root, space_id)}
    lineage: list[str] = []
    current = safe_component(folder_id, "folder id")
    seen: set[str] = set()
    while current and current not in seen and current in folders:
        lineage.append(current)
        seen.add(current)
        current = folders[current].get("parent_folder_id", "")
    if "root" not in lineage:
        lineage.append("root")
    return lineage


def knowledge_resource_targets(root: Path, space_id: str, resource_type: str, resource_id: str) -> tuple[list[str], list[str]]:
    resource_type = safe_component(resource_type, "knowledge resource type")
    resource_id = safe_component(resource_id, "knowledge resource id")
    if resource_type == "folder":
        if not knowledge_folder_path(root, space_id, resource_id).exists():
            print(f"office-system: folder not found: {resource_id}", file=sys.stderr)
            raise SystemExit(2)
        return [], folder_lineage(root, space_id, resource_id)
    if resource_type == "item":
        item_path = knowledge_item_path(root, space_id, resource_id)
        if not item_path.exists():
            print(f"office-system: item not found: {resource_id}", file=sys.stderr)
            raise SystemExit(2)
        item = read_json(item_path)
        return [resource_id], folder_lineage(root, space_id, item.get("folder_id", "root"))
    print(f"office-system: invalid knowledge resource type: {resource_type}", file=sys.stderr)
    raise SystemExit(2)


def share_applies_to_resource(share: dict[str, Any], *, item_ids: list[str], folder_ids: list[str]) -> bool:
    resource_type = share.get("resource_type")
    resource_id = share.get("resource_id")
    if resource_type == "item":
        return resource_id in item_ids
    if resource_type == "folder":
        if resource_id not in folder_ids:
            return False
        if share.get("inherit", True):
            return True
        return not item_ids and bool(folder_ids) and resource_id == folder_ids[0]
    return False


def knowledge_access_allowed(
    root: Path,
    *,
    space: dict[str, Any],
    resource_type: str,
    resource_id: str,
    permission: str,
    user: str,
    role: str,
    agent: str = "",
    project: str = "",
    workflow_run: str = "",
) -> dict[str, Any]:
    permission = safe_component(permission, "knowledge permission")
    if permission not in KNOWLEDGE_PERMISSIONS:
        print(f"office-system: invalid knowledge permission: {permission}", file=sys.stderr)
        raise SystemExit(2)
    user = safe_claim(user, "user id")
    role = safe_component(role, "user role")
    agent = registered_agent(root, agent) if agent else ""
    project = safe_component(project, "project id") if project else space.get("project_id", "")
    workflow_run = safe_component(workflow_run, "workflow run id") if workflow_run else ""
    space_id = space["space_id"]
    item_ids, folder_ids = knowledge_resource_targets(root, space_id, resource_type, resource_id)
    matching_shares = [
        share
        for share in load_knowledge_shares(root, space_id)
        if target_matches_share(share, user=user, role=role, agent=agent, project=project, workflow_run=workflow_run)
        and share_applies_to_resource(share, item_ids=item_ids, folder_ids=folder_ids)
        and permission_rank(share.get("permission", "")) >= permission_rank(permission)
    ]
    denies = [share for share in matching_shares if share.get("effect") == "deny"]
    if denies:
        return {"allowed": False, "outcome": "deny", "reasons": ["explicit deny share matched"], "matched_shares": [share.get("share_id") for share in denies]}
    allows = [share for share in matching_shares if share.get("effect") == "allow"]
    if allows:
        return {"allowed": True, "outcome": "allow", "reasons": ["explicit allow share matched"], "matched_shares": [share.get("share_id") for share in allows]}

    space_type = space.get("space_type")
    owner_user_id = space.get("owner_user_id", "")
    if space_type == "personal":
        if owner_user_id and user == owner_user_id:
            return {"allowed": True, "outcome": "allow", "reasons": ["personal space owner"], "matched_shares": []}
        return {"allowed": False, "outcome": "deny", "reasons": ["personal space is private by default"], "matched_shares": []}
    if space_type == "project":
        auth = compute_authorization_decision(root, tenant_id="local", deployment_id="local", user_id=user, user_role=role, action="knowledge.read" if permission == "read" else "knowledge.manage", resource_type="knowledge", resource_id=space_id, project_id=space.get("project_id", ""), audit=False)
        return {"allowed": auth["allowed"], "outcome": auth["outcome"], "reasons": auth["reasons"] or [f"project space {permission} allowed by role"], "matched_shares": []}
    if space_type in {"team", "company", "workflow_artifacts"}:
        action = "knowledge.read" if permission == "read" else "knowledge.manage"
        allowed = role_allows_action(role, action)
        return {"allowed": allowed, "outcome": "allow" if allowed else "deny", "reasons": [f"role {role} {'can' if allowed else 'cannot'} perform {action}"], "matched_shares": []}
    return {"allowed": False, "outcome": "deny", "reasons": ["unsupported knowledge space type"], "matched_shares": []}


def append_knowledge_acl_log(root: Path, event: dict[str, Any]) -> dict[str, Any]:
    event = {"time": now_iso(), "event": "knowledge_acl_access", **event}
    append_jsonl(root / "logs" / "knowledge-access.jsonl", event)
    append_log(root, {"event": "knowledge_acl_access", "space_id": event.get("space_id", ""), "resource_type": event.get("resource_type", ""), "resource_id": event.get("resource_id", ""), "decision": event.get("outcome", "")})
    return event


def authorize_knowledge_manage(root: Path, space: dict[str, Any], *, user: str, role: str) -> dict[str, Any]:
    user = safe_claim(user, "user id")
    role = safe_component(role, "user role")
    if space.get("space_type") == "personal" and space.get("owner_user_id") == user:
        return {"allowed": True, "outcome": "allow", "reasons": ["personal space owner"]}
    return compute_authorization_decision(root, tenant_id="local", deployment_id="local", user_id=user, user_role=role, action="knowledge.manage", resource_type="knowledge", resource_id=space.get("space_id", ""), project_id=space.get("project_id", ""), audit=False)


def knowledge_folder_create(args: argparse.Namespace) -> int:
    root = system_root()
    space = space_from_args(root, args, create=True)
    auth = authorize_knowledge_manage(root, space, user=args.created_by, role=args.role)
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    folder_id = safe_component(args.folder_id or slugify(args.title), "folder id")
    parent_id = safe_component(args.parent_folder, "parent folder id")
    if not knowledge_folder_path(root, space["space_id"], parent_id).exists():
        print(f"office-system: parent folder not found: {parent_id}", file=sys.stderr)
        return 2
    target = knowledge_folder_path(root, space["space_id"], folder_id)
    if target.exists() and not args.replace:
        print(f"office-system: folder already exists: {folder_id}; pass --replace to replace", file=sys.stderr)
        return 2
    record = {
        "version": "1.0.0",
        "kind": "digital-office-knowledge-folder",
        "space_id": space["space_id"],
        "folder_id": folder_id,
        "parent_folder_id": parent_id,
        "title": safe_claim(args.title, "folder title"),
        "created_by": safe_claim(args.created_by, "created by"),
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    write_json(target, record)
    append_audit_event(root, "knowledge_folder_created", actor_id=args.created_by, actor_role=args.role, project_id=space.get("project_id", ""), resource_type="knowledge_folder", resource_id=folder_id, outcome="created", extra={"space_id": space["space_id"]})
    print(json.dumps({"space": space, "folder": record}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def knowledge_item_add(args: argparse.Namespace) -> int:
    root = system_root()
    space = space_from_args(root, args, create=True)
    auth = authorize_knowledge_manage(root, space, user=args.created_by, role=args.role)
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    folder_id = safe_component(args.folder_id, "folder id")
    if not knowledge_folder_path(root, space["space_id"], folder_id).exists():
        print(f"office-system: folder not found: {folder_id}", file=sys.stderr)
        return 2
    item_id = safe_component(args.item_id or f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{slugify(args.title)}", "item id")
    target = knowledge_item_path(root, space["space_id"], item_id)
    if target.exists() and not args.replace:
        print(f"office-system: item already exists: {item_id}; pass --replace to replace", file=sys.stderr)
        return 2
    record = {
        "version": "1.0.0",
        "kind": "digital-office-knowledge-item",
        "space_id": space["space_id"],
        "item_id": item_id,
        "folder_id": folder_id,
        "title": safe_claim(args.title, "item title"),
        "source_ref": safe_claim(args.source_ref, "source reference"),
        "content_type": safe_claim(args.content_type or "application/octet-stream", "content type"),
        "connector_id": safe_claim(args.connector_id, "connector id", required=False),
        "snapshot_mode_default": True,
        "created_by": safe_claim(args.created_by, "created by"),
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    write_json(target, record)
    append_audit_event(root, "knowledge_item_added", actor_id=args.created_by, actor_role=args.role, project_id=space.get("project_id", ""), resource_type="knowledge_item", resource_id=item_id, outcome="created", extra={"space_id": space["space_id"], "folder_id": folder_id})
    print(json.dumps({"space": space, "item": record}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def knowledge_share(args: argparse.Namespace) -> int:
    root = system_root()
    space = space_from_args(root, args, create=True)
    auth = authorize_knowledge_manage(root, space, user=args.shared_by, role=args.role)
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    resource_type = safe_component(args.resource_type, "knowledge resource type")
    resource_id = safe_component(args.resource_id, "knowledge resource id")
    knowledge_resource_targets(root, space["space_id"], resource_type, resource_id)
    target_type = safe_component(args.target_type, "share target type")
    if target_type not in KNOWLEDGE_TARGET_TYPES:
        print(f"office-system: invalid share target type: {target_type}", file=sys.stderr)
        return 2
    permission = safe_component(args.permission, "knowledge permission")
    if permission not in KNOWLEDGE_PERMISSIONS:
        print(f"office-system: invalid knowledge permission: {permission}", file=sys.stderr)
        return 2
    share_id = safe_component(args.share_id or f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}", "share id")
    record = {
        "version": "1.0.0",
        "kind": "digital-office-knowledge-share",
        "share_id": share_id,
        "space_id": space["space_id"],
        "resource_type": resource_type,
        "resource_id": resource_id,
        "effect": args.effect,
        "permission": permission,
        "target_type": target_type,
        "target_id": safe_claim(args.target_id, "share target id"),
        "inherit": not args.no_inherit,
        "created_by": safe_claim(args.shared_by, "shared by"),
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "reason": safe_claim(args.reason, "share reason", required=False),
    }
    write_json(knowledge_share_path(root, space["space_id"], share_id), record)
    append_audit_event(root, "knowledge_share_updated", actor_id=args.shared_by, actor_role=args.role, project_id=space.get("project_id", ""), resource_type="knowledge_share", resource_id=share_id, outcome=args.effect, reason=args.reason or "", extra={"space_id": space["space_id"]})
    print(json.dumps({"space": space, "share": record}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def knowledge_access_check(args: argparse.Namespace) -> int:
    root = system_root()
    space = space_from_args(root, args, create=False)
    decision = knowledge_access_allowed(root, space=space, resource_type=args.resource_type, resource_id=args.resource_id, permission=args.permission, user=args.user, role=args.role, agent=args.agent or "", project=args.project or "", workflow_run=args.workflow_run or "")
    event = append_knowledge_acl_log(
        root,
        {
            "space_id": space["space_id"],
            "space_type": space.get("space_type", ""),
            "resource_type": args.resource_type,
            "resource_id": args.resource_id,
            "permission": args.permission,
            "user_id": args.user,
            "user_role": args.role,
            "agent_id": args.agent or "",
            "project_id": args.project or space.get("project_id", ""),
            "workflow_run_id": args.workflow_run or "",
            "outcome": decision["outcome"],
            "reasons": decision["reasons"],
        },
    )
    print(json.dumps({"space": space, "decision": decision, "access_event": event}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0 if decision["allowed"] else 2


def folder_descendants(folders: list[dict[str, Any]], folder_id: str) -> set[str]:
    children: dict[str, list[str]] = {}
    for folder in folders:
        children.setdefault(folder.get("parent_folder_id", ""), []).append(folder.get("folder_id", ""))
    seen: set[str] = set()
    stack = [folder_id]
    while stack:
        current = stack.pop()
        if current in seen:
            continue
        seen.add(current)
        stack.extend(child for child in children.get(current, []) if child)
    return seen


def knowledge_scope_resolve(args: argparse.Namespace) -> int:
    root = system_root()
    space = space_from_args(root, args, create=False)
    folders = load_knowledge_folders(root, space["space_id"])
    items = load_knowledge_items(root, space["space_id"])
    folder_id = safe_component(args.folder_id, "folder id") if args.folder_id else ""
    item_id = safe_component(args.item_id, "item id") if args.item_id else ""
    snapshot_mode = not args.live_mode
    resolved_items: list[dict[str, Any]] = []
    resolved_folders: list[dict[str, Any]] = []
    if item_id:
        candidates = [item for item in items if item.get("item_id") == item_id]
    else:
        selected_folder = folder_id or "root"
        descendant_ids = folder_descendants(folders, selected_folder)
        resolved_folders = [folder for folder in folders if folder.get("folder_id") in descendant_ids]
        candidates = [item for item in items if item.get("folder_id") in descendant_ids]
    for item in candidates:
        decision = knowledge_access_allowed(root, space=space, resource_type="item", resource_id=item.get("item_id", ""), permission="read", user=args.user, role=args.role, agent=args.agent or "", project=args.project or "", workflow_run=args.workflow_run or "")
        append_knowledge_acl_log(root, {"space_id": space["space_id"], "resource_type": "item", "resource_id": item.get("item_id", ""), "permission": "read", "user_id": args.user, "user_role": args.role, "agent_id": args.agent or "", "project_id": args.project or space.get("project_id", ""), "workflow_run_id": args.workflow_run or "", "outcome": decision["outcome"], "reasons": decision["reasons"], "source": "knowledge_scope_resolve"})
        if decision["allowed"]:
            resolved_items.append(item)
    snapshot = {
        "snapshot_id": f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}" if snapshot_mode else "",
        "mode": "snapshot" if snapshot_mode else "live",
        "created_at": now_iso() if snapshot_mode else "",
        "item_ids": [item.get("item_id") for item in resolved_items],
        "folder_ids": [folder.get("folder_id") for folder in resolved_folders],
    }
    print(json.dumps({"kind": "digital-office-knowledge-scope", "space": space, "snapshot": snapshot, "items": resolved_items, "folders": resolved_folders}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def knowledge_tree(args: argparse.Namespace) -> int:
    root = system_root()
    if args.space_type == "shared_with_me":
        spaces: list[dict[str, Any]] = []
        visible_items: list[dict[str, Any]] = []
        for space_file in sorted(knowledge_spaces_root(root).glob("*/space.json")):
            space = read_json(space_file)
            for item in load_knowledge_items(root, space["space_id"]):
                decision = knowledge_access_allowed(root, space=space, resource_type="item", resource_id=item.get("item_id", ""), permission="read", user=args.user, role=args.role, agent=args.agent or "", project=args.project or "", workflow_run=args.workflow_run or "")
                if decision["allowed"] and not (space.get("space_type") == "personal" and space.get("owner_user_id") == args.user):
                    spaces.append(space)
                    visible_items.append(item)
        unique_spaces = {space["space_id"]: space for space in spaces}
        print(json.dumps({"kind": "digital-office-knowledge-tree", "space_type": "shared_with_me", "spaces": list(unique_spaces.values()), "items": visible_items}, ensure_ascii=False, indent=2, sort_keys=True))
        return 0
    space = space_from_args(root, args, create=False)
    folders = []
    for folder in load_knowledge_folders(root, space["space_id"]):
        decision = knowledge_access_allowed(root, space=space, resource_type="folder", resource_id=folder.get("folder_id", ""), permission="read", user=args.user, role=args.role, agent=args.agent or "", project=args.project or "", workflow_run=args.workflow_run or "")
        if decision["allowed"]:
            folders.append(folder)
    items = []
    for item in load_knowledge_items(root, space["space_id"]):
        decision = knowledge_access_allowed(root, space=space, resource_type="item", resource_id=item.get("item_id", ""), permission="read", user=args.user, role=args.role, agent=args.agent or "", project=args.project or "", workflow_run=args.workflow_run or "")
        if decision["allowed"]:
            items.append(item)
    print(json.dumps({"kind": "digital-office-knowledge-tree", "space": space, "folders": folders, "items": items}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def knowledge_space_summaries(root: Path, *, user: str = "", role: str = "", limit: int = 50) -> list[dict[str, Any]]:
    summaries: list[dict[str, Any]] = []
    for space_file in sorted(knowledge_spaces_root(root).glob("*/space.json")):
        try:
            space = read_json(space_file)
        except Exception:
            continue
        folders = load_knowledge_folders(root, space["space_id"])
        items = load_knowledge_items(root, space["space_id"])
        visible = True
        if space.get("space_type") == "personal" and user and space.get("owner_user_id") != user:
            visible = False
        if user and role and items:
            visible = any(
                knowledge_access_allowed(root, space=space, resource_type="item", resource_id=item.get("item_id", ""), permission="read", user=user, role=role, project=space.get("project_id", ""))["allowed"]
                for item in items
            )
        elif space.get("space_type") == "personal" and user:
            visible = space.get("owner_user_id") == user
        if visible:
            summaries.append(
                {
                    "space_id": space["space_id"],
                    "space_type": space.get("space_type", ""),
                    "owner_user_id": space.get("owner_user_id", ""),
                    "project_id": space.get("project_id", ""),
                    "folder_count": len(folders),
                    "item_count": len(items),
                    "updated_at": space.get("updated_at", space.get("created_at", "")),
                }
            )
        if len(summaries) >= limit:
            break
    return summaries


def workbench_state(args: argparse.Namespace) -> int:
    root = system_root()
    limit = max(1, min(args.limit, 100))
    project = safe_component(args.project, "project id") if args.project else ""
    if project:
        ensure_project(root, project)
    auth = compute_authorization_decision(root, tenant_id=args.tenant, deployment_id=args.deployment, user_id=args.user, user_role=args.role, action="workbench.read", resource_type="workbench", resource_id=args.role, project_id=project, audit=True)
    if not auth["allowed"]:
        print(json.dumps({"authorization": auth, "status": "denied"}, ensure_ascii=False, indent=2, sort_keys=True))
        return 2
    runs = [run for run in read_run_records(root) if not project or run.get("project_id") == project]
    tasks = [task for task in read_records(root / "tasks", "digital-office-task") if not project or task.get("project_id") == project]
    approvals = [approval for approval in read_records(root / "approvals", "digital-office-approval") if not project or approval.get("project_id") == project]
    notifications = [item for item in read_records(root / "notifications", "digital-office-notification") if item.get("user_id") in {"", args.user}]
    if args.role not in TENANT_ADMIN_ROLES and args.role != "project_manager":
        runs = [run for run in runs if run.get("requested_by") in {"", args.user}]
        tasks = [task for task in tasks if task.get("requested_by") in {"", args.user} or task.get("assigned_user") in {"", args.user}]
        approvals = [approval for approval in approvals if approval.get("requested_by") in {"", args.user} or approval.get("approver_role") == args.role]
    agent_load: dict[str, int] = {}
    for task in tasks:
        agent = task.get("assigned_agent") or task.get("agent_id") or "unassigned"
        agent_load[agent] = agent_load.get(agent, 0) + 1
    if args.role in TENANT_ADMIN_ROLES:
        view = "owner_global"
        sections = {
            "project_health": project_summaries(root, limit),
            "blocked_workflows": compact_records([run for run in runs if run.get("status") in {"blocked", "waiting_user_confirmation", "paused_after_current_node"}], ["run_id", "status", "project_id", "agent_id", "updated_at"], limit),
            "cost_and_load": {"agent_load": dict(sorted(agent_load.items())), "queued_tasks": sum(1 for item in tasks if item.get("status") == "queued")},
            "pending_approvals": compact_records([item for item in approvals if item.get("status") == "pending"], ["approval_id", "title", "project_id", "approver_role", "updated_at"], limit),
            "knowledge_spaces": knowledge_space_summaries(root, user=args.user, role=args.role, limit=limit),
            "system_health": health_checks(root),
        }
    elif args.role == "project_manager":
        view = "project_lead"
        sections = {
            "project_progress": compact_records(runs, ["run_id", "status", "project_id", "agent_id", "workflow", "updated_at"], limit),
            "team_tasks": compact_records(tasks, ["task_id", "title", "status", "priority", "assigned_agent", "assigned_user", "updated_at"], limit),
            "blocked_items": compact_records([item for item in tasks if item.get("status") in {"blocked", "waiting_approval", "failed"}], ["task_id", "title", "status", "workflow_run_id", "updated_at"], limit),
            "knowledge_spaces": knowledge_space_summaries(root, user=args.user, role=args.role, limit=limit),
        }
    elif args.role == "professional_reviewer":
        view = "approver"
        sections = {
            "pending_approvals": compact_records([item for item in approvals if item.get("status") == "pending" and item.get("approver_role") == "professional_reviewer"], ["approval_id", "title", "risk", "project_id", "workflow_run_id", "updated_at"], limit),
            "recent_decisions": compact_records([item for item in approvals if item.get("status") != "pending"], ["approval_id", "title", "status", "updated_at"], limit),
        }
    elif args.role == "viewer":
        view = "viewer"
        visible_projects = [item for item in project_summaries(root, limit) if not project or item.get("project_id") == project]
        sections = {
            "visible_projects": visible_projects,
            "recent_outputs": compact_records([run for run in runs if run.get("status") == "completed"], ["run_id", "project_id", "agent_id", "workflow", "updated_at"], limit),
            "notifications": compact_records(notifications, ["notification_id", "title", "topic", "status", "created_at"], limit),
        }
    else:
        view = "member"
        sections = {
            "my_tasks": compact_records(tasks, ["task_id", "title", "status", "priority", "project_id", "workflow_run_id", "updated_at"], limit),
            "my_workflows": compact_records(runs, ["run_id", "status", "project_id", "agent_id", "workflow", "updated_at"], limit),
            "notifications": compact_records(notifications, ["notification_id", "title", "topic", "status", "created_at"], limit),
            "knowledge_spaces": knowledge_space_summaries(root, user=args.user, role=args.role, limit=limit),
        }
    payload = {
        "kind": "digital-office-workbench-state",
        "version": "1.0.0",
        "generated_at": now_iso(),
        "tenant_id": args.tenant,
        "deployment_id": args.deployment,
        "user_id": args.user,
        "role": args.role,
        "project_id": project,
        "view": view,
        "authorization": auth,
        "sections": sections,
    }
    print(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def loop_manifest(root: Path) -> dict[str, Any]:
    return read_json(root / "ai-native-loop.manifest.json")


def loop_run_path(root: Path, run_id: str) -> Path:
    return root / "runs" / safe_component(run_id, "run id") / "run.json"


def loop_start(args: argparse.Namespace) -> int:
    root = system_root()
    manifest = loop_manifest(root)
    project = safe_component(args.project, "project id") if args.project else ""
    agent = registered_agent(root, args.agent) if args.agent else ""
    if project:
        ensure_project(root, project)
    run_id = safe_component(args.run_id, "run id") if args.run_id else f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
    task_id = f"work-{run_id}"
    task = args.task if args.task is not None else sys.stdin.read()
    judgment_eval = evaluate_judgment(root, task=task.strip(), stage="context", agent_id=agent, workflow_run_id=run_id, action="loop.start")
    judgment_required = judgment_eval["decision"] == "pause"
    control = initial_loop_control(
        root,
        {
            "max_cycles": args.max_cycles,
            "max_stage_retries": args.max_stage_retries,
            "max_stagnant_cycles": args.max_stagnant_cycles,
            "max_duration_seconds": args.max_duration_seconds,
            "max_tool_calls": args.max_tool_calls,
            "max_model_calls": args.max_model_calls,
            "max_cost_microunits": args.max_cost_microunits,
        },
    )
    if any(int(value) < 0 for value in control.get("budgets", {}).values()):
        print("office-system: loop budgets must be non-negative", file=sys.stderr)
        return 2
    run = {
        "version": "2.0.0",
        "kind": "digital-office-loop-run",
        "run_id": run_id,
        "context_id": run_id,
        "task_id": task_id,
        "status": "waiting_human_judgment" if judgment_required else "created",
        "current_stage": "context",
        "project_id": project,
        "agent_id": agent,
        "workflow": safe_claim(args.workflow, "workflow", required=False),
        "requested_by": safe_claim(args.requested_by, "requested by", required=False),
        "task": task.strip(),
        "task_sha256": hashlib.sha256(task.encode("utf-8")).hexdigest(),
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "loop_contract_version": "2.0.0",
        "control": control,
        "stages": loop_stage_records(root),
        "hard_rules": manifest.get("hard_rules", []),
        "judgment_cases": [],
        "judgment_evaluation": judgment_eval,
        "blockers": ["human_judgment_pending"] if judgment_required else [],
        "events": [{"time": now_iso(), "event": "loop_start", "status": "waiting_human_judgment" if judgment_required else "created"}],
    }
    if args.dry_run:
        print(json.dumps(run, ensure_ascii=False, indent=2, sort_keys=True))
        return 0
    target = loop_run_path(root, run_id)
    if target.exists():
        print(f"office-system: loop run already exists: {run_id}", file=sys.stderr)
        return 2
    write_json(target, run)
    judgment_case = None
    if judgment_required:
        judgment_case = create_judgment_case(root, judgment_eval, task=task.strip(), reason="loop-start requires human judgment before execution", created_by=args.requested_by or "", created_by_role="")
    append_run_ledger_event(
        root,
        run_id,
        "loop_start",
        stage="context",
        action="loop.start",
        agent_id=agent,
        actor_id=args.requested_by or "",
        input_payload={"task": task.strip(), "project_id": project, "workflow": run["workflow"]},
        output_payload={"status": run["status"], "judgment": judgment_eval, "judgment_case_id": (judgment_case or {}).get("case_id", "")},
    )
    append_log(root, {"event": "loop_start", "run_id": run_id, "project": project, "agent": agent})
    print(json.dumps({"run_id": run_id, "path": str(target.relative_to(root)), "status": run["status"], "judgment": judgment_eval, "judgment_case": judgment_case}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def loop_stage(args: argparse.Namespace) -> int:
    root = system_root()
    run_id = safe_component(args.run_id, "run id")
    stage = normalize_loop_stage(args.stage)
    target = loop_run_path(root, run_id)
    if not target.exists():
        print(f"office-system: loop run not found: {run_id}", file=sys.stderr)
        return 2
    run = load_run_record(root, run_id)
    if run.get("status") in {"completed", "failed", "cancelled", "budget_exhausted"}:
        print(f"office-system: terminal loop cannot update stages: {run.get('status')}", file=sys.stderr)
        return 2
    open_cases = open_judgment_cases(root, run)
    if open_cases:
        print(json.dumps({"status": "blocked", "reason": "human_judgment_pending", "open_judgments": open_cases}, ensure_ascii=False, indent=2, sort_keys=True), file=sys.stderr)
        return 2
    budget_blockers = loop_budget_blockers(run)
    if budget_blockers and not (stage == "evaluate" or (stage == "act" and args.status == "completed")):
        run.setdefault("blockers", []).extend(item for item in budget_blockers if item not in run.get("blockers", []))
        run["updated_at"] = now_iso()
        write_json(target, run)
        print(json.dumps({"status": "blocked", "reason": "budget_reached", "blockers": budget_blockers, "allowed_next": ["complete current act", "evaluate", "loop-control"]}, ensure_ascii=False, indent=2, sort_keys=True), file=sys.stderr)
        return 2
    current_index = LOOP_STAGES.index(stage)
    for previous in LOOP_STAGES[:current_index]:
        if run.get("stages", {}).get(previous, {}).get("status") not in {"completed", "skipped"}:
            print(f"office-system: cannot update {stage} before {previous} is completed", file=sys.stderr)
            return 2
    stage_state = run.setdefault("stages", {}).setdefault(stage, fresh_loop_stage_state(root, stage))
    prior_status = str(stage_state.get("status", "pending"))
    if args.status == "started" and prior_status not in {"pending", "blocked", "failed", "started"}:
        print(f"office-system: cannot start {stage} from status {prior_status}", file=sys.stderr)
        return 2
    if args.status == "completed" and prior_status != "started":
        print(f"office-system: cannot complete {stage} before it is started", file=sys.stderr)
        return 2
    if args.status == "skipped":
        if stage != "decide":
            print("office-system: only the decide stage may be skipped for an explicitly simple task profile", file=sys.stderr)
            return 2
        if prior_status not in {"pending", "started"}:
            print(f"office-system: cannot skip {stage} from status {prior_status}", file=sys.stderr)
            return 2
        if not args.summary:
            print("office-system: skipping decide requires --summary with the deterministic reason", file=sys.stderr)
            return 2
    stage_state["status"] = args.status
    if args.summary:
        stage_state.setdefault("notes", []).append({"time": now_iso(), "summary": args.summary})
    body = args.body if args.body is not None else ""
    if body:
        stage_state.setdefault("notes", []).append({"time": now_iso(), "body": body})
    for artifact in args.artifact or []:
        artifact_value = safe_claim(artifact, "artifact")
        artifact_id = ""
        if ":" in artifact_value:
            artifact_id = artifact_value.split(":", 1)[0]
        elif "=" in artifact_value:
            artifact_id = artifact_value.split("=", 1)[0]
        stage_state.setdefault("artifacts", []).append({"artifact": artifact_value, "artifact_id": artifact_id or artifact_value, "time": now_iso()})
    for gate in args.gate or []:
        if ":" not in gate:
            print("office-system: --gate must be gate_id:status", file=sys.stderr)
            return 2
        gate_id, gate_status = gate.split(":", 1)
        gate_id = safe_claim(gate_id, "gate id")
        gate_status = safe_component(gate_status, "gate status")
        updated = False
        for item in stage_state.setdefault("gates", []):
            if str(item.get("gate")) == gate_id:
                item["status"] = gate_status
                item["updated_at"] = now_iso()
                updated = True
                break
        if not updated:
            print(f"office-system: unknown gate for {stage}: {gate_id}", file=sys.stderr)
            return 2
        stage_state.setdefault("gate_updates", []).append(
            {
                "time": now_iso(),
                "gate": gate_id,
                "status": gate_status,
            }
        )
    if args.status == "completed":
        missing_artifacts = [
            str(item)
            for item in stage_state.get("required_artifacts", []) or []
            if not stage_artifact_satisfies(str(item), stage_state.get("artifacts", []) or [])
        ]
        incomplete_gates = [
            str(item.get("gate", ""))
            for item in stage_state.get("gates", []) or []
            if str(item.get("status", "pending")) not in PASSING_GATE_STATUSES
        ]
        if missing_artifacts or incomplete_gates:
            stage_state["status"] = prior_status
            print(json.dumps({"status": "blocked", "stage": stage, "missing_artifacts": missing_artifacts, "incomplete_gates": incomplete_gates}, ensure_ascii=False, indent=2, sort_keys=True), file=sys.stderr)
            return 2
    if args.status == "started":
        run["status"] = LOOP_STATUS_BY_STAGE[stage]
        run["current_stage"] = stage
    elif args.status in {"failed", "blocked"}:
        run["status"] = "blocked"
        run["current_stage"] = stage
    elif args.status in {"completed", "skipped"}:
        next_stage = LOOP_STAGES[current_index + 1] if current_index + 1 < len(LOOP_STAGES) else None
        run["current_stage"] = next_stage or stage
        if next_stage:
            run["status"] = LOOP_STATUS_BY_STAGE[next_stage]
        else:
            run["status"] = "awaiting_control"
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "loop_stage", "stage": stage, "status": args.status})
    write_json(target, run)
    append_run_ledger_event(
        root,
        run_id,
        "loop_stage",
        stage=stage,
        action=f"loop.stage.{args.status}",
        agent_id=str(run.get("agent_id", "")),
        input_payload={"stage": stage, "status": args.status, "artifacts": args.artifact or [], "gates": args.gate or []},
        output_payload={"run_status": run["status"], "current_stage": run.get("current_stage", "")},
        artifact_refs=[str(item) for item in (args.artifact or [])],
    )
    append_log(root, {"event": "loop_stage", "run_id": run_id, "stage": stage, "status": args.status})
    print(json.dumps({"run_id": run_id, "stage": stage, "status": args.status, "run_status": run["status"]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def loop_status(args: argparse.Namespace) -> int:
    root = system_root()
    target = loop_run_path(root, args.run_id)
    if not target.exists():
        print(f"office-system: loop run not found: {args.run_id}", file=sys.stderr)
        return 2
    print(json.dumps(load_run_record(root, args.run_id), ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def loop_usage_add(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    additions = {
        "tool_calls": args.tool_calls,
        "model_calls": args.model_calls,
        "input_tokens": args.input_tokens,
        "output_tokens": args.output_tokens,
        "cost_microunits": args.cost_microunits,
    }
    if any(int(value) < 0 for value in additions.values()):
        print("office-system: loop usage values must be non-negative", file=sys.stderr)
        return 2
    usage = run.setdefault("control", initial_loop_control(root)).setdefault("usage", {})
    for key, value in additions.items():
        usage[key] = int(usage.get(key, 0) or 0) + int(value)
    blockers = loop_budget_blockers(run)
    if blockers:
        run.setdefault("blockers", []).extend(item for item in blockers if item not in run.get("blockers", []))
    run["updated_at"] = now_iso()
    write_json(run_record_path(root, args.run_id), run)
    append_run_ledger_event(
        root,
        args.run_id,
        "loop_usage_recorded",
        stage=normalize_loop_stage(args.stage or str(run.get("current_stage", "context"))),
        action="loop.usage.add",
        input_payload=additions,
        output_payload={"usage": usage, "budget_blockers": blockers},
    )
    payload = {
        "status": "budget_reached" if blockers else "success",
        "summary": "Loop usage recorded; no further resource-consuming transition is allowed before evaluation." if blockers else "Loop usage recorded.",
        "next_actions": ["complete current act if needed", "evaluate", "loop-control"] if blockers else ["continue current work node"],
        "artifacts": [str(run_record_path(root, args.run_id).relative_to(root))],
        "run_id": args.run_id,
        "usage": usage,
        "budget_blockers": blockers,
    }
    print(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True))
    return 2 if blockers else 0


def loop_control(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    decision = args.decision
    if decision not in LOOP_CONTROL_DECISIONS:
        print(f"office-system: unsupported loop control decision: {decision}", file=sys.stderr)
        return 2
    if run.get("status") in {"completed", "failed", "cancelled", "budget_exhausted"}:
        print(f"office-system: terminal loop cannot transition: {run.get('status')}", file=sys.stderr)
        return 2
    if decision == "cancel" and not args.confirmed:
        print("office-system: cancel requires --confirmed", file=sys.stderr)
        return 2
    progress_score = float(args.progress_score)
    if progress_score < 0.0 or progress_score > 1.0:
        print("office-system: --progress-score must be between 0 and 1", file=sys.stderr)
        return 2
    evaluate_state = run.get("stages", {}).get("evaluate", {})
    if decision not in {"cancel", "budget_exhausted"} and evaluate_state.get("status") != "completed":
        print("office-system: loop control requires a completed evaluate stage", file=sys.stderr)
        return 2
    budget_blockers = loop_budget_blockers(run)
    if budget_blockers and decision in {"continue", "replan", "retry"}:
        print(json.dumps({"status": "blocked", "reason": "budget_exhausted", "blockers": budget_blockers}, ensure_ascii=False, indent=2, sort_keys=True), file=sys.stderr)
        return 2

    control = run.setdefault("control", initial_loop_control(root))
    budgets = control.setdefault("budgets", {})
    previous_score = float(control.get("last_progress_score", 0.0) or 0.0)
    minimum_delta = float(loop_manifest(root).get("controller", {}).get("progress_policy", {}).get("minimum_progress_delta", 0.02))
    progress_delta = progress_score - previous_score
    effective_decision = decision
    target_stage = ""
    reason = safe_claim(args.reason, "loop control reason", required=False)

    if decision == "complete":
        if not args.acceptance_passed:
            print("office-system: complete requires --acceptance-passed", file=sys.stderr)
            return 2
        blockers = loop_completion_blockers(root, run, require_all_stages=True)
        if blockers:
            print(json.dumps({"status": "blocked", "reason": "completion_gates_failed", "blockers": blockers}, ensure_ascii=False, indent=2, sort_keys=True), file=sys.stderr)
            return 2
        run["status"] = "completed"
        run["blockers"] = []
    elif decision in {"continue", "replan", "retry"}:
        target_stage = {"continue": "context", "replan": "decide", "retry": "act"}[decision]
        next_cycle = int(control.get("cycle_index", 1)) + 1
        max_cycles = int(budgets.get("max_cycles", 0) or 0)
        if max_cycles > 0 and next_cycle > max_cycles:
            effective_decision = "budget_exhausted"
            run["status"] = "budget_exhausted"
            run.setdefault("blockers", []).append("max_cycles_exceeded")
        else:
            if progress_delta < minimum_delta:
                control["stagnant_cycles"] = int(control.get("stagnant_cycles", 0)) + 1
            else:
                control["stagnant_cycles"] = 0
            max_stagnant = int(budgets.get("max_stagnant_cycles", 0) or 0)
            if max_stagnant >= 0 and int(control.get("stagnant_cycles", 0)) > max_stagnant:
                effective_decision = "wait_human"
                run["status"] = "waiting_user_input"
                run.setdefault("blockers", []).append("loop_progress_stagnant")
            else:
                if decision == "retry":
                    retries = control.setdefault("stage_retries", {})
                    retries[target_stage] = int(retries.get(target_stage, 0)) + 1
                    max_retries = int(budgets.get("max_stage_retries", 0) or 0)
                    if max_retries >= 0 and retries[target_stage] > max_retries:
                        print("office-system: act retry budget exhausted; replan, wait for human input, or fail", file=sys.stderr)
                        return 2
                control.setdefault("cycle_history", []).append(
                    {
                        "cycle_index": int(control.get("cycle_index", 1)),
                        "ended_at": now_iso(),
                        "decision": decision,
                        "progress_score": progress_score,
                        "stage_statuses": {stage: run.get("stages", {}).get(stage, {}).get("status", "missing") for stage in LOOP_STAGES},
                        "stage_artifact_hashes": {stage: canonical_hash(run.get("stages", {}).get(stage, {}).get("artifacts", [])) for stage in LOOP_STAGES},
                    }
                )
                control["cycle_index"] = next_cycle
                reset_loop_from_stage(root, run, target_stage)
    elif decision == "wait_human":
        run["status"] = "waiting_user_input"
        run.setdefault("blockers", []).append("loop_human_input_required")
    elif decision == "fail":
        if not reason:
            print("office-system: fail requires --reason", file=sys.stderr)
            return 2
        run["status"] = "failed"
        run.setdefault("blockers", []).append("permanent_failure")
    elif decision == "cancel":
        run["status"] = "cancelled"
    elif decision == "budget_exhausted":
        run["status"] = "budget_exhausted"
        run.setdefault("blockers", []).extend(item for item in budget_blockers if item not in run.get("blockers", []))

    control["last_progress_score"] = progress_score
    record = {
        "time": now_iso(),
        "requested_decision": decision,
        "effective_decision": effective_decision,
        "target_stage": target_stage,
        "progress_score": progress_score,
        "progress_delta": round(progress_delta, 6),
        "acceptance_passed": bool(args.acceptance_passed),
        "reason": reason,
        "cycle_index": int(control.get("cycle_index", 1)),
    }
    control.setdefault("decision_history", []).append(record)
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "loop_control", **record})
    write_json(run_record_path(root, args.run_id), run)
    append_run_ledger_event(
        root,
        args.run_id,
        "loop_control",
        stage="evaluate",
        action=f"loop.control.{effective_decision}",
        input_payload={"decision": decision, "progress_score": progress_score, "acceptance_passed": bool(args.acceptance_passed), "reason": reason},
        output_payload={"status": run["status"], "current_stage": run.get("current_stage", ""), "record": record},
    )
    payload = {
        "status": run["status"],
        "summary": f"Loop controller applied {effective_decision}.",
        "next_actions": ["loop-status"] if run["status"] in {"completed", "failed", "cancelled", "budget_exhausted"} else [f"continue at {run.get('current_stage', 'evaluate')}"],
        "artifacts": [str(run_record_path(root, args.run_id).relative_to(root))],
        "run_id": args.run_id,
        "control_decision": record,
    }
    print(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def run_ledger_add(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    input_payload = parse_json_value(args.input_json, args.input_file, "ledger input", required=False)
    output_payload = parse_json_value(args.output_json, args.output_file, "ledger output", required=False)
    parameters = parse_json_value(args.parameters_json, args.parameters_file, "ledger parameters", required=False)
    event = append_run_ledger_event(
        root,
        args.run_id,
        args.event,
        stage=args.stage or str(run.get("current_stage", "")),
        action=args.action or "",
        agent_id=args.agent or str(run.get("agent_id", "")),
        actor_id=args.actor or "",
        actor_role=args.role or "",
        input_payload=input_payload,
        output_payload=output_payload,
        artifact_refs=args.artifact or [],
        checkpoint_id=args.checkpoint_id or "",
        handoff_id=args.handoff_id or "",
        model=args.model or "",
        provider=args.provider or "",
        parameters=parameters,
    )
    print(json.dumps(event, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def run_ledger_list(args: argparse.Namespace) -> int:
    root = system_root()
    load_run_record(root, args.run_id)
    rows = list_run_ledger(root, args.run_id, limit=args.limit)
    print(json.dumps({"run_id": args.run_id, "events": rows}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def run_ledger_verify(args: argparse.Namespace) -> int:
    root = system_root()
    load_run_record(root, args.run_id)
    result = verify_run_ledger(root, args.run_id)
    print(json.dumps(result, ensure_ascii=False, indent=2, sort_keys=True))
    return 0 if result["status"] == "valid" else 2


def checkpoint_create(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    stage = normalize_loop_stage(args.stage or str(run.get("current_stage", "context")))
    if stage not in LOOP_STAGES:
        print(f"office-system: invalid checkpoint stage: {stage}", file=sys.stderr)
        return 2
    if args.requires_human and not args.create_judgment:
        print("office-system: --requires-human checkpoints must also pass --create-judgment", file=sys.stderr)
        return 2
    if args.create_judgment and not args.requires_human:
        print("office-system: --create-judgment is only valid with --requires-human", file=sys.stderr)
        return 2
    state = parse_json_value(args.state_json, args.state_file, "checkpoint state", required=False) or {
        "run_status": run.get("status", ""),
        "current_stage": run.get("current_stage", ""),
        "blockers": run.get("blockers", []),
    }
    checkpoint_id = safe_component(args.checkpoint_id, "checkpoint id") if args.checkpoint_id else f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}-{stage}"
    if checkpoint_path(root, args.run_id, checkpoint_id).exists():
        print(f"office-system: checkpoint already exists: {checkpoint_id}", file=sys.stderr)
        return 2
    record = {
        "version": "1.0.0",
        "kind": "digital-office-run-checkpoint",
        "checkpoint_id": checkpoint_id,
        "run_id": safe_component(args.run_id, "run id"),
        "stage": stage,
        "label": safe_claim(args.label, "checkpoint label", required=False),
        "status": "requires_human_judgment" if args.requires_human else "ready_to_resume",
        "resume_cursor": safe_claim(args.resume_cursor, "resume cursor", required=False),
        "state": state,
        "state_hash": canonical_hash(state),
        "artifacts": [safe_claim(item, "checkpoint artifact") for item in (args.artifact or [])],
        "requires_human": bool(args.requires_human),
        "created_at": now_iso(),
    }
    judgment_case = None
    if args.requires_human and args.create_judgment:
        evaluation = evaluate_judgment(
            root,
            task=args.reason or record["label"] or f"Checkpoint {checkpoint_id} requires human judgment",
            stage=stage,
            agent_id=str(run.get("agent_id", "")),
            workflow_run_id=args.run_id,
            action="workflow_continue",
            signal={"decision": "pause", "risk_score": 0.74, "reason": args.reason or "checkpoint requires human judgment"},
        )
        judgment_case = create_judgment_case(root, evaluation, task=args.reason or record["label"] or checkpoint_id, reason=args.reason or "checkpoint requires human judgment", created_by=args.created_by or "", created_by_role=args.role or "")
        record["judgment_case_id"] = judgment_case["case_id"]
    write_json(checkpoint_path(root, args.run_id, checkpoint_id), record)
    run = load_run_record(root, args.run_id)
    run.setdefault("checkpoints", []).append({"checkpoint_id": checkpoint_id, "stage": stage, "path": str(checkpoint_path(root, args.run_id, checkpoint_id).relative_to(root)), "state_hash": record["state_hash"], "created_at": record["created_at"]})
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "checkpoint_created", "checkpoint_id": checkpoint_id, "stage": stage})
    write_json(run_record_path(root, args.run_id), run)
    append_run_ledger_event(
        root,
        args.run_id,
        "checkpoint_created",
        stage=stage,
        action="checkpoint.create",
        agent_id=str(run.get("agent_id", "")),
        actor_id=args.created_by or "",
        actor_role=args.role or "",
        input_payload={"label": record["label"], "resume_cursor": record["resume_cursor"], "requires_human": args.requires_human},
        output_payload={"checkpoint_id": checkpoint_id, "state_hash": record["state_hash"], "judgment_case_id": (judgment_case or {}).get("case_id", "")},
        artifact_refs=record["artifacts"],
        checkpoint_id=checkpoint_id,
    )
    print(json.dumps({"checkpoint": record, "judgment_case": judgment_case}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def checkpoint_list(args: argparse.Namespace) -> int:
    root = system_root()
    load_run_record(root, args.run_id)
    records = read_records(checkpoint_dir(root, args.run_id), "digital-office-run-checkpoint")
    print(json.dumps({"run_id": args.run_id, "checkpoints": records[: args.limit]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def forbidden_context_keys(value: Any, path: str = "$") -> list[str]:
    forbidden = {
        "chain_of_thought",
        "chain-of-thought",
        "private_reasoning",
        "reasoning_trace",
        "hidden_reasoning",
        "credential",
        "credentials",
        "password",
        "secret",
        "api_key",
        "access_token",
        "refresh_token",
    }
    found: list[str] = []
    if isinstance(value, dict):
        for key, item in value.items():
            item_path = f"{path}.{key}"
            if str(key).lower() in forbidden:
                found.append(item_path)
            found.extend(forbidden_context_keys(item, item_path))
    elif isinstance(value, list):
        for index, item in enumerate(value):
            found.extend(forbidden_context_keys(item, f"{path}[{index}]"))
    return found


def validate_json_schema_instance(value: Any, schema: dict[str, Any], root_schema: dict[str, Any], path: str = "$") -> list[str]:
    if "$ref" in schema:
        reference = str(schema["$ref"])
        if not reference.startswith("#/"):
            return [f"{path}: unsupported schema reference {reference}"]
        resolved: Any = root_schema
        try:
            for part in reference[2:].split("/"):
                resolved = resolved[part.replace("~1", "/").replace("~0", "~")]
        except (KeyError, TypeError):
            return [f"{path}: unresolved schema reference {reference}"]
        return validate_json_schema_instance(value, resolved, root_schema, path)

    errors: list[str] = []
    if "const" in schema and value != schema["const"]:
        errors.append(f"{path}: must equal {schema['const']!r}")
    if "enum" in schema and value not in schema["enum"]:
        errors.append(f"{path}: must be one of {schema['enum']}")

    expected_type = schema.get("type")
    type_ok = True
    if expected_type == "object":
        type_ok = isinstance(value, dict)
    elif expected_type == "array":
        type_ok = isinstance(value, list)
    elif expected_type == "string":
        type_ok = isinstance(value, str)
    elif expected_type == "integer":
        type_ok = isinstance(value, int) and not isinstance(value, bool)
    elif expected_type == "number":
        type_ok = isinstance(value, (int, float)) and not isinstance(value, bool)
    elif expected_type == "boolean":
        type_ok = isinstance(value, bool)
    if expected_type and not type_ok:
        return errors + [f"{path}: must be {expected_type}"]

    if isinstance(value, dict):
        properties = schema.get("properties", {})
        for key in schema.get("required", []):
            if key not in value:
                errors.append(f"{path}: missing required field {key}")
        if schema.get("additionalProperties") is False:
            for key in value:
                if key not in properties:
                    errors.append(f"{path}.{key}: unsupported field")
        for key, item in value.items():
            if key in properties:
                errors.extend(validate_json_schema_instance(item, properties[key], root_schema, f"{path}.{key}"))
    elif isinstance(value, list):
        minimum_items = schema.get("minItems")
        if isinstance(minimum_items, int) and len(value) < minimum_items:
            errors.append(f"{path}: must contain at least {minimum_items} item(s)")
        item_schema = schema.get("items")
        if isinstance(item_schema, dict):
            for index, item in enumerate(value):
                errors.extend(validate_json_schema_instance(item, item_schema, root_schema, f"{path}[{index}]"))
    elif isinstance(value, str):
        minimum_length = schema.get("minLength")
        if isinstance(minimum_length, int) and len(value) < minimum_length:
            errors.append(f"{path}: must contain at least {minimum_length} character(s)")
    elif isinstance(value, (int, float)) and not isinstance(value, bool):
        if "minimum" in schema and value < schema["minimum"]:
            errors.append(f"{path}: must be greater than or equal to {schema['minimum']}")
        if "maximum" in schema and value > schema["maximum"]:
            errors.append(f"{path}: must be less than or equal to {schema['maximum']}")
    return errors


def validate_context_envelope(
    root: Path,
    context: Any,
    *,
    run_id: str,
    run: dict[str, Any],
    from_agent: str,
    to_agent: str,
    handoff_id: str,
) -> tuple[dict[str, Any], list[str]]:
    if not isinstance(context, dict):
        return {}, ["context must be a JSON object"]
    normalized = copy.deepcopy(context)
    normalized["handoff_id"] = handoff_id
    if normalized.get("current_stage"):
        normalized["current_stage"] = normalize_loop_stage(str(normalized["current_stage"]))
    schema = read_json(root / "context-envelope.schema.json")
    errors = validate_json_schema_instance(normalized, schema, schema)
    if normalized.get("kind") != "digital-office-context-envelope":
        errors.append("kind must be digital-office-context-envelope")
    if normalized.get("version") != "2.0.0":
        errors.append("version must be 2.0.0")
    if normalized.get("run_id") != run_id:
        errors.append("run_id does not match the handoff run")
    for key in ("context_id", "task_id", "goal", "summary", "handoff_reason", "state_hash"):
        if key in normalized and not str(normalized.get(key, "")).strip():
            errors.append(f"{key} must not be empty")
    if len(str(normalized.get("state_hash", ""))) < 12:
        errors.append("state_hash must contain at least 12 characters")
    for key, minimum in (("context_version", 1), ("cycle_index", 0)):
        value = normalized.get(key)
        if not isinstance(value, int) or isinstance(value, bool) or value < minimum:
            errors.append(f"{key} must be an integer greater than or equal to {minimum}")
    if normalized.get("current_stage") not in LOOP_STAGES:
        errors.append(f"current_stage must be one of: {', '.join(LOOP_STAGES)}")
    for key, expected_agent in (("from", from_agent), ("to", to_agent)):
        actor = normalized.get(key)
        if not isinstance(actor, dict) or actor.get("id") != expected_agent:
            errors.append(f"{key}.id must match {expected_agent}")
        elif actor.get("type") not in {"secretary", "agent"}:
            errors.append(f"{key}.type must be secretary or agent for Agent handoff")
    list_fields = (
        "constraints",
        "acceptance_criteria",
        "facts",
        "source_refs",
        "artifact_refs",
        "decisions",
        "open_questions",
        "omissions",
        "risk_flags",
    )
    for key in list_fields:
        if key in normalized and not isinstance(normalized[key], list):
            errors.append(f"{key} must be an array")
    if isinstance(normalized.get("acceptance_criteria"), list) and not normalized["acceptance_criteria"]:
        errors.append("acceptance_criteria must contain at least one criterion")
    for index, fact in enumerate(normalized.get("facts", []) if isinstance(normalized.get("facts"), list) else []):
        if not isinstance(fact, dict):
            errors.append(f"facts[{index}] must be an object")
            continue
        if fact.get("status") not in {"verified", "assumption", "uncertain", "conflict"}:
            errors.append(f"facts[{index}].status is invalid")
        confidence = fact.get("confidence")
        if not isinstance(confidence, (int, float)) or isinstance(confidence, bool) or not 0 <= confidence <= 1:
            errors.append(f"facts[{index}].confidence must be between 0 and 1")
        for key in ("fact_id", "statement", "basis_refs"):
            if key not in fact:
                errors.append(f"facts[{index}] is missing {key}")
    for collection in ("source_refs", "artifact_refs"):
        for index, ref in enumerate(normalized.get(collection, []) if isinstance(normalized.get(collection), list) else []):
            if not isinstance(ref, dict):
                errors.append(f"{collection}[{index}] must be an object")
                continue
            for key in ("ref_id", "type", "uri", "authority", "retrievable"):
                if key not in ref:
                    errors.append(f"{collection}[{index}] is missing {key}")
            if "retrievable" in ref and not isinstance(ref["retrievable"], bool):
                errors.append(f"{collection}[{index}].retrievable must be boolean")
    permissions = normalized.get("permissions")
    if not isinstance(permissions, dict):
        errors.append("permissions must be an object")
    else:
        for key in ("tenant_id", "project_id", "requested_by", "role", "allowed_actions"):
            if key not in permissions:
                errors.append(f"permissions is missing {key}")
    budget = normalized.get("context_budget")
    if not isinstance(budget, dict):
        errors.append("context_budget must be an object")
    else:
        if budget.get("strategy") not in {"inline", "reference", "hybrid"}:
            errors.append("context_budget.strategy is invalid")
        estimated = budget.get("estimated_tokens")
        maximum = budget.get("max_tokens")
        if not isinstance(estimated, int) or isinstance(estimated, bool) or estimated < 0:
            errors.append("context_budget.estimated_tokens must be a non-negative integer")
        if not isinstance(maximum, int) or isinstance(maximum, bool) or maximum < 1:
            errors.append("context_budget.max_tokens must be a positive integer")
        if isinstance(estimated, int) and isinstance(maximum, int) and estimated > maximum:
            errors.append("context envelope exceeds context_budget.max_tokens")
        if not isinstance(budget.get("compacted"), bool):
            errors.append("context_budget.compacted must be boolean")
    forbidden = forbidden_context_keys(normalized)
    errors.extend(f"forbidden sensitive or private reasoning field: {item}" for item in forbidden)
    run_context_id = str(run.get("context_id", ""))
    if run_context_id and normalized.get("context_id") != run_context_id:
        errors.append("context_id does not match the run context_id")
    run_task_id = str(run.get("task_id", ""))
    if run_task_id and normalized.get("task_id") != run_task_id:
        errors.append("task_id does not match the run task_id")
    if isinstance(permissions, dict):
        for context_key, run_key in (("tenant_id", "tenant_id"), ("project_id", "project_id"), ("requested_by", "requested_by"), ("role", "requested_by_role")):
            run_value = str(run.get(run_key, ""))
            if run_value and permissions.get(context_key) != run_value:
                errors.append(f"permissions.{context_key} does not match run {run_key}")
    return normalized, errors


def handoff_create(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    from_agent = registered_agent(root, args.from_agent)
    to_agent = registered_agent(root, args.to_agent)
    if from_agent == to_agent:
        print("office-system: handoff requires different source and target agents", file=sys.stderr)
        return 2
    input_schema = parse_json_value(args.input_schema_json, args.input_schema_file, "handoff input schema", required=False)
    context = parse_json_value(args.context_json, args.context_file, "handoff context", required=True)
    acceptance = [safe_claim(item, "acceptance criterion") for item in (args.acceptance_criterion or [])]
    if not acceptance:
        print("office-system: handoff requires at least one --acceptance-criterion", file=sys.stderr)
        return 2
    handoff_id = safe_component(args.handoff_id, "handoff id") if args.handoff_id else f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}-{from_agent}-to-{to_agent}"
    if handoff_path(root, args.run_id, handoff_id).exists():
        print(f"office-system: handoff already exists: {handoff_id}", file=sys.stderr)
        return 2
    context, context_errors = validate_context_envelope(
        root,
        context,
        run_id=safe_component(args.run_id, "run id"),
        run=run,
        from_agent=from_agent,
        to_agent=to_agent,
        handoff_id=handoff_id,
    )
    if context_errors:
        print("office-system: invalid context envelope:", file=sys.stderr)
        for error in context_errors:
            print(f"  - {error}", file=sys.stderr)
        return 2
    if set(context.get("acceptance_criteria", [])) != set(acceptance):
        print("office-system: handoff acceptance criteria must match context.acceptance_criteria", file=sys.stderr)
        return 2
    stage = normalize_loop_stage(args.stage or str(run.get("current_stage", "context")))
    if context.get("current_stage") != stage:
        print("office-system: handoff stage must match context.current_stage", file=sys.stderr)
        return 2
    envelope = {
        "version": "2.0.0",
        "kind": "digital-office-typed-handoff",
        "handoff_id": handoff_id,
        "run_id": safe_component(args.run_id, "run id"),
        "from_agent": from_agent,
        "to_agent": to_agent,
        "context_id": context["context_id"],
        "task_id": context["task_id"],
        "context_version": context["context_version"],
        "stage": stage,
        "reason": safe_claim(args.reason, "handoff reason"),
        "input_schema": input_schema,
        "input_schema_hash": canonical_hash(input_schema),
        "context": context,
        "context_hash": canonical_hash(context),
        "artifacts": [safe_claim(item, "handoff artifact") for item in (args.artifact or [])],
        "acceptance_criteria": acceptance,
        "status": "pending_acceptance",
        "acknowledgments": [],
        "created_by": safe_claim(args.created_by, "created by", required=False),
        "created_at": now_iso(),
    }
    envelope["handoff_contract_hash"] = canonical_hash({key: value for key, value in envelope.items() if key != "handoff_contract_hash"})
    write_json(handoff_path(root, args.run_id, handoff_id), envelope)
    run.setdefault("handoffs", []).append({"handoff_id": handoff_id, "from_agent": from_agent, "to_agent": to_agent, "path": str(handoff_path(root, args.run_id, handoff_id).relative_to(root)), "contract_hash": envelope["handoff_contract_hash"], "created_at": envelope["created_at"]})
    run["updated_at"] = now_iso()
    run.setdefault("events", []).append({"time": run["updated_at"], "event": "handoff_created", "handoff_id": handoff_id, "from_agent": from_agent, "to_agent": to_agent})
    write_json(run_record_path(root, args.run_id), run)
    append_run_ledger_event(
        root,
        args.run_id,
        "handoff_created",
        stage=envelope["stage"],
        action="handoff.create",
        agent_id=from_agent,
        actor_id=args.created_by or "",
        actor_role=args.role or "",
        input_payload={"to_agent": to_agent, "reason": envelope["reason"], "input_schema": input_schema, "context": context},
        output_payload={"handoff_id": handoff_id, "contract_hash": envelope["handoff_contract_hash"], "acceptance_criteria": acceptance},
        artifact_refs=envelope["artifacts"],
        handoff_id=handoff_id,
    )
    print(json.dumps(envelope, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def handoff_ack(args: argparse.Namespace) -> int:
    root = system_root()
    run = load_run_record(root, args.run_id)
    handoff_id = safe_component(args.handoff_id, "handoff id")
    path = handoff_path(root, args.run_id, handoff_id)
    if not path.exists():
        print(f"office-system: handoff not found: {handoff_id}", file=sys.stderr)
        return 2
    record = read_json(path)
    recipient = registered_agent(root, args.received_by)
    if record.get("to_agent") != recipient:
        print("office-system: only the registered target Agent may acknowledge this handoff", file=sys.stderr)
        return 2
    current_status = str(record.get("status", ""))
    if current_status not in {"pending_acceptance", "needs_context"}:
        print(f"office-system: handoff cannot be acknowledged from status {current_status}", file=sys.stderr)
        return 2
    expected_hash = safe_claim(args.expected_context_hash, "expected context hash")
    actual_hash = str(record.get("context_hash", ""))
    if not hmac.compare_digest(expected_hash, actual_hash):
        print("office-system: context hash mismatch; refusing acknowledgment", file=sys.stderr)
        return 2
    missing_fields = [safe_claim(item, "missing context field") for item in (args.missing_field or [])]
    message = safe_claim(args.message, "acknowledgment message", required=False)
    if args.decision == "request_context" and (not missing_fields or not message):
        print("office-system: request_context requires --missing-field and --message", file=sys.stderr)
        return 2
    if args.decision == "reject" and not message:
        print("office-system: reject requires --message", file=sys.stderr)
        return 2
    status_by_decision = {"accept": "accepted", "request_context": "needs_context", "reject": "rejected"}
    acknowledged_at = now_iso()
    acknowledgment = {
        "decision": args.decision,
        "received_by": recipient,
        "verified_context_hash": expected_hash,
        "missing_fields": missing_fields,
        "message": message,
        "confirmed": bool(args.confirmed),
        "acknowledged_at": acknowledged_at,
    }
    record["status"] = status_by_decision[args.decision]
    record.setdefault("acknowledgments", []).append(acknowledgment)
    record["updated_at"] = acknowledged_at
    write_json(path, record)
    for summary in run.get("handoffs", []):
        if summary.get("handoff_id") == handoff_id:
            summary["status"] = record["status"]
            summary["acknowledged_at"] = acknowledged_at
            summary["received_by"] = recipient
    run["updated_at"] = acknowledged_at
    run.setdefault("events", []).append({
        "time": acknowledged_at,
        "event": "handoff_acknowledged",
        "handoff_id": handoff_id,
        "decision": args.decision,
        "received_by": recipient,
    })
    write_json(run_record_path(root, args.run_id), run)
    append_run_ledger_event(
        root,
        args.run_id,
        "handoff_acknowledged",
        stage=str(record.get("stage", run.get("current_stage", "context"))),
        action="handoff.ack",
        agent_id=recipient,
        actor_id=args.created_by or "",
        actor_role=args.role or "",
        input_payload={"decision": args.decision, "expected_context_hash": expected_hash, "missing_fields": missing_fields},
        output_payload={"handoff_id": handoff_id, "status": record["status"], "message": message},
        artifact_refs=record.get("artifacts", []),
        handoff_id=handoff_id,
    )
    print(json.dumps(record, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def handoff_list(args: argparse.Namespace) -> int:
    root = system_root()
    load_run_record(root, args.run_id)
    records = read_records(handoff_dir(root, args.run_id), "digital-office-typed-handoff")
    if args.agent:
        agent = registered_agent(root, args.agent)
        records = [item for item in records if item.get("from_agent") == agent or item.get("to_agent") == agent]
    print(json.dumps({"run_id": args.run_id, "handoffs": records[: args.limit]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def coordination_policy(root: Path) -> dict[str, Any]:
    path = root / "coordination.policy.json"
    if path.exists():
        return read_json(path)
    return {
        "version": "0.0.0",
        "kind": "digital-office-coordination-policy",
        "max_parallel_agents": 3,
        "high_risk_modes": ["human_gated", "secretary_centralized"],
    }


def infer_task_risk(task: str, supplied: str = "") -> str:
    if supplied:
        return supplied
    text = task.lower()
    critical_terms = ["delete", "overwrite", "deploy", "publish", "payment", "export customer data", "删除", "覆盖", "部署", "发布", "付款", "导出客户数据"]
    high_terms = ["legal", "medical", "tax", "investment", "contract", "法律", "医疗", "税务", "投资", "合同"]
    if any(term in text for term in critical_terms):
        return "critical"
    if any(term in text for term in high_terms):
        return "high"
    return "medium" if len(task) > 240 else "low"


def build_coordination_plan(
    root: Path,
    *,
    task: str,
    candidate_agents: list[str],
    complexity: str = "",
    risk: str = "",
    parallelizable: bool = False,
    sequential: bool = False,
    run_id: str = "",
) -> dict[str, Any]:
    policy = coordination_policy(root)
    agents = [registered_agent(root, item) for item in candidate_agents if item]
    inferred_risk = infer_task_risk(task, risk)
    text = task.lower()
    if not complexity:
        complexity = "high" if len(task) > 400 else "medium" if len(task) > 160 else "low"
    if any(term in text for term in ["parallel", "independent", "in parallel", "并行", "分别", "多个方向"]):
        parallelizable = True
    if not parallelizable and any(term in text for term in ["then", "after", "approval", "handoff", "依次", "之后", "审批", "交接"]):
        sequential = True
    if inferred_risk in {"high", "critical"}:
        mode = "human_gated"
        stop_required = True
        reason = "High-risk or regulated work requires secretary mediation and human judgment before execution."
    elif parallelizable and len(agents) > 1 and not sequential:
        mode = "parallel_expert_dag"
        stop_required = False
        reason = "Task has independent workstreams and multiple candidate specialists."
    elif sequential or len(agents) > 1:
        mode = "sequential_specialist_chain"
        stop_required = False
        reason = "Task contains ordered dependencies or specialist handoffs."
    elif len(agents) == 1:
        mode = "single_agent"
        stop_required = False
        reason = "One specialist is sufficient and risk is not high."
    else:
        mode = "secretary_centralized"
        stop_required = False
        reason = "No concrete specialist set was supplied; secretary should coordinate intake and routing."
    max_parallel = int(policy.get("max_parallel_agents", 3))
    return {
        "version": "1.0.0",
        "kind": "digital-office-coordination-plan",
        "task_hash": canonical_hash(task),
        "run_id": safe_component(run_id, "run id") if run_id else "",
        "mode": mode,
        "risk": inferred_risk,
        "complexity": complexity,
        "candidate_agents": agents,
        "max_parallel_agents": max_parallel if mode == "parallel_expert_dag" else 1,
        "stop_required_before_execution": stop_required,
        "required_gates": ["human_judgment_closed", "typed_handoff_contracts", "checkpoint_before_resume"] if stop_required else ["typed_handoff_contracts", "checkpoint_before_resume"],
        "reason": reason,
        "drift_controls": [
            "run_ledger_hash_chain",
            "checkpoint_resume_cursor",
            "typed_handoff_contract_hash",
            "deterministic_eval_before_completion",
        ],
        "created_at": now_iso(),
    }


def coordination_plan(args: argparse.Namespace) -> int:
    root = system_root()
    task = args.task if args.task is not None else sys.stdin.read()
    agents = args.agent or []
    plan = build_coordination_plan(
        root,
        task=task.strip(),
        candidate_agents=agents,
        complexity=args.complexity or "",
        risk=args.risk or "",
        parallelizable=args.parallelizable,
        sequential=args.sequential,
        run_id=args.run_id or "",
    )
    if args.run_id:
        load_run_record(root, args.run_id)
        append_run_ledger_event(
            root,
            args.run_id,
            "coordination_plan_created",
            action="coordination.plan",
            input_payload={"task": task.strip(), "candidate_agents": agents, "risk": args.risk or "", "complexity": args.complexity or ""},
            output_payload=plan,
        )
    print(json.dumps(plan, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def eval_suite_path(root: Path, suite: str) -> Path:
    candidate = Path(suite).expanduser()
    if candidate.exists():
        return candidate
    if not candidate.suffix:
        candidate = candidate.with_suffix(".json")
    path = root / "evals" / candidate.name
    if path.exists():
        return path
    print(f"office-system: eval suite not found: {suite}", file=sys.stderr)
    raise SystemExit(2)


def grade_eval_case(root: Path, case: dict[str, Any]) -> dict[str, Any]:
    case_type = str(case.get("type", ""))
    case_id = str(case.get("id", "unnamed"))
    failures: list[str] = []
    observed: dict[str, Any] = {}
    if case_type == "judgment":
        observed = evaluate_judgment(
            root,
            task=str(case.get("task", "")),
            stage=normalize_loop_stage(str(case.get("stage", "context"))),
            action=str(case.get("action", "")),
            route=case.get("route") if isinstance(case.get("route"), dict) else {},
            signal=case.get("signal") if isinstance(case.get("signal"), dict) else {},
        )
        if "expect_decision" in case and observed.get("decision") != case["expect_decision"]:
            failures.append(f"decision expected {case['expect_decision']!r}, got {observed.get('decision')!r}")
        for action in case.get("expect_blocked_actions", []) or []:
            if action not in observed.get("blocked_actions", []):
                failures.append(f"blocked action missing: {action}")
    elif case_type == "coordination":
        observed = build_coordination_plan(
            root,
            task=str(case.get("task", "")),
            candidate_agents=[str(item) for item in case.get("candidate_agents", [])],
            complexity=str(case.get("complexity", "")),
            risk=str(case.get("risk", "")),
            parallelizable=bool(case.get("parallelizable", False)),
            sequential=bool(case.get("sequential", False)),
        )
        if "expect_mode" in case and observed.get("mode") != case["expect_mode"]:
            failures.append(f"mode expected {case['expect_mode']!r}, got {observed.get('mode')!r}")
        if "expect_stop_required" in case and observed.get("stop_required_before_execution") is not bool(case["expect_stop_required"]):
            failures.append("stop_required_before_execution mismatch")
    elif case_type == "rule_scope":
        observed = infer_rule_scope(root, title=str(case.get("title", "")), body=str(case.get("body", "")), project=str(case.get("project_id", "")), agent=str(case.get("agent_id", "")))
        if "expect_scope" in case and observed.get("scope") != case["expect_scope"]:
            failures.append(f"scope expected {case['expect_scope']!r}, got {observed.get('scope')!r}")
    else:
        failures.append(f"unknown eval case type: {case_type}")
    return {
        "case_id": case_id,
        "type": case_type,
        "language": case.get("language", ""),
        "status": "pass" if not failures else "fail",
        "failures": failures,
        "observed": observed,
    }


def eval_run(args: argparse.Namespace) -> int:
    root = system_root()
    path = eval_suite_path(root, args.suite)
    suite = read_json(path)
    cases = suite.get("cases", [])
    if not isinstance(cases, list):
        print("office-system: eval suite cases must be a list", file=sys.stderr)
        return 2
    if not cases:
        print("office-system: eval suite must include at least one case", file=sys.stderr)
        return 2
    if any(not isinstance(case, dict) for case in cases):
        print("office-system: every eval suite case must be an object", file=sys.stderr)
        return 2
    results = [grade_eval_case(root, case) for case in cases]
    passed = sum(1 for item in results if item["status"] == "pass")
    report = {
        "version": "1.0.0",
        "kind": "digital-office-eval-report",
        "suite_id": suite.get("suite_id", path.stem),
        "suite_path": str(path.relative_to(root) if path.is_relative_to(root) else path),
        "status": "success" if passed == len(results) else "error",
        "passed": passed,
        "total": len(results),
        "results": results,
        "generated_at": now_iso(),
    }
    if not args.no_write:
        report_path = root / "evals" / "reports" / f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{slugify(str(report['suite_id']))}.json"
        write_json(report_path, report)
        report["report_path"] = str(report_path.relative_to(root))
    print(json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True))
    return 0 if report["status"] == "success" else 1


def iteration_proposal_id(title: str) -> str:
    return f"{dt.datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}-{slugify(title)}"


def iteration_status_path(root: Path, proposal_id: str) -> Path:
    return root / "iterations" / "status" / f"{safe_component(proposal_id, 'proposal id')}.json"


def iteration_proposal_create(args: argparse.Namespace) -> int:
    root = system_root()
    project = safe_component(args.project, "project id") if args.project else ""
    agent = registered_agent(root, args.agent) if args.agent else ""
    if project:
        ensure_project(root, project)
    if args.run_id:
        run_path = loop_run_path(root, args.run_id)
        if not run_path.exists():
            print(f"office-system: loop run not found: {args.run_id}", file=sys.stderr)
            return 2
    body = args.body if args.body is not None else sys.stdin.read()
    proposal_id = iteration_proposal_id(args.title)
    report = root / "iterations" / "proposals" / f"{proposal_id}.md"
    status_file = iteration_status_path(root, proposal_id)
    proposal = {
        "version": "1.0.0",
        "kind": "digital-office-iteration-proposal",
        "proposal_id": proposal_id,
        "status": "pending_user_confirmation",
        "title": args.title,
        "target": args.target,
        "run_id": safe_component(args.run_id, "run id") if args.run_id else "",
        "project_id": project,
        "agent_id": agent,
        "summary": args.summary,
        "body": body.strip(),
        "expected_impact": args.expected_impact,
        "risk": args.risk,
        "rollback": args.rollback,
        "source_refs": [safe_claim(item, "source ref") for item in (args.source_ref or [])],
        "regression_checks": [safe_claim(item, "regression check") for item in (args.regression_check or [])],
        "requires_user_confirmation": True,
        "auto_apply_allowed": False,
        "created_at": now_iso(),
        "report": str(report.relative_to(root)),
        "allowed_decisions": ["confirm", "tune", "pause", "reject"],
    }
    lines = [
        f"# Iteration Proposal: {args.title}",
        "",
        "- State: pending_user_confirmation",
        f"- Proposal ID: {proposal_id}",
        f"- Target: {args.target}",
        f"- Run ID: {proposal['run_id']}",
        f"- Project: {project}",
        f"- Agent: {agent}",
        f"- Created at: {proposal['created_at']}",
        "",
        "## What Will Change",
        "",
        body.strip() or args.summary,
        "",
        "## Why This Is Suggested",
        "",
        args.summary,
        "",
        "## Expected Impact",
        "",
        args.expected_impact,
        "",
        "## Risk",
        "",
        args.risk,
        "",
        "## Rollback",
        "",
        args.rollback,
        "",
        "## Regression Checks",
        "",
    ]
    if args.regression_check:
        lines.extend(f"- {item}" for item in args.regression_check)
    else:
        lines.append("- No regression check provided yet. This proposal cannot be called production-ready until one is added.")
    lines.extend(
        [
            "",
            "## User Decision Required",
            "",
            "This proposal is not applied automatically. The GUI must show the change, risk, rollback, and regression checks, then wait for one of: Confirm, Tune Through Conversation, Pause, Reject.",
        ]
    )
    write_text(report, "\n".join(lines) + "\n")
    write_json(status_file, proposal)
    append_log(root, {"event": "iteration_proposal_create", "proposal_id": proposal_id, "target": args.target, "run_id": proposal["run_id"]})
    print(json.dumps({"proposal_id": proposal_id, "status": proposal["status"], "report": proposal["report"], "status_file": str(status_file.relative_to(root))}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def iteration_proposal_decision(args: argparse.Namespace) -> int:
    root = system_root()
    proposal_id = safe_component(args.proposal_id, "proposal id")
    status_file = iteration_status_path(root, proposal_id)
    if not status_file.exists():
        print(f"office-system: iteration proposal not found: {proposal_id}", file=sys.stderr)
        return 2
    state = read_json(status_file)
    state["status"] = ITERATION_DECISION_TO_STATUS[args.decision]
    state["decision"] = args.decision
    state["decision_updated_at"] = now_iso()
    if args.message:
        state.setdefault("decision_messages", []).append({"time": now_iso(), "message": args.message})
    if args.decision == "confirm":
        state["next_action"] = f"iteration-proposal-apply --proposal-id {proposal_id} --confirmed"
    elif args.decision == "tune":
        state["next_action"] = "secretary continues the iteration conversation and regenerates or updates the proposal"
    elif args.decision == "pause":
        state["next_action"] = "do nothing until the user resumes"
    elif args.decision == "reject":
        state["next_action"] = "do not apply this proposal"
    write_json(status_file, state)
    append_log(root, {"event": "iteration_proposal_decision", "proposal_id": proposal_id, "decision": args.decision})
    print(json.dumps(state, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def iteration_proposal_apply(args: argparse.Namespace) -> int:
    root = system_root()
    proposal_id = safe_component(args.proposal_id, "proposal id")
    status_file = iteration_status_path(root, proposal_id)
    if not status_file.exists():
        print(f"office-system: iteration proposal not found: {proposal_id}", file=sys.stderr)
        return 2
    if not args.confirmed:
        print("office-system: iteration application requires --confirmed after the user reviews the proposal", file=sys.stderr)
        return 2
    state = read_json(status_file)
    if state.get("status") != "confirmed_for_application":
        print("office-system: iteration proposal must be confirmed before application", file=sys.stderr)
        return 2
    record = {
        "version": "1.0.0",
        "kind": "digital-office-applied-iteration-record",
        "proposal_id": proposal_id,
        "target": state.get("target"),
        "applied_at": now_iso(),
        "applied_by": safe_claim(args.applied_by, "applied by", required=False),
        "artifacts": [safe_claim(item, "artifact") for item in (args.artifact or [])],
        "regression_result": args.regression_result or "",
        "rollback": state.get("rollback", ""),
        "note": args.note or "",
    }
    state["status"] = "applied_verified" if args.regression_result else "applied_pending_regression"
    state["applied_at"] = record["applied_at"]
    state["applied_record"] = f"iterations/applied/{proposal_id}.json"
    write_json(root / "iterations" / "applied" / f"{proposal_id}.json", record)
    write_json(status_file, state)
    append_log(root, {"event": "iteration_proposal_apply", "proposal_id": proposal_id, "status": state["status"]})
    print(json.dumps({"proposal_id": proposal_id, "status": state["status"], "applied_record": state["applied_record"]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def health_checks(root: Path) -> dict[str, Any]:
    return {
        "root": str(root),
        "agents_registry": (root / "agents.registry.json").exists(),
        "digital_employees_registry": (root / "digital-employees.registry.json").exists(),
        "workflow_packs_registry": (root / "workflow-packs.registry.json").exists(),
        "context_envelope_schema": (root / "context-envelope.schema.json").exists(),
        "context_handoff_policy": (root / "context-handoff.policy.json").exists(),
        "skill_installations_registry": (root / "skill-installations.registry.json").exists(),
        "local_skill_sources": (root.parent / "skills" / "_imported" / "claude-for-legal-ZH").exists(),
        "knowledge_registry": (root / "knowledge.registry.json").exists(),
        "identity_access_registry": (root / "identity.access.registry.json").exists(),
        "industry_solutions_registry": (root / "industry-solutions.registry.json").exists(),
        "external_knowledge_sources_registry": (root / "external-knowledge-sources.registry.json").exists(),
        "ai_native_loop_manifest": (root / "ai-native-loop.manifest.json").exists(),
        "onboarding_presets": (root / "onboarding.presets.json").exists(),
        "settings_dir": (root / "settings").exists(),
        "user_preferences_configured": onboarding_preferences_path(root).exists(),
        "rules_registry": (root / "rules" / "rules.registry.json").exists(),
        "judgment_policy": (root / "judgment.policy.json").exists(),
        "coordination_policy": (root / "coordination.policy.json").exists(),
        "multimodal_pipeline": (root / "multimodal.pipeline.json").exists(),
        "rag_pipeline": (root / "rag.pipeline.json").exists(),
        "workflow_runs_dir": (root / "runs").exists(),
        "evals_dir": (root / "evals").exists(),
        "eval_reports_dir": (root / "evals" / "reports").exists(),
        "task_inbox_dir": (root / "tasks").exists(),
        "approval_center_dir": (root / "approvals").exists(),
        "judgment_cases_dir": (root / "judgments").exists(),
        "rule_proposals_dir": (root / "rule-proposals").exists(),
        "notifications_dir": (root / "notifications").exists(),
        "audit_logs_dir": (root / "logs").exists(),
        "knowledge_spaces_dir": (root / "knowledge" / "spaces").exists(),
        "web_app_dir": (root / "web" / "app").exists(),
        "web_index": (root / "web" / "app" / "index.html").exists(),
        "pwa_manifest": (root / "web" / "app" / "manifest.webmanifest").exists(),
        "service_worker": (root / "web" / "app" / "service-worker.js").exists(),
        "tesseract": bool(shutil.which("tesseract")),
        "pdftotext": bool(shutil.which("pdftotext")),
        "rapidocr_onnxruntime": bool(importlib.util.find_spec("rapidocr_onnxruntime")),
        "pypdf": bool(importlib.util.find_spec("pypdf")),
        "python_docx": bool(importlib.util.find_spec("docx")),
        "sentence_transformers": bool(importlib.util.find_spec("sentence_transformers")),
        "python_docx_builtin": True,
    }


def required_health_keys() -> tuple[str, ...]:
    return (
        "agents_registry",
        "digital_employees_registry",
        "workflow_packs_registry",
        "context_envelope_schema",
        "context_handoff_policy",
        "skill_installations_registry",
        "local_skill_sources",
        "ai_native_loop_manifest",
        "onboarding_presets",
        "settings_dir",
        "workflow_runs_dir",
        "task_inbox_dir",
        "approval_center_dir",
        "judgment_policy",
        "coordination_policy",
        "evals_dir",
        "eval_reports_dir",
        "notifications_dir",
        "audit_logs_dir",
        "knowledge_spaces_dir",
        "web_app_dir",
        "web_index",
        "pwa_manifest",
        "service_worker",
    )


def optional_health_keys() -> tuple[str, ...]:
    return (
        "knowledge_registry",
        "identity_access_registry",
        "industry_solutions_registry",
        "external_knowledge_sources_registry",
        "rules_registry",
        "multimodal_pipeline",
    )




def detailed_health(root: Path) -> dict[str, Any]:
    """Return detailed component-level health status for the /api/health endpoint."""
    checks = health_checks(root)
    components: dict[str, Any] = {}

    # knowledge_base
    kb_reg = root / 'knowledge.registry.json'
    if kb_reg.exists():
        try:
            kb_data = json.loads(kb_reg.read_text(encoding='utf-8'))
            spaces_count = len(kb_data.get('knowledge_bases', {}))
            components['knowledge_base'] = {'status': 'ok', 'spaces_count': spaces_count}
        except (json.JSONDecodeError, OSError):
            components['knowledge_base'] = {'status': 'degraded', 'spaces_count': 0}
    else:
        components['knowledge_base'] = {'status': 'not_found', 'spaces_count': 0}

    # rag_index
    rag_pipe = root / 'rag.pipeline.json'
    if rag_pipe.exists():
        try:
            rag_data = json.loads(rag_pipe.read_text(encoding='utf-8'))
            stores = rag_data.get('stores', {})
            default_store = rag_data.get('default_store', '')
            if default_store and default_store in stores:
                store_cfg = stores[default_store]
                index_path = root / store_cfg.get('index_dir', 'data/rag-index')
                if index_path.exists() and any(index_path.iterdir()):
                    components['rag_index'] = {'status': 'ok'}
                else:
                    components['rag_index'] = {'status': 'not_indexed', 'optional': True}
            else:
                components['rag_index'] = {'status': 'not_indexed', 'optional': True}
        except (json.JSONDecodeError, OSError):
            components['rag_index'] = {'status': 'degraded'}
    else:
        components['rag_index'] = {'status': 'not_found', 'optional': True}

    # agents
    agents_reg = root / 'agents.registry.json'
    if agents_reg.exists():
        try:
            agents_data = effective_agents_registry(root)
            registered_count = len(agents_data.get('agents', {}))
            components['agents'] = {'status': 'ok', 'registered_count': registered_count}
        except (json.JSONDecodeError, OSError):
            components['agents'] = {'status': 'degraded', 'registered_count': 0}
    else:
        components['agents'] = {'status': 'not_found', 'registered_count': 0}

    # workflows
    runs_dir = root / 'runs'
    if runs_dir.exists():
        active_count = sum(1 for p in runs_dir.iterdir() if p.is_dir())
        components['workflows'] = {'status': 'ok', 'active_count': active_count}
    else:
        components['workflows'] = {'status': 'not_found', 'active_count': 0}

    # disk
    try:
        disk_usage = shutil.disk_usage(str(root))
        available_mb = disk_usage.free // (1024 * 1024)
        disk_status = 'ok' if available_mb > 500 else 'low'
        components['disk'] = {'status': disk_status, 'available_mb': available_mb}
    except OSError:
        components['disk'] = {'status': 'unknown', 'available_mb': 0}

    # Determine overall status
    any_degraded = any(
        c.get('status') not in ('ok', 'low')
        for c in components.values()
        if not c.get('optional')
    )
    any_down = any(
        c.get('status') in ('not_found',)
        for c in components.values()
        if not c.get('optional')
    )
    required = required_health_keys()
    all_required_ok = all(checks.get(k) for k in required)

    if not all_required_ok or any_down:
        overall = 'down'
    elif any_degraded:
        overall = 'degraded'
    else:
        overall = 'ok'

    return {
        'status': overall,
        'timestamp': now_iso(),
        'checks': checks,
        'components': components,
    }
def status_counts(records: list[dict[str, Any]], field: str = "status") -> dict[str, int]:
    counts: dict[str, int] = {}
    for record in records:
        key = str(record.get(field) or "unknown")
        counts[key] = counts.get(key, 0) + 1
    return dict(sorted(counts.items()))


def compact_records(records: list[dict[str, Any]], fields: list[str], limit: int) -> list[dict[str, Any]]:
    compacted = []
    for record in records[:limit]:
        compacted.append({field: record.get(field, "") for field in fields})
    return compacted


def project_summaries(root: Path, limit: int) -> list[dict[str, Any]]:
    projects: list[dict[str, Any]] = []
    for path in sorted((root / "projects").glob("*/project.json")):
        if path.parent.name == "_template":
            continue
        try:
            data = read_json(path)
        except Exception:
            continue
        assessment = assess_project_context(root, path.parent.name, data)
        projects.append(
            {
                "project_id": data.get("project_id", path.parent.name),
                "name": data.get("name", path.parent.name),
                "status": data.get("status", ""),
                "agent_roster": data.get("agent_roster", []),
                "updated_at": data.get("updated_at", data.get("created_at", "")),
                "context_readiness": {key: assessment[key] for key in ("required", "readiness_score", "readiness_threshold", "ready", "confirmed", "context_version", "blockers", "suggestions", "question_policy", "intent", "context")},
            }
        )
    projects.sort(key=lambda item: str(item.get("updated_at") or ""), reverse=True)
    return projects[:limit]


def usage_total(value: Any, needle: str) -> int:
    total = 0
    if isinstance(value, dict):
        for key, item in value.items():
            if needle in str(key).lower() and isinstance(item, (int, float)):
                total += int(item)
            total += usage_total(item, needle)
    elif isinstance(value, list):
        for item in value:
            total += usage_total(item, needle)
    return total


def employee_performance_summaries(agents: list[dict[str, Any]], runs: list[dict[str, Any]], tasks: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    success_statuses = {"completed", "delivered", "approved"}
    issue_statuses = {"failed", "cancelled", "stopped", "blocked", "rejected"}
    metrics: dict[str, dict[str, Any]] = {}

    def ensure(agent_id: str) -> dict[str, Any]:
        agent_id = str(agent_id or "unassigned")
        if agent_id not in metrics:
            metrics[agent_id] = {
                "agent_id": agent_id,
                "run_count": 0,
                "task_count": 0,
                "success_count": 0,
                "issue_count": 0,
                "active_count": 0,
                "token_estimate": 0,
                "model_calls": 0,
                "tool_calls": 0,
                "last_active_at": "",
                "success_rate": 0,
            }
        return metrics[agent_id]

    for agent in agents:
        ensure(str(agent.get("agent_id", "")))

    for run in runs:
        agent_id = str(run.get("agent_id") or "unassigned")
        item = ensure(agent_id)
        status = str(run.get("status", ""))
        item["run_count"] += 1
        if status in success_statuses:
            item["success_count"] += 1
        elif status in issue_statuses:
            item["issue_count"] += 1
        else:
            item["active_count"] += 1
        usage = run.get("control", {}).get("usage", {})
        item["token_estimate"] += usage_total(usage, "token")
        item["model_calls"] += usage_total(usage, "model_call")
        item["tool_calls"] += usage_total(usage, "tool_call")
        updated = str(run.get("updated_at") or run.get("created_at") or "")
        if updated > str(item.get("last_active_at", "")):
            item["last_active_at"] = updated

    for task in tasks:
        agent_id = str(task.get("assigned_agent") or "unassigned")
        item = ensure(agent_id)
        item["task_count"] += 1
        status = str(task.get("status", ""))
        if status in success_statuses:
            item["success_count"] += 1
        elif status in issue_statuses:
            item["issue_count"] += 1
        updated = str(task.get("updated_at") or task.get("created_at") or "")
        if updated > str(item.get("last_active_at", "")):
            item["last_active_at"] = updated

    for item in metrics.values():
        total_closed = int(item["success_count"]) + int(item["issue_count"])
        item["success_rate"] = round((int(item["success_count"]) / total_closed) * 100) if total_closed else 0
    return metrics


def employee_gap_suggestions(agents: list[dict[str, Any]], runs: list[dict[str, Any]], tasks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    present = {str(agent.get("agent_id", "")) for agent in agents}
    text = " ".join(str(item.get("title", "")) for item in [*runs, *tasks]).lower()
    suggestions = [
        {
            "suggested_agent_id": "finance",
            "display_name": "财务助理",
            "template_agent_id": "planner",
            "reason": "适合处理预算、报价、发票、付款节点和成本核对。",
            "skills": ["finance-billing-ops", "cost-tracking"],
            "keywords": ["预算", "报价", "发票", "付款", "成本"],
        },
        {
            "suggested_agent_id": "sales-ops",
            "display_name": "销售运营助理",
            "template_agent_id": "writer",
            "reason": "适合沉淀客户跟进、商机材料、报价说明和销售话术。",
            "skills": ["content-strategy", "customer-billing-ops"],
            "keywords": ["客户", "销售", "商机", "报价", "跟进"],
        },
        {
            "suggested_agent_id": "customer-success",
            "display_name": "客户成功助理",
            "template_agent_id": "researcher",
            "reason": "适合汇总客户问题、服务记录、续约风险和满意度反馈。",
            "skills": ["carrier-relationship-management", "customer-billing-ops"],
            "keywords": ["客户问题", "续约", "满意度", "服务"],
        },
        {
            "suggested_agent_id": "operations",
            "display_name": "运营助理",
            "template_agent_id": "planner",
            "reason": "适合把重复流程、排期、物料和跨部门事项标准化。",
            "skills": ["automation-workflows", "inventory-demand-planning"],
            "keywords": ["排期", "流程", "运营", "物料"],
        },
    ]
    result = []
    for suggestion in suggestions:
        if suggestion["suggested_agent_id"] in present:
            continue
        priority = "medium"
        if any(keyword.lower() in text for keyword in suggestion["keywords"]):
            priority = "high"
        result.append({**suggestion, "priority": priority})
    result.sort(key=lambda item: 0 if item["priority"] == "high" else 1)
    return result[:4]


def custom_agent_dependencies(root: Path, agent_id: str) -> dict[str, list[str]]:
    active_runs = [
        str(run.get("run_id", ""))
        for run in read_run_records(root)
        if run.get("agent_id") == agent_id and run.get("status") not in {"completed", "cancelled", "stopped", "failed"}
    ]
    projects = []
    for path in sorted((root / "projects").glob("*/project.json")):
        if path.parent.name == "_template":
            continue
        try:
            project = read_json(path)
        except (OSError, json.JSONDecodeError):
            continue
        if agent_id in project.get("agent_roster", []):
            projects.append(str(project.get("project_id", path.parent.name)))
    return {"active_runs": active_runs, "projects": projects}


def custom_agent_view(agent_id: str, record: dict[str, Any], root: Path | None = None) -> dict[str, Any]:
    config = record.get("config", {})
    lifecycle = record.get("lifecycle", {})
    view = {
        "agent_id": agent_id,
        "display_name": config.get("display_name", agent_id),
        "user_visible_role": record.get("user_visible_role", ""),
        "status": record.get("status", "inactive"),
        "origin": "custom",
        "editable": True,
        "template_agent_id": record.get("template_agent_id", ""),
        "profile": config.get("profile", ""),
        "provider": config.get("provider", ""),
        "model": config.get("model", ""),
        "skills": config.get("skills", []),
        "workflow_packs": record.get("workflow_packs", []),
        "created_at": lifecycle.get("created_at", ""),
        "updated_at": lifecycle.get("updated_at", ""),
        "created_by": lifecycle.get("created_by", ""),
    }
    if root is not None:
        view["dependencies"] = custom_agent_dependencies(root, agent_id)
    return view


def agent_lifecycle_create(args: argparse.Namespace) -> int:
    root = system_root()
    actor_role = require_agent_admin(args.role)
    if not args.confirmed:
        print("office-system: agent-create requires --confirmed", file=sys.stderr)
        return 2
    agent_id = safe_component(args.agent, "agent id")
    display_name = safe_claim(args.display_name, "Agent display name")
    visible_role = safe_claim(args.role_description, "Agent role description")
    template_id = safe_component(args.template, "template Agent id")
    base_registry = read_json(root / "agents.registry.json")
    template = base_registry.get("agents", {}).get(template_id)
    if not isinstance(template, dict):
        print(f"office-system: unknown built-in Agent template: {template_id}", file=sys.stderr)
        return 2
    if agent_id in base_registry.get("agents", {}):
        print("office-system: built-in Agent ids cannot be replaced", file=sys.stderr)
        return 2
    requested_skills = [safe_component(item, "skill id") for item in (args.skill or [])]
    requested_keywords = [safe_claim(item, "routing keyword") for item in (args.keyword or [])]
    path = custom_agents_path(root)
    with JsonFileLock(path.with_name(f".{path.name}.lock")):
        custom = load_custom_agents_registry(root)
        if agent_id in custom.get("agents", {}):
            print(f"office-system: custom Agent already exists: {agent_id}", file=sys.stderr)
            return 2
        config = copy.deepcopy(template)
        config["display_name"] = display_name
        config["source_template"] = template_id
        config["custom_agent"] = True
        config["skills"] = list(dict.fromkeys([*config.get("skills", []), *requested_skills]))
        if requested_keywords:
            config.setdefault("routing", {})["keywords"] = [
                {"term": keyword, "weight": 6} for keyword in requested_keywords
            ]
        record = {
            "version": "1.0.0",
            "status": "active",
            "template_agent_id": template_id,
            "user_visible_role": visible_role,
            "workflow_packs": [safe_component(item, "workflow pack") for item in (args.workflow_pack or [])],
            "config": config,
            "lifecycle": {
                "created_at": now_iso(),
                "updated_at": now_iso(),
                "created_by": safe_claim(args.requested_by, "requesting user"),
                "updated_by": safe_claim(args.requested_by, "requesting user"),
            },
        }
        custom.setdefault("agents", {})[agent_id] = record
        custom["updated_at"] = now_iso()
        write_json_unlocked(path, custom)
    audit = append_audit_event(
        root,
        "agent_created",
        actor_id=args.requested_by,
        actor_role=actor_role,
        agent_id=agent_id,
        resource_type="agent",
        resource_id=agent_id,
        outcome="created",
        extra={"template_agent_id": template_id, "skills": requested_skills},
    )
    payload = custom_agent_view(agent_id, record, root)
    payload["audit_event_id"] = audit["event_id"]
    print(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def agent_lifecycle_status(args: argparse.Namespace) -> int:
    root = system_root()
    actor_role = require_agent_admin(args.role)
    if not args.confirmed:
        print("office-system: agent-status requires --confirmed", file=sys.stderr)
        return 2
    agent_id = safe_component(args.agent, "agent id")
    status = safe_component(args.status, "Agent status")
    if status not in CUSTOM_AGENT_STATUSES:
        print(f"office-system: unsupported Agent status: {status}", file=sys.stderr)
        return 2
    path = custom_agents_path(root)
    with JsonFileLock(path.with_name(f".{path.name}.lock")):
        custom = load_custom_agents_registry(root)
        record = custom.get("agents", {}).get(agent_id)
        if not isinstance(record, dict):
            print("office-system: only custom Agents can change lifecycle status", file=sys.stderr)
            return 2
        previous = str(record.get("status", "inactive"))
        if previous == status:
            print(json.dumps(custom_agent_view(agent_id, record, root), ensure_ascii=False, indent=2, sort_keys=True))
            return 0
        if status == "archived" and custom_agent_dependencies(root, agent_id)["active_runs"]:
            print("office-system: stop or complete active Agent runs before archiving", file=sys.stderr)
            return 4
        record["status"] = status
        record.setdefault("lifecycle", {})["updated_at"] = now_iso()
        record["lifecycle"]["updated_by"] = safe_claim(args.requested_by, "requesting user")
        record["lifecycle"]["status_reason"] = safe_claim(args.reason, "status reason", required=False)
        custom["updated_at"] = now_iso()
        write_json_unlocked(path, custom)
    audit = append_audit_event(
        root,
        "agent_status_changed",
        actor_id=args.requested_by,
        actor_role=actor_role,
        agent_id=agent_id,
        resource_type="agent",
        resource_id=agent_id,
        outcome=status,
        reason=args.reason or "",
        extra={"previous_status": previous, "status": status},
    )
    payload = custom_agent_view(agent_id, record, root)
    payload["audit_event_id"] = audit["event_id"]
    print(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def agent_lifecycle_delete(args: argparse.Namespace) -> int:
    root = system_root()
    actor_role = require_agent_admin(args.role)
    if not args.confirmed:
        print("office-system: agent-delete requires --confirmed", file=sys.stderr)
        return 2
    agent_id = safe_component(args.agent, "agent id")
    path = custom_agents_path(root)
    with JsonFileLock(path.with_name(f".{path.name}.lock")):
        custom = load_custom_agents_registry(root)
        record = custom.get("agents", {}).get(agent_id)
        if not isinstance(record, dict):
            print("office-system: only custom Agents can be permanently deleted", file=sys.stderr)
            return 2
        if record.get("status") != "archived":
            print("office-system: archive the Agent before permanent deletion", file=sys.stderr)
            return 4
        dependencies = custom_agent_dependencies(root, agent_id)
        if dependencies["active_runs"]:
            print("office-system: Agent still has active workflow runs", file=sys.stderr)
            return 4
        deleted_at = now_iso()
        custom["agents"].pop(agent_id)
        custom.setdefault("tombstones", []).append(
            {
                "agent_id": agent_id,
                "display_name": record.get("config", {}).get("display_name", agent_id),
                "deleted_at": deleted_at,
                "deleted_by": safe_claim(args.requested_by, "requesting user"),
                "history_preserved": True,
            }
        )
        custom["updated_at"] = deleted_at
        write_json_unlocked(path, custom)
    audit = append_audit_event(
        root,
        "agent_deleted",
        actor_id=args.requested_by,
        actor_role=actor_role,
        agent_id=agent_id,
        resource_type="agent",
        resource_id=agent_id,
        outcome="deleted",
        reason=args.reason or "",
        extra={"history_preserved": True, "project_references": dependencies["projects"]},
    )
    print(json.dumps({"status": "deleted", "agent_id": agent_id, "history_preserved": True, "audit_event_id": audit["event_id"]}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def agent_lifecycle_list(args: argparse.Namespace) -> int:
    root = system_root()
    items = agent_summaries(root)
    if not args.include_archived:
        items = [item for item in items if item.get("status") != "archived"]
    print(json.dumps({"kind": "digital-office-agent-lifecycle-list", "count": len(items), "items": items}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def agent_summaries(root: Path) -> list[dict[str, Any]]:
    registry = read_json(root / "agents.registry.json")
    agents = []
    for agent_id, agent in sorted(registry.get("agents", {}).items()):
        agents.append(
            {
                "agent_id": agent_id,
                "display_name": agent.get("display_name", agent_id),
                "profile": agent.get("profile", ""),
                "portable_role": agent.get("portable_role", ""),
                "orchestration_roles": agent.get("orchestration_roles", []),
                "skills": agent.get("skills", []),
                "provider": agent.get("provider", ""),
                "model": agent.get("model", ""),
                "status": "active",
                "origin": "built_in",
                "editable": False,
            }
        )
    custom = load_custom_agents_registry(root)
    for agent_id, record in sorted(custom.get("agents", {}).items()):
        agents.append(custom_agent_view(agent_id, record, root))
    return agents


def digital_employee_summaries(root: Path) -> list[dict[str, Any]]:
    path = root / "digital-employees.registry.json"
    if not path.exists():
        return []
    registry = read_json(path)
    employees = []
    for employee_id, employee in sorted(registry.get("employees", {}).items()):
        employees.append(
            {
                "employee_id": employee_id,
                "agent_id": employee.get("agent_id", employee_id),
                "display_name": employee.get("display_name", employee_id),
                "display_name_zh": employee.get("display_name_zh", ""),
                "user_visible_role": employee.get("user_visible_role", ""),
                "represents_department_owner": bool(employee.get("represents_department_owner", False)),
                "workflow_packs": employee.get("workflow_packs", []),
                "skill_staff": employee.get("skill_staff", []),
                "status": "active",
                "origin": "built_in",
                "editable": False,
            }
        )
    custom = load_custom_agents_registry(root)
    for employee_id, record in sorted(custom.get("agents", {}).items()):
        config = record.get("config", {})
        employees.append(
            {
                "employee_id": employee_id,
                "agent_id": employee_id,
                "display_name": config.get("display_name", employee_id),
                "display_name_zh": config.get("display_name", employee_id),
                "user_visible_role": record.get("user_visible_role", ""),
                "represents_department_owner": True,
                "workflow_packs": record.get("workflow_packs", []),
                "skill_staff": config.get("skills", []),
                "status": record.get("status", "inactive"),
                "origin": "custom",
                "editable": True,
                "template_agent_id": record.get("template_agent_id", ""),
            }
        )
    return employees


def workflow_pack_summaries(root: Path) -> list[dict[str, Any]]:
    path = root / "workflow-packs.registry.json"
    if not path.exists():
        return []
    registry = read_json(path)
    packs = []
    for pack_id, pack in sorted(registry.get("packs", {}).items()):
        packs.append(
            {
                "pack_id": pack_id,
                "owner_agent": pack.get("owner_agent", ""),
                "workflows": pack.get("workflows", []),
                "skill_lanes": pack.get("skill_lanes", []),
                "context_envelope_required": bool(pack.get("context_envelope_required", False)),
            }
        )
    return packs


def context_contract_summary(root: Path) -> dict[str, Any]:
    path = root / "context-envelope.schema.json"
    if not path.exists():
        return {"configured": False, "required": []}
    schema = read_json(path)
    handoff_policy_path = root / "context-handoff.policy.json"
    handoff_policy = read_json(handoff_policy_path) if handoff_policy_path.exists() else {}
    return {
        "configured": True,
        "schema": "context-envelope.schema.json",
        "schema_version": schema.get("properties", {}).get("version", {}).get("const", ""),
        "required": schema.get("required", []),
        "actor_types": schema.get("$defs", {}).get("actor", {}).get("properties", {}).get("type", {}).get("enum", []),
        "handoff_policy": "context-handoff.policy.json" if handoff_policy else "",
        "default_transfer_mode": handoff_policy.get("default_mode", ""),
        "acknowledgment_decisions": handoff_policy.get("acknowledgment", {}).get("decisions", []),
    }


def loop_runtime_summary(root: Path) -> dict[str, Any]:
    path = root / "ai-native-loop.manifest.json"
    if not path.exists():
        return {"configured": False, "work_nodes": [], "controller_decisions": [], "default_budgets": {}}
    manifest = read_json(path)
    return {
        "configured": True,
        "version": manifest.get("version", ""),
        "work_nodes": list(manifest.get("stages", {}).keys()),
        "controller_decisions": list(manifest.get("controller", {}).get("decisions", {}).keys()),
        "default_budgets": manifest.get("controller", {}).get("default_budgets", {}),
        "task_profiles": manifest.get("architecture", {}).get("task_profiles", {}),
        "legacy_stage_aliases": manifest.get("architecture", {}).get("legacy_stage_aliases", {}),
    }


def skill_installation_summaries(root: Path) -> list[dict[str, Any]]:
    path = root / "skill-installations.registry.json"
    if not path.exists():
        return []
    registry = read_json(path)
    items = []
    for name, cfg in sorted(registry.get("installations", {}).items()):
        install_path = cfg.get("install_path", "")
        resolved = root / install_path if install_path else None
        if install_path and not Path(str(install_path)).is_absolute():
            resolved = (root / str(install_path)).resolve()
        skill_files = len(list(resolved.glob("**/SKILL.md"))) if resolved and resolved.exists() else 0
        items.append(
            {
                "name": name,
                "status": cfg.get("status", ""),
                "license": cfg.get("license", ""),
                "used_by": cfg.get("used_by", []),
                "install_path": install_path,
                "skill_files": skill_files,
                "activation": cfg.get("activation", ""),
            }
        )
    return items


def model_runtime_summary(root: Path) -> dict[str, Any]:
    gateway = root / "bin" / "model-gateway"
    if not gateway.exists():
        return {"status": "missing", "providers": [], "runtime": {"default_mode": "host", "agents": {}}}
    try:
        proc = subprocess.run([str(gateway), "status"], text=True, capture_output=True, timeout=10)
    except (OSError, subprocess.TimeoutExpired) as exc:
        return {"status": "error", "providers": [], "runtime": {}, "error": str(exc)}
    if proc.returncode != 0:
        return {"status": "error", "providers": [], "runtime": {}, "error": (proc.stderr or proc.stdout)[-1000:]}
    try:
        payload = json.loads(proc.stdout)
    except json.JSONDecodeError:
        return {"status": "error", "providers": [], "runtime": {}, "error": "model gateway returned invalid JSON"}
    return {
        "status": payload.get("status", "ready"),
        "providers": payload.get("providers", []),
        "runtime": payload.get("runtime", {}),
        "local_runtimes": payload.get("local_runtimes", []),
        "secret_storage": payload.get("secret_storage", "local_private_file_or_environment"),
        "configure_command": "model-gateway connection-set --provider PROVIDER --base-url URL --model MODEL --secret-stdin --confirmed",
    }


def gui_capabilities() -> list[dict[str, Any]]:
    return [
        {"id": "global_settings", "status": "ready", "commands": ["settings-options", "settings-status", "settings-update"]},
        {"id": "first_run_onboarding", "status": "ready", "commands": ["onboarding-options", "onboarding-apply"]},
        {"id": "web_ui_pwa", "status": "ready", "commands": ["digital-office-gui", "web-config", "web-serve"], "routes": ["/", "/manifest.webmanifest", "/service-worker.js", "/api/health", "/api/gui-state", "/api/web-app"]},
        {"id": "workflow_control_plane", "status": "ready", "commands": ["workflow-start", "workflow-status", "workflow-list", "workflow-cancel", "workflow-resume", "workflow-retry", "workflow-control"]},
        {"id": "direct_agent_invocation", "status": "ready", "commands": ["agent-invoke"]},
        {"id": "direct_model_api_gateway", "status": "ready", "commands": ["model-gateway status", "model-gateway connection-set", "model-gateway test", "model-gateway resolve", "model-gateway invoke"]},
        {"id": "workflow_canvas_revisions", "status": "ready", "commands": ["workflow-draft-create", "workflow-draft-patch", "workflow-draft-validate", "workflow-draft-activate", "workflow-node-context"]},
        {"id": "workflow_runtime_controls", "status": "ready", "commands": ["workflow-control", "workflow-node-context"]},
        {"id": "runtime_replay_and_checkpoints", "status": "ready", "commands": ["run-ledger-add", "run-ledger-list", "run-ledger-verify", "checkpoint-create", "checkpoint-list"]},
        {"id": "typed_agent_handoffs", "status": "ready", "commands": ["handoff-create", "handoff-ack", "handoff-list"]},
        {"id": "loop_runtime", "status": "ready", "commands": ["loop-start", "loop-stage", "loop-usage-add", "loop-control", "loop-status"]},
        {"id": "coordination_policy", "status": "ready", "commands": ["coordination-plan"]},
        {"id": "agent_eval_harness", "status": "ready", "commands": ["eval-run"]},
        {"id": "task_inbox", "status": "ready", "commands": ["task-list", "task-status", "task-update"]},
        {"id": "approval_center", "status": "ready", "commands": ["approval-create", "approval-list", "approval-decision"]},
        {"id": "human_judgment_gates", "status": "ready", "commands": ["judgment-evaluate", "judgment-list", "judgment-decision", "judgment-resume"]},
        {"id": "collaborative_rule_intake", "status": "ready", "commands": ["rule-elicit", "rule-suggest", "rule-proposal-list", "rule-proposal-decision"]},
        {"id": "notification_center", "status": "ready", "commands": ["notification-list", "notification-mark-read"]},
        {"id": "audit_events", "status": "ready", "commands": ["audit-events"]},
        {"id": "project_knowledge", "status": "ready", "commands": ["knowledge-add", "knowledge-add-text", "rag-index", "rag-search"]},
        {"id": "knowledge_spaces", "status": "ready", "commands": ["knowledge-tree", "knowledge-folder-create", "knowledge-item-add", "knowledge-share", "knowledge-scope-resolve", "knowledge-access-check"]},
        {"id": "role_workbenches", "status": "ready", "commands": ["workbench-state"]},
        {"id": "agent_registry", "status": "ready", "commands": ["agent-list", "agent-create", "agent-status", "agent-delete", "agent-plugin-report", "agent-plugin-decision", "agent-plugin-activate"]},
        {"id": "agent_lifecycle", "status": "ready", "commands": ["agent-list", "agent-create", "agent-status", "agent-delete"]},
        {"id": "digital_employee_registry", "status": "ready", "commands": ["gui-state"]},
        {"id": "workflow_packs", "status": "ready", "commands": ["gui-state", "workflow-start"]},
        {"id": "context_envelope", "status": "ready", "commands": ["handoff-create", "handoff-ack", "checkpoint-create", "gui-state"]},
        {"id": "local_skill_installations", "status": "ready", "commands": ["install-skill-sources"]},
        {"id": "product_updates", "status": "ready", "commands": ["iteration-proposal-create", "iteration-proposal-decision", "iteration-proposal-apply"]},
        {"id": "data_sharing", "status": "ready", "commands": ["telemetry-status", "telemetry-export", "telemetry-send"]},
        {"id": "external_knowledge_sources", "status": "ready", "commands": ["knowledge-source-mount", "knowledge-access-log"]},
    ]


def build_gui_state_payload(root: Path, *, project: str = "", user: str = "", role: str = "", limit: int = 20) -> dict[str, Any]:
    limit = max(1, min(limit, 100))
    project = safe_component(project, "project id") if project else ""
    user = safe_claim(user, "user", required=False)
    role = safe_component(role, "role") if role else ""
    preferences = current_onboarding_preferences(root)
    runs = [
        run
        for run in read_run_records(root)
        if (not project or run.get("project_id") == project) and (not user or run.get("requested_by") in {"", user})
    ]
    tasks = [
        task
        for task in read_records(root / "tasks", "digital-office-task")
        if (not project or task.get("project_id") == project) and (not user or task.get("requested_by") in {"", user} or task.get("assigned_user") in {"", user})
    ]
    approvals = [
        approval
        for approval in read_records(root / "approvals", "digital-office-approval")
        if (not project or approval.get("project_id") == project) and (not user or approval.get("requested_by") in {"", user})
    ]
    judgments = [
        case
        for case in read_records(root / "judgments", "digital-office-judgment-case")
        if not project or case.get("project_id") in {"", project}
    ]
    rule_proposals = [
        proposal
        for proposal in read_records(root / "rule-proposals", "digital-office-rule-proposal")
        if not project or proposal.get("project_id") in {"", project}
    ]
    notifications = [
        notification
        for notification in read_records(root / "notifications", "digital-office-notification")
        if not user or notification.get("user_id") in {"", user}
    ]
    audit_rows = read_jsonl(root / "logs" / "audit-events.jsonl", limit=limit)
    company_entries = load_entries(root, root / "knowledge" / "company" / "entries")
    mounts = read_records(root / "knowledge" / "mounts")
    checks = health_checks(root)
    agents = agent_summaries(root)
    digital_employees = digital_employee_summaries(root)
    workflow_packs = workflow_pack_summaries(root)
    context_contract = context_contract_summary(root)
    loop_runtime = loop_runtime_summary(root)
    skill_installations = skill_installation_summaries(root)
    model_runtime = model_runtime_summary(root)
    projects = project_summaries(root, 1000)
    employee_performance = employee_performance_summaries(digital_employees, runs, tasks)
    employee_suggestions = employee_gap_suggestions(digital_employees, runs, tasks)
    project_knowledge_entries: dict[str, dict[str, Any]] = {}
    for item in projects:
        project_id = str(item.get("project_id", ""))
        if not project_id:
            continue
        entries = load_entries(root, project_path(root, project_id) / "knowledge" / "entries")
        project_knowledge_entries[project_id] = {
            "count": len(entries),
            "items": compact_records(entries, ["entry_id", "title", "kind", "status", "created_at", "source_file"], min(limit, 12)),
        }
    active_workflows = [run for run in runs if run.get("status") not in {"completed", "cancelled", "stopped"}]
    draft_revision_count = sum(1 for run in runs for revision in run.get("revisions", []) if revision.get("status") == "draft")
    knowledge_spaces = knowledge_space_summaries(root, user=user, role=role, limit=limit)
    recent_run_runtime = [
        {
            "run_id": run.get("run_id", ""),
            "ledger_events": count_jsonl_records(run_ledger_path(root, str(run.get("run_id", "")))) if run.get("run_id") else 0,
            "checkpoints": len(run.get("checkpoints", []) or []),
            "handoffs": len(run.get("handoffs", []) or []),
            "pending_handoffs": sum(1 for item in run.get("handoffs", []) or [] if item.get("status", "pending_acceptance") in {"pending_acceptance", "needs_context"}),
            "context_id": run.get("context_id", ""),
            "task_id": run.get("task_id", ""),
            "current_stage": run.get("current_stage", ""),
            "cycle_index": run.get("control", {}).get("cycle_index", 1),
            "budget_usage": run.get("control", {}).get("usage", {}),
            "budgets": run.get("control", {}).get("budgets", {}),
            "last_control_decision": (run.get("control", {}).get("decision_history", []) or [{}])[-1],
        }
        for run in runs[:limit]
    ]
    payload = {
        "kind": "digital-office-gui-state",
        "version": "2.0.0",
        "generated_at": now_iso(),
        "scope": {"project_id": project, "user_id": user, "role": role},
        "health": {"status": "ok" if all(checks[k] for k in required_health_keys()) else "degraded", "checks": checks},
        "settings": {
            "configured": preferences is not None,
            "outputs": ONBOARDING_OUTPUTS,
            "preferences": preferences or {},
            "presets": onboarding_presets(root),
            "options_command": "settings-options",
            "update_command": "settings-update --confirmed",
        },
        "capabilities": gui_capabilities(),
        "agents": {"count": len(agents), "items": agents},
        "digital_employees": {"count": len(digital_employees), "items": digital_employees},
        "employee_performance": {"items": employee_performance, "suggestions": employee_suggestions},
        "workflow_packs": {"count": len(workflow_packs), "items": workflow_packs},
        "context_contract": context_contract,
        "loop_runtime": loop_runtime,
        "model_runtime": model_runtime,
        "skill_installations": {
            "count": len(skill_installations),
            "by_status": status_counts(skill_installations),
            "items": skill_installations,
        },
        "projects": {"count": len(projects), "items": projects[:limit]},
        "workflows": {
            "count": len(runs),
            "active_count": len(active_workflows),
            "draft_revision_count": draft_revision_count,
            "by_status": status_counts(runs),
            "recent": compact_records(runs, ["run_id", "title", "status", "project_id", "agent_id", "workflow", "invocation_mode", "active_revision_id", "requested_by", "created_at", "updated_at"], limit),
        },
        "tasks": {
            "count": len(tasks),
            "by_status": status_counts(tasks),
            "recent": compact_records(tasks, ["task_id", "title", "status", "priority", "project_id", "assigned_agent", "assigned_user", "workflow_run_id", "updated_at"], limit),
        },
        "approvals": {
            "count": len(approvals),
            "by_status": status_counts(approvals),
            "recent": compact_records(approvals, ["approval_id", "title", "status", "approver_role", "project_id", "workflow_run_id", "task_id", "updated_at"], limit),
        },
        "judgments": {
            "count": len(judgments),
            "by_status": status_counts(judgments),
            "recent": compact_records(judgments, ["case_id", "status", "risk_label", "required_human_role", "workflow_run_id", "task_id", "updated_at"], limit),
        },
        "rule_proposals": {
            "count": len(rule_proposals),
            "by_status": status_counts(rule_proposals),
            "recent": compact_records(rule_proposals, ["proposal_id", "title", "status", "proposed_scope", "scope_confidence", "project_id", "agent_id", "updated_at"], limit),
        },
        "notifications": {
            "count": len(notifications),
            "unread": sum(1 for item in notifications if item.get("status") == "unread"),
            "by_status": status_counts(notifications),
            "recent": compact_records(notifications, ["notification_id", "title", "topic", "status", "severity", "resource_type", "resource_id", "created_at"], limit),
        },
        "knowledge": {
            "company_entries": len(company_entries),
            "project_entries": project_knowledge_entries,
            "external_mounts": len(mounts),
            "spaces": {"count": len(knowledge_spaces), "items": knowledge_spaces},
            "rag_index_configured": (root / "rag.pipeline.json").exists(),
        },
        "runtime_replay": {
            "coordination_policy": (root / "coordination.policy.json").exists(),
            "eval_suites": [path.stem for path in sorted((root / "evals").glob("*.json"))[:limit]],
            "recent_runs": recent_run_runtime,
        },
        "workbench": {
            "command": "workbench-state",
            "entry_points": ["owner_global", "project_lead", "member", "approver", "admin", "viewer"],
        },
        "audit": {"recent": audit_rows[-limit:]},
        "next_gui_actions": [
            "Show settings status and prompt for configuration when settings.configured is false.",
            "Surface pending approvals, unread notifications, blocked workflows, and waiting tasks on the home screen.",
            "Use agent-invoke for explicit @Agent dispatch and keep every dispatch linked to a workflow run and task.",
            "Show open judgment cases before workflow actions; resume only after judgment-decision closes the case.",
            "Use rule-elicit during collaboration and rule-suggest when the user states a durable operating rule.",
            "Use workflow draft revisions for canvas edits; activate only after validation and explicit confirmation.",
            "Render Context, Decide, Act, and Evaluate from loop_runtime; expose controller decisions and budget usage without inventing client-side state transitions.",
            "Treat pending_acceptance and needs_context handoffs as incomplete; call handoff-ack only after recipient and context hash verification.",
            "Resolve knowledge folders through knowledge-scope-resolve before running Agent nodes.",
            "Route all mutating actions through the matching command and require explicit GUI confirmation where commands expose --confirmed.",
        ],
    }
    return payload


def gui_state(args: argparse.Namespace) -> int:
    root = system_root()
    payload = build_gui_state_payload(root, project=args.project or "", user=args.user or "", role=args.role or "", limit=args.limit)
    print(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def web_static_root(root: Path) -> Path:
    return root / "web" / "app"


def web_app_config(root: Path, *, public_url: str = "", user: str = "", role: str = "", project: str = "", tenant: str = "", deployment: str = "") -> dict[str, Any]:
    checks = health_checks(root)
    return {
        "kind": "digital-office-web-app-config",
        "version": "1.0.0",
        "generated_at": now_iso(),
        "app": {
            "name": "Digital Office",
            "short_name": "Office",
            "mode": "web_ui_pwa",
            "public_url": safe_claim(public_url, "public url", required=False),
            "static_root": str(web_static_root(root)),
            "manifest": "/manifest.webmanifest",
            "service_worker": "/service-worker.js",
            "start_url": "/",
        },
        "default_scope": {
            "tenant_id": safe_claim(tenant, "tenant id", required=False),
            "deployment_id": safe_claim(deployment, "deployment id", required=False),
            "user_id": safe_claim(user, "user id", required=False),
            "role": safe_component(role, "role") if role else "",
            "project_id": safe_component(project, "project id") if project else "",
        },
        "api": {
            "read_routes": [
                {"method": "GET", "path": "/api/health", "description": "Return office-system health checks and PWA readiness."},
                {"method": "GET", "path": "/api/gui-state", "description": "Return the home-screen GUI state snapshot."},
                {"method": "GET", "path": "/api/web-app", "description": "Return this Web/PWA configuration contract."},
            ],
            "mutation_routes": [
                {"method": "POST", "path": "/api/workflows", "description": "Create a governed workflow through the secretary."},
                {"method": "POST", "path": "/api/projects", "description": "Create a project folder for task context, knowledge, work records, and deliverables."},
                {"method": "POST", "path": "/api/projects/{project_id}/intent/confirm", "description": "Confirm the secretary's versioned understanding of user intent."},
                {"method": "POST", "path": "/api/projects/{project_id}/context", "description": "Update the governed project context intake."},
                {"method": "POST", "path": "/api/projects/{project_id}/context/confirm", "description": "Confirm the current project context hash before execution."},
                {"method": "POST", "path": "/api/knowledge/uploads", "description": "Upload text or files into company or project knowledge."},
                {"method": "POST", "path": "/api/settings", "description": "Persist user-facing secretary and work preferences."},
                {"method": "POST", "path": "/api/model-connections/{provider_id}", "description": "Configure a model API or Token Plan connection without returning its secret."},
                {"method": "POST", "path": "/api/model-connections/{provider_id}/test", "description": "Run a real bounded connection test."},
                {"method": "POST", "path": "/api/model-runtime", "description": "Configure local/API automatic selection policy."},
                {"method": "DELETE", "path": "/api/model-connections/{provider_id}?confirmed=true", "description": "Disconnect a model provider and remove its locally stored secret."},
                {"method": "POST", "path": "/api/agents", "description": "Create a template-based custom digital employee."},
                {"method": "POST", "path": "/api/agents/{agent_id}/status", "description": "Activate, disable, archive, or restore a custom Agent."},
                {"method": "DELETE", "path": "/api/agents/{agent_id}?confirmed=true", "description": "Permanently remove an archived custom Agent while preserving history."},
                {"method": "POST", "path": "/api/approvals/{approval_id}/decision", "description": "Approve or reject a pending approval."},
            ],
            "authentication": "Bearer token from DIGITAL_OFFICE_WEB_TOKEN is required for every non-loopback binding. Loopback development may omit the token.",
            "mutation_policy": "Mutating GUI actions use narrow dedicated API routes backed by governed office-system commands. No generic remote shell or arbitrary command endpoint is exposed.",
        },
        "pwa": {
            "installable_shell": checks.get("pwa_manifest") and checks.get("service_worker") and checks.get("web_index"),
            "offline_shell": True,
            "cache_name": "digital-office-shell-v2",
        },
        "deployment": {
            "recommended": "Serve behind HTTPS reverse proxy such as Caddy or Nginx. Bind web-serve to 127.0.0.1 for reverse proxy deployments, or to 0.0.0.0 only on trusted LAN/VPN networks.",
            "examples": ["agent-system/deploy/Caddyfile.example", "agent-system/deploy/nginx.conf.example", "agent-system/deploy/systemd/digital-office-web.service.example"],
        },
        "health": {"status": "ok" if all(checks[key] for key in required_health_keys()) else "degraded", "checks": checks},
    }


def web_config(args: argparse.Namespace) -> int:
    root = system_root()
    payload = web_app_config(root, public_url=args.public_url or "", user=args.user or "", role=args.role or "", project=args.project or "", tenant=args.tenant or "", deployment=args.deployment or "")
    print(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True))
    return 0 if payload["health"]["status"] == "ok" else 1


def web_json_response(handler: http.server.BaseHTTPRequestHandler, status: int, payload: dict[str, Any], headers: dict[str, str] | None = None) -> None:
    body = json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(body)))
    for name, value in (headers or {}).items():
        handler.send_header(name, value)
    handler.end_headers()
    handler.wfile.write(body)


def run_office_json_command(root: Path, command: list[str]) -> tuple[int, dict[str, Any]]:
    env = os.environ.copy()
    env["DIGITAL_OFFICE_SYSTEM_HOME"] = str(root)
    try:
        completed = subprocess.run(
            [sys.executable, str(Path(__file__).resolve()), *command],
            capture_output=True,
            text=True,
            env=env,
            timeout=120,
            check=False,
        )
    except subprocess.TimeoutExpired:
        return 1, {
            "status": "error",
            "error": "The requested action timed out.",
            "next_actions": ["Check runtime health and retry the action."],
        }
    output = completed.stdout.strip()
    if output:
        try:
            payload = json.loads(output)
        except json.JSONDecodeError:
            payload = {"status": "error", "summary": output[-2000:]}
    else:
        payload = {
            "status": "error" if completed.returncode else "success",
            "summary": completed.stderr.strip() or "Command completed without output.",
        }
    if completed.returncode:
        payload.setdefault("status", "error")
        payload.setdefault("error", completed.stderr.strip() or "The requested action failed.")
        payload.setdefault("next_actions", ["Review the request and retry."])
    return completed.returncode, payload


def run_model_gateway_json(root: Path, command: list[str], *, secret_input: str = "") -> tuple[int, dict[str, Any]]:
    gateway = root / "bin" / "model-gateway"
    env = os.environ.copy()
    env["DIGITAL_OFFICE_SYSTEM_HOME"] = str(root)
    try:
        completed = subprocess.run([str(gateway), *command], input=secret_input, capture_output=True, text=True, timeout=45, env=env)
    except (OSError, subprocess.TimeoutExpired) as exc:
        return 2, {"status": "error", "error": "model_gateway_unavailable", "message": str(exc)}
    raw = completed.stdout.strip() if completed.returncode == 0 else (completed.stderr.strip() or completed.stdout.strip())
    try:
        payload = json.loads(raw) if raw else {"status": "error", "error": "empty_gateway_response"}
    except json.JSONDecodeError:
        payload = {"status": "error", "error": "invalid_gateway_response", "message": raw[-1000:]}
    return completed.returncode, payload


def web_serve(args: argparse.Namespace) -> int:
    root = system_root()
    static_root = Path(args.static_dir).expanduser() if args.static_dir else web_static_root(root)
    static_root = static_root.resolve()
    if not static_root.exists():
        print(f"office-system: web static directory not found: {static_root}", file=sys.stderr)
        return 2
    public_url = args.public_url or ""
    default_scope = {
        "tenant": args.tenant or "",
        "deployment": args.deployment or "",
        "user": args.user or "",
        "role": args.role or "",
        "project": args.project or "",
    }
    allow_origin = args.allow_origin or ""
    auth_env = safe_component(args.auth_token_env, "auth token environment variable")
    auth_token = os.environ.get(auth_env, "")
    try:
        loopback_host = ipaddress.ip_address(args.host).is_loopback
    except ValueError:
        loopback_host = args.host.lower() == "localhost"
    if not loopback_host and not auth_token:
        print(f"office-system: non-loopback web binding requires a Bearer token in ${auth_env}", file=sys.stderr)
        return 2
    if auth_token and allow_origin == "*":
        print("office-system: authenticated web API cannot use wildcard --allow-origin", file=sys.stderr)
        return 2

    class DigitalOfficeHandler(http.server.SimpleHTTPRequestHandler):
        server_version = "DigitalOfficeWeb/1.0"
        extensions_map = {**http.server.SimpleHTTPRequestHandler.extensions_map, ".webmanifest": "application/manifest+json", ".js": "text/javascript; charset=utf-8", ".css": "text/css; charset=utf-8", ".svg": "image/svg+xml"}

        def log_message(self, format: str, *values: Any) -> None:
            if not args.quiet:
                super().log_message(format, *values)

        def end_headers(self) -> None:
            self.send_header("X-Content-Type-Options", "nosniff")
            self.send_header("Referrer-Policy", "same-origin")
            self.send_header("X-Frame-Options", "DENY")
            self.send_header("Permissions-Policy", "camera=(), microphone=(), geolocation=()")
            self.send_header("Content-Security-Policy", "default-src 'self'; connect-src 'self'; img-src 'self' data:; style-src 'self' 'unsafe-inline'; script-src 'self'; object-src 'none'; base-uri 'self'; frame-ancestors 'none'")
            if allow_origin:
                self.send_header("Access-Control-Allow-Origin", allow_origin)
                self.send_header("Vary", "Origin")
            super().end_headers()

        def api_authorized(self) -> bool:
            if not auth_token:
                return True
            authorization = self.headers.get("Authorization", "")
            prefix = "Bearer "
            supplied = authorization[len(prefix):] if authorization.startswith(prefix) else ""
            return bool(supplied) and hmac.compare_digest(supplied, auth_token)

        def do_OPTIONS(self) -> None:
            if not allow_origin:
                self.send_error(405)
                return
            self.send_response(204)
            self.send_header("Access-Control-Allow-Methods", "GET, POST, DELETE, OPTIONS")
            self.send_header("Access-Control-Allow-Headers", "Authorization, Content-Type")
            self.send_header("Access-Control-Max-Age", "600")
            self.end_headers()

        def request_scope(self) -> tuple[str, str, str, str]:
            tenant = default_scope["tenant"] or "local"
            deployment = default_scope["deployment"] or "local"
            user = default_scope["user"] or "local-user"
            role = default_scope["role"] or ("owner" if loopback_host else "")
            return tenant, deployment, user, role

        def read_json_body(self) -> dict[str, Any]:
            try:
                length = int(self.headers.get("Content-Length", "0"))
            except ValueError as exc:
                raise ValueError("Invalid Content-Length header.") from exc
            if length <= 0 or length > 10_485_760:
                raise ValueError("Request body must be JSON and no larger than 10 MB.")
            try:
                payload = json.loads(self.rfile.read(length).decode("utf-8"))
            except (UnicodeDecodeError, json.JSONDecodeError) as exc:
                raise ValueError("Request body must contain valid UTF-8 JSON.") from exc
            if not isinstance(payload, dict):
                raise ValueError("Request body must be a JSON object.")
            return payload

        def command_response(self, command: list[str], *, success_status: int = 200) -> None:
            code, payload = run_office_json_command(root, command)
            status = success_status if code == 0 else 403 if code == 3 else 409 if code == 4 else 400
            web_json_response(self, status, payload, headers={"Cache-Control": "no-store"})

        def knowledge_upload_response(self, body: dict[str, Any], *, user: str, role: str) -> None:
            scope = str(body.get("scope", "")).strip()
            if scope not in {"company", "project"}:
                web_json_response(self, 400, {"status": "error", "error": "scope must be company or project"})
                return
            project_id = str(body.get("project_id", "")).strip()
            if scope == "project" and not project_id:
                web_json_response(self, 400, {"status": "error", "error": "project_id is required for project uploads"})
                return
            title = str(body.get("title", "")).strip()
            if not title:
                web_json_response(self, 400, {"status": "error", "error": "title is required"})
                return
            approve = bool(body.get("approve", False))
            notes = str(body.get("notes", "")).strip() or f"Uploaded from Web by {user}."
            command: list[str]
            temporary_file: Path | None = None
            try:
                content_base64 = str(body.get("content_base64", "")).strip()
                text_body = body.get("body")
                if content_base64:
                    try:
                        decoded = base64.b64decode(content_base64, validate=True)
                    except (binascii.Error, ValueError):
                        web_json_response(self, 400, {"status": "error", "error": "content_base64 must be valid base64"})
                        return
                    if not decoded or len(decoded) > 7_340_032:
                        web_json_response(self, 400, {"status": "error", "error": "uploaded file must be between 1 byte and 7 MB"})
                        return
                    raw_filename = Path(str(body.get("filename") or f"{title}.bin")).name
                    suffix = Path(raw_filename).suffix.lower()
                    if not re.fullmatch(r"\.[a-z0-9]{1,12}", suffix):
                        suffix = ".bin"
                    safe_filename = f"{slugify(Path(raw_filename).stem or title)[:90]}{suffix}"
                    upload_dir = root / "tmp" / "web-uploads" / dt.datetime.now().strftime("%Y%m%d")
                    upload_dir.mkdir(parents=True, exist_ok=True)
                    temporary_file = upload_dir / f"{uuid.uuid4().hex}-{safe_filename}"
                    temporary_file.write_bytes(decoded)
                    command = ["knowledge-add", "--scope", scope, "--file", str(temporary_file), "--title", title, "--notes", notes]
                    kind = str(body.get("kind", "")).strip()
                    if kind in {"text", "word", "pdf", "image", "binary"}:
                        command.extend(["--kind", kind])
                elif isinstance(text_body, str) and text_body.strip():
                    command = ["knowledge-add-text", "--scope", scope, "--title", title, "--body", text_body]
                else:
                    web_json_response(self, 400, {"status": "error", "error": "body or content_base64 is required"})
                    return
                if scope == "project":
                    command.extend(["--project", project_id])
                if approve:
                    command.append("--approve")
                code, payload = run_office_json_command(root, command)
                if code == 0:
                    payload = {
                        "status": "uploaded",
                        "entry": payload,
                        "review_status": payload.get("status", "pending_review"),
                        "message": "Knowledge was added and is ready for review." if not approve else "Knowledge was approved for agent use.",
                    }
                    append_audit_event(
                        root,
                        "knowledge_uploaded_from_web",
                        actor_id=user,
                        actor_role=role,
                        project_id=project_id,
                        resource_type="knowledge_entry",
                        resource_id=str(payload.get("entry", {}).get("entry_id", "")),
                        outcome="created",
                        reason=title,
                        extra={"scope": scope, "approved_for_agent_use": approve},
                    )
                status = 201 if code == 0 else 403 if code == 3 else 409 if code == 4 else 400
                web_json_response(self, status, payload, headers={"Cache-Control": "no-store"})
            finally:
                if temporary_file is not None:
                    try:
                        temporary_file.unlink(missing_ok=True)
                    except OSError:
                        pass

        def do_POST(self) -> None:
            parsed = urllib.parse.urlparse(self.path)
            if not parsed.path.startswith("/api/"):
                web_json_response(self, 404, {"status": "not_found", "path": parsed.path})
                return
            if not self.api_authorized():
                web_json_response(self, 401, {"status": "unauthorized"}, headers={"WWW-Authenticate": "Bearer", "Cache-Control": "no-store"})
                return
            try:
                body = self.read_json_body()
            except ValueError as exc:
                web_json_response(self, 400, {"status": "error", "error": str(exc)})
                return
            tenant, deployment, user, role = self.request_scope()
            if not role:
                web_json_response(self, 403, {"status": "denied", "error": "The Web server must be started with an authorized --role for mutating actions."})
                return

            if parsed.path == "/api/workflows":
                task = str(body.get("task", "")).strip()
                if not task:
                    web_json_response(self, 400, {"status": "error", "error": "task is required"})
                    return
                generated_title = str(body.get("title", "")).strip() or (task[:48] + ("..." if len(task) > 48 else ""))
                command = ["workflow-start", "--tenant", tenant, "--deployment", deployment, "--user", user, "--role", role, "--task", task, "--title", generated_title, "--priority", str(body.get("priority", "normal"))]
                for field, flag in (("project_id", "--project"), ("agent_id", "--agent"), ("workflow", "--workflow"), ("idempotency_key", "--idempotency-key")):
                    value = str(body.get(field, "")).strip()
                    if value:
                        command.extend([flag, value])
                self.command_response(command, success_status=201)
                return

            if parsed.path == "/api/projects":
                name = str(body.get("name", "")).strip()
                if not name:
                    web_json_response(self, 400, {"status": "error", "error": "name is required"})
                    return
                command = ["project-create", "--name", name, "--guided-intake"]
                brief = str(body.get("brief", "")).strip()
                if brief:
                    command.extend(["--brief", brief])
                project_id = str(body.get("project_id", "")).strip()
                if project_id:
                    command.extend(["--project", project_id])
                agents = body.get("agent_roster")
                if isinstance(agents, list):
                    clean_agents = [str(item).strip() for item in agents if str(item).strip()]
                    if clean_agents:
                        command.extend(["--agents", ",".join(clean_agents)])
                schedule = str(body.get("methodology_schedule", "")).strip()
                if schedule in {"manual", "weekly", "monthly", "on_project_close"}:
                    command.extend(["--methodology-schedule", schedule])
                self.command_response(command, success_status=201)
                return

            project_context_match = re.fullmatch(r"/api/projects/([^/]+)/context", parsed.path)
            if project_context_match:
                project_id = urllib.parse.unquote(project_context_match.group(1))
                context = body.get("context")
                if not isinstance(context, dict):
                    web_json_response(self, 400, {"status": "error", "error": "context must be an object"})
                    return
                self.command_response(["project-context-update", "--project", project_id, "--context-json", json.dumps(context, ensure_ascii=False), "--updated-by", user])
                return

            project_confirm_match = re.fullmatch(r"/api/projects/([^/]+)/context/confirm", parsed.path)
            if project_confirm_match:
                if not bool(body.get("confirmed", False)):
                    web_json_response(self, 400, {"status": "error", "error": "confirmed=true is required"})
                    return
                project_id = urllib.parse.unquote(project_confirm_match.group(1))
                self.command_response(["project-context-confirm", "--project", project_id, "--confirmed-by", user, "--confirmed"])
                return

            project_intent_match = re.fullmatch(r"/api/projects/([^/]+)/intent/confirm", parsed.path)
            if project_intent_match:
                if not bool(body.get("confirmed", False)):
                    web_json_response(self, 400, {"status": "error", "error": "confirmed=true is required"})
                    return
                project_id = urllib.parse.unquote(project_intent_match.group(1))
                command = ["project-intent-confirm", "--project", project_id, "--confirmed-by", user, "--confirmed"]
                expected_hash = str(body.get("expected_hash", "")).strip()
                if expected_hash:
                    command.extend(["--expected-hash", expected_hash])
                self.command_response(command)
                return

            if parsed.path == "/api/knowledge/uploads":
                self.knowledge_upload_response(body, user=user, role=role)
                return

            if parsed.path == "/api/settings":
                command = ["settings-update", "--tenant", tenant, "--user", user, "--confirmed"]
                choices = body.get("choices", {}) if isinstance(body.get("choices"), dict) else {}
                for field in ONBOARDING_FIELDS:
                    value = str(choices.get(field, "")).strip()
                    if value:
                        command.extend([f"--{field.replace('_', '-')}", value])
                for field, flag in (("company_name", "--company-name"), ("secretary_name", "--secretary-name"), ("tone_note", "--tone-note")):
                    if field in body:
                        command.extend([flag, str(body.get(field, ""))])
                self.command_response(command)
                return

            model_connection_match = re.fullmatch(r"/api/model-connections/([^/]+)", parsed.path)
            if model_connection_match:
                if role not in AGENT_ADMIN_ROLES:
                    web_json_response(self, 403, {"status": "denied", "error": "Only an administrator can change model credentials."})
                    return
                provider_id = urllib.parse.unquote(model_connection_match.group(1))
                command = ["connection-set", "--provider", provider_id, "--confirmed"]
                for field, flag in (("base_url", "--base-url"), ("model", "--model"), ("protocol", "--protocol")):
                    value = str(body.get(field, "")).strip()
                    if value:
                        command.extend([flag, value])
                if body.get("enabled") is False:
                    command.append("--disabled")
                secret = str(body.get("secret", "")).strip()
                if secret:
                    command.append("--secret-stdin")
                code, payload = run_model_gateway_json(root, command, secret_input=secret)
                if code == 0:
                    append_audit_event(root, "model_connection_updated", actor_id=user, actor_role=role, resource_type="model_provider", resource_id=provider_id, outcome="configured", extra={"secret_changed": bool(secret)})
                web_json_response(self, 200 if code == 0 else 409 if code == 4 else 400, payload, headers={"Cache-Control": "no-store"})
                return

            model_test_match = re.fullmatch(r"/api/model-connections/([^/]+)/test", parsed.path)
            if model_test_match:
                if role not in AGENT_ADMIN_ROLES:
                    web_json_response(self, 403, {"status": "denied", "error": "Only an administrator can test model credentials."})
                    return
                provider_id = urllib.parse.unquote(model_test_match.group(1))
                code, payload = run_model_gateway_json(root, ["test", "--provider", provider_id])
                append_audit_event(root, "model_connection_tested", actor_id=user, actor_role=role, resource_type="model_provider", resource_id=provider_id, outcome="passed" if code == 0 else "failed")
                web_json_response(self, 200 if code == 0 else 409, payload, headers={"Cache-Control": "no-store"})
                return

            if parsed.path == "/api/model-runtime":
                if role not in AGENT_ADMIN_ROLES:
                    web_json_response(self, 403, {"status": "denied", "error": "Only an administrator can change model routing."})
                    return
                command = ["runtime-set", "--confirmed"]
                default_mode = str(body.get("default_mode", "")).strip()
                selection_policy = str(body.get("selection_policy", "")).strip()
                if default_mode:
                    command.extend(["--default-mode", default_mode])
                if selection_policy:
                    command.extend(["--selection-policy", selection_policy])
                order = body.get("provider_order")
                if isinstance(order, list):
                    command.extend(["--provider-order", ",".join(str(item) for item in order)])
                code, payload = run_model_gateway_json(root, command)
                if code == 0:
                    append_audit_event(root, "model_runtime_updated", actor_id=user, actor_role=role, resource_type="model_runtime", resource_id="default", outcome="configured")
                web_json_response(self, 200 if code == 0 else 400, payload, headers={"Cache-Control": "no-store"})
                return

            if parsed.path == "/api/agents":
                command = [
                    "agent-create",
                    "--agent", str(body.get("agent_id", "")),
                    "--display-name", str(body.get("display_name", "")),
                    "--role-description", str(body.get("role_description", "")),
                    "--template", str(body.get("template_agent_id", "")),
                    "--requested-by", user,
                    "--role", role,
                    "--confirmed",
                ]
                for value in body.get("skills", []) if isinstance(body.get("skills", []), list) else []:
                    command.extend(["--skill", str(value)])
                for value in body.get("keywords", []) if isinstance(body.get("keywords", []), list) else []:
                    command.extend(["--keyword", str(value)])
                for value in body.get("workflow_packs", []) if isinstance(body.get("workflow_packs", []), list) else []:
                    command.extend(["--workflow-pack", str(value)])
                self.command_response(command, success_status=201)
                return

            status_match = re.fullmatch(r"/api/agents/([^/]+)/status", parsed.path)
            if status_match:
                command = ["agent-status", "--agent", urllib.parse.unquote(status_match.group(1)), "--status", str(body.get("status", "")), "--requested-by", user, "--role", role, "--confirmed"]
                if body.get("reason"):
                    command.extend(["--reason", str(body["reason"])])
                self.command_response(command)
                return

            approval_match = re.fullmatch(r"/api/approvals/([^/]+)/decision", parsed.path)
            if approval_match:
                command = ["approval-decision", "--approval-id", urllib.parse.unquote(approval_match.group(1)), "--decision", str(body.get("decision", "")), "--decided-by", user, "--role", role, "--confirmed"]
                if body.get("message"):
                    command.extend(["--message", str(body["message"])])
                self.command_response(command)
                return

            web_json_response(self, 404, {"status": "not_found", "path": parsed.path})

        def do_DELETE(self) -> None:
            parsed = urllib.parse.urlparse(self.path)
            if not parsed.path.startswith("/api/"):
                web_json_response(self, 404, {"status": "not_found", "path": parsed.path})
                return
            if not self.api_authorized():
                web_json_response(self, 401, {"status": "unauthorized"}, headers={"WWW-Authenticate": "Bearer", "Cache-Control": "no-store"})
                return
            tenant, deployment, user, role = self.request_scope()
            del tenant, deployment
            if not role:
                web_json_response(self, 403, {"status": "denied", "error": "The Web server must be started with an authorized --role for mutating actions."})
                return
            agent_match = re.fullmatch(r"/api/agents/([^/]+)", parsed.path)
            query = urllib.parse.parse_qs(parsed.query)
            if agent_match and query.get("confirmed", [""])[0].lower() == "true":
                self.command_response(["agent-delete", "--agent", urllib.parse.unquote(agent_match.group(1)), "--requested-by", user, "--role", role, "--confirmed"])
                return
            model_match = re.fullmatch(r"/api/model-connections/([^/]+)", parsed.path)
            if model_match and query.get("confirmed", [""])[0].lower() == "true":
                if role not in AGENT_ADMIN_ROLES:
                    web_json_response(self, 403, {"status": "denied", "error": "Only an administrator can disconnect model credentials."})
                    return
                provider_id = urllib.parse.unquote(model_match.group(1))
                code, payload = run_model_gateway_json(root, ["connection-delete", "--provider", provider_id, "--confirmed"])
                if code == 0:
                    append_audit_event(root, "model_connection_deleted", actor_id=user, actor_role=role, resource_type="model_provider", resource_id=provider_id, outcome="disconnected")
                web_json_response(self, 200 if code == 0 else 400, payload, headers={"Cache-Control": "no-store"})
                return
            web_json_response(self, 400 if agent_match else 404, {"status": "error", "error": "Permanent deletion requires confirmed=true."} if agent_match else {"status": "not_found", "path": parsed.path})

        def translate_path(self, path: str) -> str:
            parsed = urllib.parse.urlparse(path)
            request_path = urllib.parse.unquote(parsed.path)
            if request_path in {"", "/"}:
                request_path = "/index.html"
            candidate = (static_root / request_path.lstrip("/")).resolve()
            try:
                if not candidate.is_relative_to(static_root):
                    return str(static_root / "index.html")
            except AttributeError:
                if os.path.commonpath([str(candidate), str(static_root)]) != str(static_root):
                    return str(static_root / "index.html")
            if candidate.is_dir():
                candidate = candidate / "index.html"
            if not candidate.exists() and "." not in candidate.name:
                candidate = static_root / "index.html"
            return str(candidate)

        def do_GET(self) -> None:
            parsed = urllib.parse.urlparse(self.path)
            query = urllib.parse.parse_qs(parsed.query)

            def one(name: str, default: str = "") -> str:
                values = query.get(name)
                return values[0] if values else default

            if parsed.path == "/healthz":
                detail = detailed_health(root)
                code = 200 if detail["status"] == "ok" else 503 if detail["status"] == "down" else 200
                web_json_response(self, code, {"status": detail["status"], "timestamp": detail["timestamp"]}, headers={"Cache-Control": "no-store"})
                return
            if parsed.path.startswith("/api/") and not self.api_authorized():
                web_json_response(self, 401, {"status": "unauthorized"}, headers={"WWW-Authenticate": "Bearer", "Cache-Control": "no-store"})
                return
            if parsed.path == "/api/health":
                detail = detailed_health(root)
                code = 200 if detail["status"] == "ok" else 503 if detail["status"] == "down" else 200
                web_json_response(self, code, detail, headers={"Cache-Control": "no-store"})
                return
            if parsed.path == "/api/gui-state":
                try:
                    payload = build_gui_state_payload(root, project=one("project", default_scope["project"]), user=one("user", default_scope["user"]), role=one("role", default_scope["role"]), limit=int(one("limit", "20")))
                except (SystemExit, ValueError) as exc:
                    web_json_response(self, 400, {"status": "error", "error": str(exc)})
                    return
                web_json_response(self, 200, payload, headers={"Cache-Control": "no-store"})
                return
            if parsed.path == "/api/web-app":
                payload = web_app_config(root, public_url=public_url, user=one("user", default_scope["user"]), role=one("role", default_scope["role"]), project=one("project", default_scope["project"]), tenant=one("tenant", default_scope["tenant"]), deployment=one("deployment", default_scope["deployment"]))
                web_json_response(self, 200 if payload["health"]["status"] == "ok" else 503, payload, headers={"Cache-Control": "no-store"})
                return
            if parsed.path.startswith("/api/"):
                web_json_response(self, 404, {"status": "not_found", "path": parsed.path})
                return
            super().do_GET()

    server = http.server.ThreadingHTTPServer((args.host, args.port), DigitalOfficeHandler)
    print(json.dumps({"status": "serving", "host": args.host, "port": args.port, "static_root": str(static_root), "url": f"http://{args.host}:{args.port}/", "api": ["/api/health", "/api/gui-state", "/api/web-app"], "api_authentication": "bearer" if auth_token else "loopback_only_without_token", "auth_token_env": auth_env}, ensure_ascii=False, indent=2, sort_keys=True))
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\noffice-system: web server stopped", file=sys.stderr)
    finally:
        server.server_close()
    return 0


def backup(args: argparse.Namespace) -> int:
    """Create a tar.gz backup of critical user data."""
    root = system_root()
    timestamp = dt.datetime.now().strftime('%Y%m%d-%H%M%S')
    backup_name = f'digital-office-backup-{timestamp}.tar.gz'
    backup_path = Path(args.output).expanduser() if getattr(args, "output", None) else Path.cwd() / backup_name
    backup_path = backup_path.resolve()
    if os.path.commonpath([str(backup_path), str(root.resolve())]) == str(root.resolve()):
        print('backup: output must be outside agent-system so it cannot archive itself', file=sys.stderr)
        return 1
    if backup_path.exists():
        print(f'backup: output already exists: {backup_path}', file=sys.stderr)
        return 1
    backup_path.parent.mkdir(parents=True, exist_ok=True)

    critical_dirs = [
        'settings',
        'knowledge',
        'projects',
        'approvals',
        'notifications',
        'iterations',
        'tasks',
        'agent-requests',
        'runs',
        'judgments',
        'rule-proposals',
        'logs',
    ]

    existing_dirs = [d for d in critical_dirs if (root / d).exists()]
    if not existing_dirs:
        print('backup: no critical data directories found, nothing to back up', file=sys.stderr)
        return 1

    try:
        with tarfile.open(str(backup_path), 'w:gz') as tar:
            for dirname in existing_dirs:
                tar.add(str(root / dirname), arcname=dirname)
        print(str(backup_path))
        return 0
    except (OSError, tarfile.TarError) as exc:
        print(f'backup: failed to create archive: {exc}', file=sys.stderr)
        return 1

def restore(args: argparse.Namespace) -> int:
    """Restore from a tar.gz backup file."""
    if not args.confirmed and not args.dry_run:
        print('restore: --confirmed is required because restore replaces active user-data directories', file=sys.stderr)
        return 2
    backup_file = Path(args.backup_file)
    if not backup_file.exists():
        print(f'restore: backup file not found: {backup_file}', file=sys.stderr)
        return 1
    if not backup_file.is_file():
        print(f'restore: not a file: {backup_file}', file=sys.stderr)
        return 1
    if args.max_members <= 0 or args.max_uncompressed_bytes <= 0:
        print('restore: archive limits must be positive integers', file=sys.stderr)
        return 2

    expected_dirs = {'settings', 'knowledge', 'projects', 'approvals', 'notifications', 'iterations', 'tasks', 'agent-requests', 'runs', 'judgments', 'rule-proposals', 'logs'}
    safe_members: list[tarfile.TarInfo] = []
    try:
        with tarfile.open(str(backup_file), 'r:gz') as tar:
            members = tar.getmembers()
            if len(members) > args.max_members:
                print(f'restore: archive has too many members ({len(members)} > {args.max_members})', file=sys.stderr)
                return 1
            uncompressed_bytes = sum(member.size for member in members if member.isfile())
            if uncompressed_bytes > args.max_uncompressed_bytes:
                print(f'restore: archive is too large after extraction ({uncompressed_bytes} > {args.max_uncompressed_bytes} bytes)', file=sys.stderr)
                return 1
            top_dirs = set()
            for m in members:
                normalized_name = m.name.replace('\\', '/')
                parts = [part for part in normalized_name.split('/') if part not in {'', '.'}]
                if not parts or normalized_name.startswith('/') or '..' in parts:
                    print(f'restore: unsafe archive path: {m.name}', file=sys.stderr)
                    return 1
                if m.issym() or m.islnk() or m.isdev() or m.isfifo():
                    print(f'restore: unsupported archive member type: {m.name}', file=sys.stderr)
                    return 1
                if not (m.isdir() or m.isfile()):
                    print(f'restore: unsupported archive member: {m.name}', file=sys.stderr)
                    return 1
                top_dirs.add(parts[0])
                safe_members.append(m)
    except (tarfile.TarError, OSError) as exc:
        print(f'restore: invalid backup archive: {exc}', file=sys.stderr)
        return 1

    unexpected_dirs = top_dirs - expected_dirs
    if unexpected_dirs:
        print(f"restore: archive contains unsupported top-level paths: {', '.join(sorted(unexpected_dirs))}", file=sys.stderr)
        return 1
    found_dirs = top_dirs & expected_dirs
    if not found_dirs:
        print('restore: no recognized data directories found in backup', file=sys.stderr)
        return 1

    root = system_root()
    if args.dry_run:
        print(json.dumps({"status": "valid", "backup": str(backup_file.resolve()), "components": sorted(found_dirs), "members": len(safe_members)}, ensure_ascii=False, indent=2, sort_keys=True))
        return 0

    staging = root / 'tmp' / f'restore-staging-{uuid.uuid4().hex}'
    rollback = root / 'tmp' / 'restore-rollbacks' / f"{dt.datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:8]}"
    staging.mkdir(parents=True, exist_ok=False)
    rollback.mkdir(parents=True, exist_ok=False)
    try:
        with tarfile.open(str(backup_file), 'r:gz') as tar:
            for member in safe_members:
                target = (staging / member.name).resolve()
                if os.path.commonpath([str(target), str(staging.resolve())]) != str(staging.resolve()):
                    raise ValueError(f'unsafe archive path: {member.name}')
                if member.isdir():
                    target.mkdir(parents=True, exist_ok=True)
                    continue
                target.parent.mkdir(parents=True, exist_ok=True)
                source = tar.extractfile(member)
                if source is None:
                    raise ValueError(f'archive file has no content: {member.name}')
                with source, target.open('wb') as output:
                    shutil.copyfileobj(source, output)
                os.chmod(target, member.mode & 0o777)
    except (tarfile.TarError, OSError, ValueError) as exc:
        shutil.rmtree(staging, ignore_errors=True)
        print(f'restore: extraction failed: {exc}', file=sys.stderr)
        return 1

    json_errors: list[str] = []
    for json_file in staging.glob('**/*.json'):
        try:
            json.loads(json_file.read_text(encoding='utf-8'))
        except (json.JSONDecodeError, OSError) as exc:
            json_errors.append(f'{json_file.relative_to(staging)}: {exc}')
    if json_errors:
        shutil.rmtree(staging, ignore_errors=True)
        print(f'restore: archive contains {len(json_errors)} invalid JSON file(s):', file=sys.stderr)
        for err in json_errors[:10]:
            print(f'  {err}', file=sys.stderr)
        return 1

    moved_existing: list[str] = []
    installed: list[str] = []
    try:
        with JsonFileLock(root / 'tmp' / 'restore.lock'):
            for dirname in sorted(found_dirs):
                current = root / dirname
                candidate = staging / dirname
                if current.exists():
                    shutil.move(str(current), str(rollback / dirname))
                    moved_existing.append(dirname)
                shutil.move(str(candidate), str(current))
                installed.append(dirname)
    except OSError as exc:
        for dirname in reversed(installed):
            current = root / dirname
            if current.exists():
                shutil.rmtree(current, ignore_errors=True)
        for dirname in reversed(moved_existing):
            prior = rollback / dirname
            if prior.exists():
                shutil.move(str(prior), str(root / dirname))
        shutil.rmtree(staging, ignore_errors=True)
        print(f'restore: atomic replacement failed and was rolled back: {exc}', file=sys.stderr)
        return 1

    shutil.rmtree(staging, ignore_errors=True)
    print(json.dumps({"status": "restored", "components": sorted(found_dirs), "rollback_path": str(rollback)}, ensure_ascii=False, indent=2, sort_keys=True))
    return 0

def health(args: argparse.Namespace) -> int:
    root = system_root()
    checks = health_checks(root)
    print(json.dumps(checks, ensure_ascii=False, indent=2, sort_keys=True))
    required = required_health_keys()
    return 0 if all(checks[k] for k in required) else 1


def add_knowledge_space_selector(parser: argparse.ArgumentParser) -> None:
    parser.add_argument("--space-type", choices=sorted(KNOWLEDGE_SPACE_TYPES), required=True)
    parser.add_argument("--owner")
    parser.add_argument("--project")
    parser.add_argument("--team")
    parser.add_argument("--workflow-run")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(prog="office-system")
    sub = parser.add_subparsers(dest="command", required=True)

    p = sub.add_parser("project-create")
    p.add_argument("--project")
    p.add_argument("--name", required=True)
    p.add_argument("--agents", default="")
    p.add_argument("--methodology-schedule", choices=["manual", "weekly", "monthly", "on_project_close"], default="manual")
    p.add_argument("--guided-intake", action="store_true")
    p.add_argument("--brief", default="")
    p.set_defaults(func=create_project)

    p = sub.add_parser("project-context-status")
    p.add_argument("--project", required=True)
    p.set_defaults(func=project_context_status)

    p = sub.add_parser("project-context-update")
    p.add_argument("--project", required=True)
    p.add_argument("--context-json", required=True)
    p.add_argument("--updated-by", default="")
    p.set_defaults(func=project_context_update)

    p = sub.add_parser("project-context-confirm")
    p.add_argument("--project", required=True)
    p.add_argument("--confirmed-by", required=True)
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=project_context_confirm)

    p = sub.add_parser("project-intent-confirm")
    p.add_argument("--project", required=True)
    p.add_argument("--expected-hash", default="")
    p.add_argument("--confirmed-by", required=True)
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=project_intent_confirm)

    p = sub.add_parser("knowledge-add")
    p.add_argument("--scope", choices=["company", "project"], required=True)
    p.add_argument("--project")
    p.add_argument("--file", required=True)
    p.add_argument("--title")
    p.add_argument("--kind", choices=["text", "word", "pdf", "image", "binary"])
    p.add_argument("--notes")
    p.add_argument("--approve", action="store_true")
    p.set_defaults(func=add_file_entry)

    p = sub.add_parser("knowledge-add-text")
    p.add_argument("--scope", choices=["company", "project"], required=True)
    p.add_argument("--project")
    p.add_argument("--title", required=True)
    p.add_argument("--body")
    p.add_argument("--approve", action="store_true")
    p.set_defaults(func=add_text_entry)

    p = sub.add_parser("knowledge-folder-create")
    add_knowledge_space_selector(p)
    p.add_argument("--folder-id")
    p.add_argument("--parent-folder", default="root")
    p.add_argument("--title", required=True)
    p.add_argument("--created-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--replace", action="store_true")
    p.set_defaults(func=knowledge_folder_create)

    p = sub.add_parser("knowledge-item-add")
    add_knowledge_space_selector(p)
    p.add_argument("--item-id")
    p.add_argument("--folder-id", default="root")
    p.add_argument("--title", required=True)
    p.add_argument("--source-ref", required=True)
    p.add_argument("--content-type")
    p.add_argument("--connector-id")
    p.add_argument("--created-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--replace", action="store_true")
    p.set_defaults(func=knowledge_item_add)

    p = sub.add_parser("knowledge-share")
    add_knowledge_space_selector(p)
    p.add_argument("--resource-type", choices=sorted(KNOWLEDGE_RESOURCE_TYPES), required=True)
    p.add_argument("--resource-id", required=True)
    p.add_argument("--target-type", choices=sorted(KNOWLEDGE_TARGET_TYPES), required=True)
    p.add_argument("--target-id", required=True)
    p.add_argument("--permission", choices=sorted(KNOWLEDGE_PERMISSIONS), default="read")
    p.add_argument("--effect", choices=["allow", "deny"], default="allow")
    p.add_argument("--shared-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--share-id")
    p.add_argument("--reason")
    p.add_argument("--no-inherit", action="store_true")
    p.set_defaults(func=knowledge_share)

    p = sub.add_parser("knowledge-access-check")
    add_knowledge_space_selector(p)
    p.add_argument("--resource-type", choices=sorted(KNOWLEDGE_RESOURCE_TYPES), required=True)
    p.add_argument("--resource-id", required=True)
    p.add_argument("--permission", choices=sorted(KNOWLEDGE_PERMISSIONS), default="read")
    p.add_argument("--user", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--agent")
    p.set_defaults(func=knowledge_access_check)

    p = sub.add_parser("knowledge-scope-resolve")
    add_knowledge_space_selector(p)
    p.add_argument("--folder-id")
    p.add_argument("--item-id")
    p.add_argument("--user", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--agent")
    p.add_argument("--live-mode", action="store_true")
    p.set_defaults(func=knowledge_scope_resolve)

    p = sub.add_parser("knowledge-tree")
    add_knowledge_space_selector(p)
    p.add_argument("--user", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--agent")
    p.set_defaults(func=knowledge_tree)

    p = sub.add_parser("rule-add")
    p.add_argument("--scope", choices=["global", "agent", "project"], required=True)
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--title", required=True)
    p.add_argument("--body")
    p.set_defaults(func=add_rule)

    p = sub.add_parser("rule-elicit")
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--context")
    p.add_argument("--limit", type=int, default=4)
    p.set_defaults(func=rule_elicit)

    p = sub.add_parser("rule-suggest")
    p.add_argument("--title", required=True)
    p.add_argument("--body")
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--source")
    p.add_argument("--created-by")
    p.add_argument("--role")
    p.add_argument("--proposal-id")
    p.set_defaults(func=rule_suggest)

    p = sub.add_parser("rule-proposal-list")
    p.add_argument("--status", choices=sorted(RULE_PROPOSAL_STATUSES))
    p.add_argument("--scope", choices=["global", "project", "agent"])
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--limit", type=int, default=20)
    p.set_defaults(func=rule_proposal_list)

    p = sub.add_parser("rule-proposal-decision")
    p.add_argument("--proposal-id", required=True)
    p.add_argument("--decision", choices=["approve", "tune", "reject"], required=True)
    p.add_argument("--tenant", default="local")
    p.add_argument("--deployment", default="local")
    p.add_argument("--decided-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--scope", choices=["global", "project", "agent"])
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--title")
    p.add_argument("--body")
    p.add_argument("--message")
    p.add_argument("--confirmed", action="store_true")
    p.add_argument("--override-conflicts", action="store_true")
    p.set_defaults(func=rule_proposal_decision)

    p = sub.add_parser("context")
    p.add_argument("--project")
    p.add_argument("--agent")
    p.set_defaults(func=context)

    p = sub.add_parser("gui-state")
    p.add_argument("--project")
    p.add_argument("--user")
    p.add_argument("--role")
    p.add_argument("--limit", type=int, default=20)
    p.set_defaults(func=gui_state)

    p = sub.add_parser("web-config")
    p.add_argument("--public-url")
    p.add_argument("--tenant")
    p.add_argument("--deployment")
    p.add_argument("--user")
    p.add_argument("--role")
    p.add_argument("--project")
    p.set_defaults(func=web_config)

    p = sub.add_parser("web-serve")
    p.add_argument("--host", default=os.environ.get("WEB_HOST", "127.0.0.1"))
    p.add_argument("--port", type=int, default=os.environ.get("WEB_PORT", "8787"))
    p.add_argument("--static-dir")
    p.add_argument("--public-url")
    p.add_argument("--tenant")
    p.add_argument("--deployment")
    p.add_argument("--user")
    p.add_argument("--role")
    p.add_argument("--project")
    p.add_argument("--allow-origin")
    p.add_argument("--auth-token-env", default="DIGITAL_OFFICE_WEB_TOKEN")
    p.add_argument("--quiet", action="store_true")
    p.set_defaults(func=web_serve)

    p = sub.add_parser("workbench-state")
    p.add_argument("--tenant", required=True)
    p.add_argument("--deployment", required=True)
    p.add_argument("--user", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--project")
    p.add_argument("--limit", type=int, default=20)
    p.set_defaults(func=workbench_state)

    p = sub.add_parser("onboarding-options")
    p.set_defaults(func=onboarding_options)

    p = sub.add_parser("onboarding-status")
    p.set_defaults(func=onboarding_status)

    p = sub.add_parser("onboarding-apply")
    add_preferences_arguments(p)
    p.set_defaults(func=onboarding_apply, preference_source="gui_first_run_onboarding")

    p = sub.add_parser("settings-options")
    p.set_defaults(func=onboarding_options)

    p = sub.add_parser("settings-status")
    p.set_defaults(func=onboarding_status)

    p = sub.add_parser("settings-update")
    add_preferences_arguments(p)
    p.set_defaults(func=onboarding_apply, preference_source="gui_settings_update")

    p = sub.add_parser("identity-context")
    p.add_argument("--tenant", required=True)
    p.add_argument("--deployment", required=True)
    p.add_argument("--user", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--workflow-run")
    p.add_argument("--session")
    p.set_defaults(func=identity_context)

    p = sub.add_parser("auth-decision")
    p.add_argument("--tenant", required=True)
    p.add_argument("--deployment", required=True)
    p.add_argument("--user", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--action", required=True)
    p.add_argument("--resource-type", required=True)
    p.add_argument("--resource-id", required=True)
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--workflow-run")
    p.add_argument("--reason")
    p.set_defaults(func=auth_decision)

    p = sub.add_parser("workflow-start")
    p.add_argument("--tenant", required=True)
    p.add_argument("--deployment", required=True)
    p.add_argument("--user", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--task")
    p.add_argument("--title")
    p.add_argument("--project")
    p.add_argument("--agent", default="auto")
    p.add_argument("--workflow")
    p.add_argument("--priority", choices=["low", "normal", "high", "urgent"], default="normal")
    p.add_argument("--run-id")
    p.add_argument("--task-id")
    p.add_argument("--idempotency-key")
    p.add_argument("--reason")
    p.set_defaults(func=workflow_start)

    p = sub.add_parser("agent-invoke")
    p.add_argument("--tenant", required=True)
    p.add_argument("--deployment", required=True)
    p.add_argument("--user", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--project", required=True)
    p.add_argument("--agent", required=True)
    p.add_argument("--task")
    p.add_argument("--title")
    p.add_argument("--priority", choices=["low", "normal", "high", "urgent"], default="normal")
    p.add_argument("--run-id")
    p.add_argument("--task-id")
    p.add_argument("--reason")
    p.set_defaults(func=agent_invoke)

    p = sub.add_parser("workflow-status")
    p.add_argument("--run-id", required=True)
    p.set_defaults(func=workflow_status)

    p = sub.add_parser("workflow-list")
    p.add_argument("--status")
    p.add_argument("--project")
    p.add_argument("--user")
    p.add_argument("--limit", type=int, default=20)
    p.set_defaults(func=workflow_list)

    p = sub.add_parser("workflow-cancel")
    p.add_argument("--run-id", required=True)
    p.add_argument("--requested-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--reason")
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=workflow_cancel)

    p = sub.add_parser("workflow-resume")
    p.add_argument("--run-id", required=True)
    p.add_argument("--requested-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--reason")
    p.set_defaults(func=workflow_resume)

    p = sub.add_parser("workflow-retry")
    p.add_argument("--run-id", required=True)
    p.add_argument("--stage", choices=LOOP_STAGE_CHOICES)
    p.add_argument("--requested-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--reason")
    p.set_defaults(func=workflow_retry)

    p = sub.add_parser("workflow-draft-create")
    p.add_argument("--run-id", required=True)
    p.add_argument("--created-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--revision-id")
    p.add_argument("--summary")
    p.set_defaults(func=workflow_draft_create)

    p = sub.add_parser("workflow-draft-patch")
    p.add_argument("--run-id", required=True)
    p.add_argument("--revision-id", required=True)
    p.add_argument("--updated-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--patch-json")
    p.add_argument("--patch-file")
    p.add_argument("--summary")
    p.set_defaults(func=workflow_draft_patch)

    p = sub.add_parser("workflow-draft-validate")
    p.add_argument("--run-id", required=True)
    p.add_argument("--revision-id", required=True)
    p.set_defaults(func=workflow_draft_validate)

    p = sub.add_parser("workflow-draft-activate")
    p.add_argument("--run-id", required=True)
    p.add_argument("--revision-id", required=True)
    p.add_argument("--activated-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--reason")
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=workflow_draft_activate)

    p = sub.add_parser("workflow-control")
    p.add_argument("--run-id", required=True)
    p.add_argument("--action", choices=sorted(WORKFLOW_CONTROL_ACTIONS), required=True)
    p.add_argument("--requested-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--reason")
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=workflow_control)

    p = sub.add_parser("workflow-node-context")
    p.add_argument("--run-id", required=True)
    p.add_argument("--node-id", required=True)
    p.add_argument("--revision-id")
    p.set_defaults(func=workflow_node_context)

    p = sub.add_parser("task-list")
    p.add_argument("--status", choices=sorted(TASK_STATUSES))
    p.add_argument("--project")
    p.add_argument("--assigned-agent")
    p.add_argument("--assigned-user")
    p.add_argument("--limit", type=int, default=20)
    p.set_defaults(func=task_list)

    p = sub.add_parser("task-status")
    p.add_argument("--task-id", required=True)
    p.set_defaults(func=task_status)

    p = sub.add_parser("task-update")
    p.add_argument("--task-id", required=True)
    p.add_argument("--status", choices=sorted(TASK_STATUSES))
    p.add_argument("--summary")
    p.add_argument("--artifact", action="append")
    p.add_argument("--assigned-agent")
    p.add_argument("--assigned-user")
    p.add_argument("--updated-by", required=True)
    p.add_argument("--role", required=True)
    p.set_defaults(func=task_update)

    p = sub.add_parser("judgment-evaluate")
    p.add_argument("--task")
    p.add_argument("--stage", choices=LOOP_STAGE_CHOICES, default="context")
    p.add_argument("--agent")
    p.add_argument("--workflow-run")
    p.add_argument("--task-id")
    p.add_argument("--action")
    p.add_argument("--route-json")
    p.add_argument("--route-file")
    p.add_argument("--signal-json")
    p.add_argument("--signal-file")
    p.add_argument("--create-case", action="store_true")
    p.add_argument("--case-id")
    p.add_argument("--reason")
    p.add_argument("--created-by")
    p.add_argument("--role")
    p.set_defaults(func=judgment_evaluate)

    p = sub.add_parser("judgment-list")
    p.add_argument("--status", choices=sorted(JUDGMENT_STATUSES))
    p.add_argument("--workflow-run")
    p.add_argument("--required-human-role")
    p.add_argument("--limit", type=int, default=20)
    p.set_defaults(func=judgment_list)

    p = sub.add_parser("judgment-decision")
    p.add_argument("--case-id", required=True)
    p.add_argument("--decision", choices=["approve", "reject", "request_evidence", "revise_scope", "cancel"], required=True)
    p.add_argument("--tenant", default="local")
    p.add_argument("--deployment", default="local")
    p.add_argument("--decided-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--project")
    p.add_argument("--message")
    p.add_argument("--scope-note")
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=judgment_decision)

    p = sub.add_parser("judgment-resume")
    p.add_argument("--run-id", required=True)
    p.add_argument("--requested-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--reason")
    p.set_defaults(func=judgment_resume)

    p = sub.add_parser("approval-create")
    p.add_argument("--tenant", required=True)
    p.add_argument("--deployment", required=True)
    p.add_argument("--title", required=True)
    p.add_argument("--body")
    p.add_argument("--action", required=True)
    p.add_argument("--resource-type", required=True)
    p.add_argument("--resource-id", required=True)
    p.add_argument("--requested-by", required=True)
    p.add_argument("--requested-by-role", required=True)
    p.add_argument("--approver-role", default="project_manager")
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--workflow-run")
    p.add_argument("--task-id")
    p.add_argument("--risk")
    p.add_argument("--expires-at")
    p.add_argument("--approval-id")
    p.add_argument("--idempotency-key")
    p.set_defaults(func=approval_create)

    p = sub.add_parser("approval-list")
    p.add_argument("--status", choices=sorted(APPROVAL_STATUSES))
    p.add_argument("--project")
    p.add_argument("--approver-role")
    p.add_argument("--limit", type=int, default=20)
    p.set_defaults(func=approval_list)

    p = sub.add_parser("approval-decision")
    p.add_argument("--approval-id", required=True)
    p.add_argument("--decision", choices=["approve", "reject", "cancel"], required=True)
    p.add_argument("--decided-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--message")
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=approval_decision)

    p = sub.add_parser("audit-events")
    p.add_argument("--event")
    p.add_argument("--resource-type")
    p.add_argument("--resource-id")
    p.add_argument("--limit", type=int, default=50)
    p.set_defaults(func=audit_events)

    p = sub.add_parser("notification-list")
    p.add_argument("--user")
    p.add_argument("--unread-only", action="store_true")
    p.add_argument("--limit", type=int, default=20)
    p.set_defaults(func=notification_list)

    p = sub.add_parser("notification-mark-read")
    p.add_argument("--notification-id", required=True)
    p.add_argument("--user")
    p.set_defaults(func=notification_mark_read)

    p = sub.add_parser("loop-start")
    p.add_argument("--run-id")
    p.add_argument("--task")
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--workflow", default="")
    p.add_argument("--requested-by", default="")
    p.add_argument("--max-cycles", type=int)
    p.add_argument("--max-stage-retries", type=int)
    p.add_argument("--max-stagnant-cycles", type=int)
    p.add_argument("--max-duration-seconds", type=int)
    p.add_argument("--max-tool-calls", type=int)
    p.add_argument("--max-model-calls", type=int)
    p.add_argument("--max-cost-microunits", type=int)
    p.add_argument("--dry-run", action="store_true")
    p.set_defaults(func=loop_start)

    p = sub.add_parser("loop-stage")
    p.add_argument("--run-id", required=True)
    p.add_argument("--stage", choices=LOOP_STAGE_CHOICES, required=True)
    p.add_argument("--status", choices=["started", "completed", "skipped", "failed", "blocked"], required=True)
    p.add_argument("--summary")
    p.add_argument("--body")
    p.add_argument("--artifact", action="append")
    p.add_argument("--gate", action="append")
    p.set_defaults(func=loop_stage)

    p = sub.add_parser("loop-status")
    p.add_argument("--run-id", required=True)
    p.set_defaults(func=loop_status)

    p = sub.add_parser("loop-usage-add")
    p.add_argument("--run-id", required=True)
    p.add_argument("--stage", choices=LOOP_STAGE_CHOICES)
    p.add_argument("--tool-calls", type=int, default=0)
    p.add_argument("--model-calls", type=int, default=0)
    p.add_argument("--input-tokens", type=int, default=0)
    p.add_argument("--output-tokens", type=int, default=0)
    p.add_argument("--cost-microunits", type=int, default=0)
    p.set_defaults(func=loop_usage_add)

    p = sub.add_parser("loop-control")
    p.add_argument("--run-id", required=True)
    p.add_argument("--decision", choices=sorted(LOOP_CONTROL_DECISIONS), required=True)
    p.add_argument("--progress-score", type=float, required=True)
    p.add_argument("--acceptance-passed", action="store_true")
    p.add_argument("--reason")
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=loop_control)

    p = sub.add_parser("run-ledger-add")
    p.add_argument("--run-id", required=True)
    p.add_argument("--event", required=True)
    p.add_argument("--stage", choices=LOOP_STAGE_CHOICES)
    p.add_argument("--action")
    p.add_argument("--agent")
    p.add_argument("--actor")
    p.add_argument("--role")
    p.add_argument("--input-json")
    p.add_argument("--input-file")
    p.add_argument("--output-json")
    p.add_argument("--output-file")
    p.add_argument("--parameters-json")
    p.add_argument("--parameters-file")
    p.add_argument("--artifact", action="append")
    p.add_argument("--checkpoint-id")
    p.add_argument("--handoff-id")
    p.add_argument("--model")
    p.add_argument("--provider")
    p.set_defaults(func=run_ledger_add)

    p = sub.add_parser("run-ledger-list")
    p.add_argument("--run-id", required=True)
    p.add_argument("--limit", type=int, default=50)
    p.set_defaults(func=run_ledger_list)

    p = sub.add_parser("run-ledger-verify")
    p.add_argument("--run-id", required=True)
    p.set_defaults(func=run_ledger_verify)

    p = sub.add_parser("checkpoint-create")
    p.add_argument("--run-id", required=True)
    p.add_argument("--checkpoint-id")
    p.add_argument("--stage", choices=LOOP_STAGE_CHOICES)
    p.add_argument("--label")
    p.add_argument("--resume-cursor")
    p.add_argument("--state-json")
    p.add_argument("--state-file")
    p.add_argument("--artifact", action="append")
    p.add_argument("--requires-human", action="store_true")
    p.add_argument("--create-judgment", action="store_true")
    p.add_argument("--reason")
    p.add_argument("--created-by")
    p.add_argument("--role")
    p.set_defaults(func=checkpoint_create)

    p = sub.add_parser("checkpoint-list")
    p.add_argument("--run-id", required=True)
    p.add_argument("--limit", type=int, default=20)
    p.set_defaults(func=checkpoint_list)

    p = sub.add_parser("handoff-create")
    p.add_argument("--run-id", required=True)
    p.add_argument("--handoff-id")
    p.add_argument("--from-agent", required=True)
    p.add_argument("--to-agent", required=True)
    p.add_argument("--stage", choices=LOOP_STAGE_CHOICES)
    p.add_argument("--reason", required=True)
    p.add_argument("--input-schema-json")
    p.add_argument("--input-schema-file")
    p.add_argument("--context-json")
    p.add_argument("--context-file")
    p.add_argument("--artifact", action="append")
    p.add_argument("--acceptance-criterion", action="append")
    p.add_argument("--created-by")
    p.add_argument("--role")
    p.set_defaults(func=handoff_create)

    p = sub.add_parser("handoff-list")
    p.add_argument("--run-id", required=True)
    p.add_argument("--agent")
    p.add_argument("--limit", type=int, default=20)
    p.set_defaults(func=handoff_list)

    p = sub.add_parser("handoff-ack")
    p.add_argument("--run-id", required=True)
    p.add_argument("--handoff-id", required=True)
    p.add_argument("--received-by", required=True)
    p.add_argument("--decision", choices=["accept", "request_context", "reject"], required=True)
    p.add_argument("--expected-context-hash", required=True)
    p.add_argument("--missing-field", action="append")
    p.add_argument("--message")
    p.add_argument("--confirmed", action="store_true")
    p.add_argument("--created-by")
    p.add_argument("--role")
    p.set_defaults(func=handoff_ack)

    p = sub.add_parser("coordination-plan")
    p.add_argument("--task")
    p.add_argument("--run-id")
    p.add_argument("--agent", action="append")
    p.add_argument("--complexity", choices=["low", "medium", "high"])
    p.add_argument("--risk", choices=["low", "medium", "high", "critical"])
    p.add_argument("--parallelizable", action="store_true")
    p.add_argument("--sequential", action="store_true")
    p.set_defaults(func=coordination_plan)

    p = sub.add_parser("eval-run")
    p.add_argument("--suite", required=True)
    p.add_argument("--no-write", action="store_true")
    p.set_defaults(func=eval_run)

    p = sub.add_parser("iteration-proposal-create")
    p.add_argument("--title", required=True)
    p.add_argument("--target", choices=["agent", "workflow", "rule", "knowledge", "harness", "release", "model", "product"], required=True)
    p.add_argument("--summary", required=True)
    p.add_argument("--body")
    p.add_argument("--expected-impact", required=True)
    p.add_argument("--risk", required=True)
    p.add_argument("--rollback", required=True)
    p.add_argument("--run-id")
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--source-ref", action="append")
    p.add_argument("--regression-check", action="append")
    p.set_defaults(func=iteration_proposal_create)

    p = sub.add_parser("iteration-proposal-decision")
    p.add_argument("--proposal-id", required=True)
    p.add_argument("--decision", choices=["confirm", "tune", "pause", "reject"], required=True)
    p.add_argument("--message")
    p.set_defaults(func=iteration_proposal_decision)

    p = sub.add_parser("iteration-proposal-apply")
    p.add_argument("--proposal-id", required=True)
    p.add_argument("--confirmed", action="store_true")
    p.add_argument("--applied-by", default="")
    p.add_argument("--artifact", action="append")
    p.add_argument("--regression-result")
    p.add_argument("--note")
    p.set_defaults(func=iteration_proposal_apply)

    p = sub.add_parser("methodology-draft")
    p.add_argument("--project", required=True)
    p.add_argument("--period", default="manual")
    p.set_defaults(func=methodology_draft)

    p = sub.add_parser("methodology-approve")
    p.add_argument("--project", required=True)
    p.add_argument("--draft", required=True)
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=methodology_approve)

    p = sub.add_parser("relay-add")
    p.add_argument("--project", required=True)
    p.add_argument("--subproject")
    p.add_argument("--agent", required=True)
    p.add_argument("--title", required=True)
    p.add_argument("--body")
    p.add_argument("--status", choices=["active", "blocked", "done", "superseded"], default="active")
    p.add_argument("--source-ref", action="append")
    p.add_argument("--next-action", action="append")
    p.set_defaults(func=relay_add)

    p = sub.add_parser("agent-request-submit")
    p.add_argument("--title", required=True)
    p.add_argument("--body")
    p.add_argument("--project")
    p.add_argument("--requested-by")
    p.add_argument("--priority", choices=["low", "normal", "high", "urgent"], default="normal")
    p.add_argument("--confirmed", action="store_true")
    p.add_argument("--timeout", type=int, default=30)
    p.set_defaults(func=agent_request_submit)

    p = sub.add_parser("agent-request-status")
    p.add_argument("--request-id", required=True)
    p.set_defaults(func=agent_request_status)

    p = sub.add_parser("agent-request-set-status")
    p.add_argument("--request-id", required=True)
    p.add_argument(
        "--status",
        choices=["received", "in_progress", "completed", "downloaded_deployed", "pending_user_confirmation", "needs_tuning", "paused_by_user", "needs_more_info", "rejected"],
        required=True,
    )
    p.add_argument("--message")
    p.add_argument("--package")
    p.set_defaults(func=agent_request_set_status)

    p = sub.add_parser("agent-improvement-draft")
    p.add_argument("--agent", required=True)
    p.add_argument("--kind", choices=["soul", "workflow"], required=True)
    p.add_argument("--title", required=True)
    p.add_argument("--body")
    p.add_argument("--project")
    p.set_defaults(func=agent_improvement_draft)

    p = sub.add_parser("agent-improvement-approve")
    p.add_argument("--agent", required=True)
    p.add_argument("--draft", required=True)
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=agent_improvement_approve)

    p = sub.add_parser("agent-plugin-report")
    p.add_argument("--package", required=True)
    p.add_argument("--project")
    p.add_argument("--request-id")
    p.set_defaults(func=agent_plugin_report)

    p = sub.add_parser("agent-plugin-decision")
    p.add_argument("--report-id", required=True)
    p.add_argument("--decision", choices=["confirm", "tune", "pause"], required=True)
    p.add_argument("--message")
    p.set_defaults(func=agent_plugin_decision)

    p = sub.add_parser("agent-plugin-activate")
    p.add_argument("--package", required=True)
    p.add_argument("--project")
    p.add_argument("--report")
    p.add_argument("--report-id")
    p.add_argument("--request-id")
    p.add_argument("--confirmed", action="store_true")
    p.add_argument("--replace-agent", action="store_true")
    p.add_argument("--replace-profile", action="store_true")
    p.set_defaults(func=agent_plugin_activate)

    p = sub.add_parser("rag-index")
    p.add_argument("--scope", choices=["company", "project"], required=True)
    p.add_argument("--project")
    p.add_argument("--mode", choices=["auto", "lexical", "embedding"], default="auto")
    p.add_argument("--embedding-model", default="BAAI/bge-small-zh-v1.5")
    p.add_argument("--chunk-chars", type=int, default=1200)
    p.add_argument("--chunk-overlap", type=int, default=180)
    p.add_argument("--include-pending", action="store_true")
    p.set_defaults(func=rag_index)

    p = sub.add_parser("rag-search")
    p.add_argument("--scope", choices=["company", "project"], required=True)
    p.add_argument("--project")
    p.add_argument("--query", required=True)
    p.add_argument("--limit", type=int, default=5)
    p.add_argument("--embedding-model", default="BAAI/bge-small-zh-v1.5")
    p.set_defaults(func=rag_search)

    p = sub.add_parser("knowledge-source-mount")
    p.add_argument("--source-class", choices=["customer_owned_external_kb", "provider_sold_industry_kb"], required=True)
    p.add_argument("--source-id", required=True)
    p.add_argument("--display-name")
    p.add_argument("--provider", default="")
    p.add_argument("--tenant", required=True)
    p.add_argument("--deployment", required=True)
    p.add_argument("--created-by", required=True)
    p.add_argument("--mount-target", choices=["company_knowledge", "project_knowledge", "agent_specialist_context", "licensed_company_reference", "licensed_project_reference", "licensed_agent_reference"], required=True)
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--entitlement")
    p.add_argument("--license-sku")
    p.add_argument("--sync-mode", choices=["metadata_only", "selected_items", "excerpt_index", "embedding_index", "full_customer_owned_sync", "controlled_remote_retrieval"], default="excerpt_index")
    p.add_argument("--allowed-user", action="append")
    p.add_argument("--allowed-role", action="append")
    p.add_argument("--mount-id")
    p.add_argument("--replace", action="store_true")
    p.set_defaults(func=knowledge_source_mount)

    p = sub.add_parser("knowledge-access-log")
    p.add_argument("--tenant", required=True)
    p.add_argument("--deployment", required=True)
    p.add_argument("--user", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--project")
    p.add_argument("--agent")
    p.add_argument("--workflow-run")
    p.add_argument("--source-class", choices=["customer_owned_external_kb", "provider_sold_industry_kb"], required=True)
    p.add_argument("--source-id", required=True)
    p.add_argument("--mount-id", required=True)
    p.add_argument("--knowledge-pack")
    p.add_argument("--entitlement")
    p.add_argument("--query")
    p.add_argument("--result-source-id", action="append")
    p.add_argument("--snippet-count", type=int, default=0)
    p.add_argument("--decision", choices=["allow", "deny"], required=True)
    p.add_argument("--deny-reason")
    p.set_defaults(func=knowledge_access_log)

    p = sub.add_parser("telemetry-status")
    p.set_defaults(func=telemetry_status)

    p = sub.add_parser("telemetry-export")
    p.set_defaults(func=telemetry_export)

    p = sub.add_parser("telemetry-send")
    p.add_argument("--bundle", required=True)
    p.add_argument("--confirmed", action="store_true")
    p.add_argument("--timeout", type=int, default=30)
    p.set_defaults(func=telemetry_send)

    p = sub.add_parser("agent-list")
    p.add_argument("--include-archived", action="store_true")
    p.set_defaults(func=agent_lifecycle_list)

    p = sub.add_parser("agent-create")
    p.add_argument("--agent", required=True)
    p.add_argument("--display-name", required=True)
    p.add_argument("--role-description", required=True)
    p.add_argument("--template", required=True)
    p.add_argument("--skill", action="append")
    p.add_argument("--keyword", action="append")
    p.add_argument("--workflow-pack", action="append")
    p.add_argument("--requested-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=agent_lifecycle_create)

    p = sub.add_parser("agent-status")
    p.add_argument("--agent", required=True)
    p.add_argument("--status", choices=sorted(CUSTOM_AGENT_STATUSES), required=True)
    p.add_argument("--requested-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--reason")
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=agent_lifecycle_status)

    p = sub.add_parser("agent-delete")
    p.add_argument("--agent", required=True)
    p.add_argument("--requested-by", required=True)
    p.add_argument("--role", required=True)
    p.add_argument("--reason")
    p.add_argument("--confirmed", action="store_true")
    p.set_defaults(func=agent_lifecycle_delete)

    p = sub.add_parser("backup")
    p.add_argument("--output")
    p.set_defaults(func=backup)

    p = sub.add_parser("restore")
    p.add_argument("backup_file", help="Path to the backup tar.gz file")
    p.add_argument("--dry-run", action="store_true")
    p.add_argument("--confirmed", action="store_true")
    p.add_argument("--max-members", type=int, default=DEFAULT_RESTORE_MAX_MEMBERS)
    p.add_argument("--max-uncompressed-bytes", type=int, default=DEFAULT_RESTORE_MAX_UNCOMPRESSED_BYTES)
    p.set_defaults(func=restore)

    p = sub.add_parser("health")
    p.set_defaults(func=health)
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    read_only_run_commands = {
        "workflow-status",
        "workflow-list",
        "loop-status",
        "run-ledger-list",
        "run-ledger-verify",
        "checkpoint-list",
        "handoff-list",
    }
    run_id = str(getattr(args, "run_id", "") or "")
    if not run_id and args.command in {"workflow-start", "agent-invoke", "loop-start"}:
        run_id = "_create"
    if not run_id and args.command == "task-update":
        task_id = str(getattr(args, "task_id", "") or "")
        task_file = task_path(system_root(), task_id) if task_id else None
        if task_file and task_file.exists():
            run_id = str(read_json(task_file).get("workflow_run_id", ""))
    if run_id and args.command not in read_only_run_commands:
        lock_id = safe_component(run_id, "run lock id")
        lock_path = system_root() / "runs" / f".{lock_id}.command.lock"
        with JsonFileLock(lock_path):
            return args.func(args)
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
